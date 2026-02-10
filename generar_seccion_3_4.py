"""
Generador de Documentacion - Seccion 3.4: Integracion del Modelo de Machine Learning
Trabajo de Grado - Ivan Condori Choquehuanca
Genera un archivo .docx con la documentacion completa del pipeline ML
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONFIGURACION GENERAL
# ============================================================

AUTOR = "Ivan Condori Choquehuanca"
SISTEMA = "Sistema de Intermediacion Laboral con PLN"

# Contadores globales (independientes para seccion 3.4)
tabla_counter = {"value": 0}
figura_counter = {"value": 0}


# ============================================================
# FUNCIONES AUXILIARES DE FORMATO
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_table(doc, headers, rows, caption, nota, fuente="Elaboracion propia"):
    tabla_counter["value"] += 1
    num = tabla_counter["value"]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(table)

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Tabla {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    nota_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    fuente_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_figure_placeholder(doc, descripcion_captura, caption, nota, fuente="Elaboracion propia"):
    figura_counter["value"] += 1
    num = figura_counter["value"]

    p_box = doc.add_paragraph()
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_box.add_run(f"[INSERTAR CAPTURA DE PANTALLA]\n{descripcion_captura}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 0, 0)
    run.italic = True

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_code_block(doc, code_text, caption, nota, fuente="Elaboracion propia"):
    figura_counter["value"] += 1
    num = figura_counter["value"]

    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Consolas" w:hAnsi="Consolas" w:cs="Consolas"/>')
        rPr.append(rFonts)

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_paragraph_text(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_formula_block(doc, formula_text, caption, nota, fuente="Elaboracion propia"):
    """Agrega una formula matematica con formato especial."""
    figura_counter["value"] += 1
    num = figura_counter["value"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(formula_text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    run.italic = True
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>')
    rPr.append(rFonts)

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


# ============================================================
# SECCION 3.4: INTEGRACION DEL MODELO DE MACHINE LEARNING
# ============================================================

def generate_section_3_4(doc):
    """Genera toda la seccion 3.4"""

    doc.add_heading("3.4. INTEGRACION DEL MODELO DE MACHINE LEARNING", level=1)

    add_paragraph_text(doc,
        "Esta seccion documenta el aporte ingenieril principal del proyecto: el diseno, "
        "implementacion y evaluacion del modelo de Machine Learning que permite la intermediacion "
        "laboral automatizada. El sistema utiliza un enfoque hibrido que combina ingenieria de "
        "caracteristicas basada en dominio (feature engineering) con un modelo supervisado de "
        "regresion Ridge para predecir el grado de correspondencia entre perfiles profesionales "
        "y requisitos institucionales."
    )

    add_paragraph_text(doc,
        "El pipeline ML consta de cinco fases principales: (1) extraccion de datos del CV mediante "
        "un LLM (Google Gemini), (2) ingenieria de caracteristicas con 5 scorers especializados, "
        "(3) generacion de dataset sintetico con reglas expertas, (4) entrenamiento del modelo Ridge "
        "con validacion cruzada y busqueda de hiperparametros, y (5) evaluacion con metricas de "
        "regresion y clasificacion."
    )

    # --- Tabla resumen del pipeline ---
    add_formatted_table(doc,
        headers=["Fase", "Componente", "Entrada", "Salida", "Tecnologia"],
        rows=[
            ["1. Extraccion", "LLM Extractor", "PDF del CV (texto)", "JSON estructurado", "Google Gemini API"],
            ["2. Scoring", "5 Scorers especializados", "JSON del CV + Requisitos", "5 scores [0-1]", "spaCy, TF-IDF, CEFR"],
            ["3. Features", "Feature Extractor", "Scores + Pesos inst.", "Vector 18 dimensiones", "NumPy"],
            ["4. Dataset", "Synthetic Generator", "Reglas expertas (15)", "5000 ejemplos CSV", "NumPy, Pandas, Beta/Gamma"],
            ["5. Modelo", "Ridge Regression", "Dataset sintetico", "Modelo .joblib", "scikit-learn"],
            ["6. Prediccion", "Match Predictor", "Vector 18-dim", "Score + Clasificacion", "scikit-learn, joblib"],
        ],
        caption="Fases del pipeline de Machine Learning",
        nota="Cada fase tiene modulos independientes y desacoplados que permiten la evolucion individual de cada componente.",
    )

    # ================================================================
    # 3.4.1. PREPARACION DE LOS DATOS
    # ================================================================

    doc.add_heading("3.4.1. Preparacion de los datos", level=2)

    add_paragraph_text(doc,
        "La preparacion de datos constituye la base fundamental del sistema de ML. Se implemento "
        "un pipeline de ingenieria de caracteristicas que transforma datos no estructurados (texto "
        "del CV) en un vector numerico de 18 dimensiones, listo para ser procesado por el modelo. "
        "Esta seccion documenta cada componente del pipeline."
    )

    # --------------------------------------------------------
    # 3.4.1.1 Extraccion de datos con LLM
    # --------------------------------------------------------

    doc.add_heading("3.4.1.1. Extraccion de datos del CV con LLM (Google Gemini)", level=3)

    add_paragraph_text(doc,
        "El primer paso del pipeline consiste en la extraccion de informacion estructurada desde "
        "el CV en formato PDF. Se utiliza la API de Google Gemini (modelo gemini-flash-latest) para "
        "analizar el texto del CV y extraer las cinco dimensiones clave del perfil profesional: "
        "habilidades tecnicas, habilidades blandas, formacion academica, experiencia laboral e idiomas."
    )

    add_paragraph_text(doc,
        "El proceso consta de dos etapas: (1) extraccion del texto crudo del PDF usando pdfplumber, "
        "y (2) procesamiento del texto con Gemini mediante un prompt especializado que actua como un "
        "reclutador experto en tecnologia y RRHH. El prompt instruye al LLM a retornar un JSON "
        "estructurado con campos especificos para cada dimension."
    )

    # Tabla de estructura del JSON de Gemini
    add_formatted_table(doc,
        headers=["Campo", "Tipo", "Descripcion", "Ejemplo"],
        rows=[
            ["personal_info.summary", "string", "Resumen profesional (max 3 lineas)", "Ingeniero de software con 3 anios..."],
            ["personal_info.languages", "list[str]", "Idiomas con nivel", "['Espanol (Nativo)', 'Ingles (B2)']"],
            ["hard_skills", "list[str]", "Habilidades tecnicas", "['Python', 'React', 'SQL']"],
            ["soft_skills", "list[str]", "Habilidades blandas", "['Liderazgo', 'Comunicacion']"],
            ["education[].degree", "string", "Titulo obtenido", "Ingenieria de Sistemas"],
            ["education[].institution", "string", "Institucion educativa", "EMI"],
            ["education[].year", "string", "Anio de obtencion", "2020"],
            ["experience[].role", "string", "Cargo desempenado", "Desarrollador Full Stack"],
            ["experience[].company", "string", "Empresa", "AGETIC"],
            ["experience[].duration", "string", "Duracion del cargo", "2021 - 2023"],
            ["experience[].description", "string", "Descripcion de responsabilidades", "Desarrollo de APIs REST..."],
        ],
        caption="Estructura del JSON extraido por Google Gemini",
        nota="El LLM recibe un prompt estructurado que solicita la extraccion en este formato exacto. Se limitan los primeros 4000 caracteres del CV.",
    )

    # Codigo del prompt
    add_code_block(doc, '''# llm_extractor.py - Extraccion con Google Gemini
def extract_skills_with_llm(text: str) -> Dict[str, Any]:
    model = genai.GenerativeModel('gemini-flash-latest')
    prompt = f"""
    Actua como un reclutador experto en tecnologia y RRHH.
    Analiza el siguiente texto extraido de un CV y extrae
    la informacion relevante en un formato estructurado.
    TEXTO DEL CV: {text[:4000]}
    Instrucciones:
    1. Extrae Habilidades Tecnicas (Hard Skills) y Blandas
    2. Extrae Historial Academico con titulo, institucion, anio
    3. Extrae Experiencia Laboral con cargo, empresa, duracion
    4. Extrae los Idiomas con nivel
    5. Redacta un Breve Resumen Profesional (max 3 lineas)
    Retorna SOLAMENTE un JSON valido...
    """
    response = model.generate_content(prompt)
    data = json.loads(response.text.strip())
    return data''',
        caption="Codigo fuente del extractor LLM (llm_extractor.py)",
        nota="Se utiliza Google Gemini con prompt engineering para convertir texto no estructurado del CV en un JSON con 5 dimensiones.",
    )

    # Diagrama de flujo placeholder
    add_figure_placeholder(doc,
        "Diagrama de flujo: PDF -> pdfplumber (texto) -> Gemini API (prompt) -> JSON estructurado -> Validacion",
        "Flujo de extraccion de datos del CV con LLM",
        "El flujo muestra como el CV en PDF se transforma en datos estructurados. Si el JSON no contiene los campos requeridos, se asignan valores vacios por defecto.",
    )

    # --------------------------------------------------------
    # 3.4.1.2 Pipeline de ingenieria de caracteristicas
    # --------------------------------------------------------

    doc.add_heading("3.4.1.2. Pipeline de ingenieria de caracteristicas", level=3)

    add_paragraph_text(doc,
        "El Feature Extractor es el componente central que orquesta los 5 scorers especializados "
        "y construye el vector de 18 dimensiones que alimenta al modelo ML. Recibe como entrada "
        "el JSON del CV (extraido por Gemini) y la configuracion institucional (pesos y requisitos), "
        "y produce un vector numerico estandarizado."
    )

    add_paragraph_text(doc,
        "El vector de features se divide en cuatro grupos semanticos: (1) Scores del CV (5 dimensiones), "
        "(2) Pesos institucionales (5 dimensiones), (3) Features de interaccion - producto cruzado "
        "entre score y peso (5 dimensiones), y (4) Features de contexto - experiencia absoluta "
        "(3 dimensiones). Esta estructura permite al modelo capturar tanto la calidad individual "
        "del candidato como su alineacion con los requisitos especificos de cada institucion."
    )

    # Tabla del vector de features
    add_formatted_table(doc,
        headers=["Indice", "Grupo", "Nombre", "Descripcion", "Rango"],
        rows=[
            ["0", "CV Scores", "hard_skills_score", "Score de habilidades tecnicas", "[0, 1]"],
            ["1", "CV Scores", "soft_skills_score", "Score de habilidades blandas", "[0, 1]"],
            ["2", "CV Scores", "experience_score", "Score de experiencia laboral", "[0, 1]"],
            ["3", "CV Scores", "education_score", "Score de formacion academica", "[0, 1]"],
            ["4", "CV Scores", "languages_score", "Score de idiomas", "[0, 1]"],
            ["5", "Pesos Inst.", "inst_weight_hard", "Peso institucional para hard skills", "[0, 1]"],
            ["6", "Pesos Inst.", "inst_weight_soft", "Peso institucional para soft skills", "[0, 1]"],
            ["7", "Pesos Inst.", "inst_weight_exp", "Peso institucional para experiencia", "[0, 1]"],
            ["8", "Pesos Inst.", "inst_weight_edu", "Peso institucional para educacion", "[0, 1]"],
            ["9", "Pesos Inst.", "inst_weight_lang", "Peso institucional para idiomas", "[0, 1]"],
            ["10", "Interaccion", "interaction_hard", "hard_skills_score x weight_hard", "[0, 1]"],
            ["11", "Interaccion", "interaction_soft", "soft_skills_score x weight_soft", "[0, 1]"],
            ["12", "Interaccion", "interaction_exp", "experience_score x weight_exp", "[0, 1]"],
            ["13", "Interaccion", "interaction_edu", "education_score x weight_edu", "[0, 1]"],
            ["14", "Interaccion", "interaction_lang", "languages_score x weight_lang", "[0, 1]"],
            ["15", "Contexto", "total_experience_years", "Anios de experiencia totales", "[0, 10+]"],
            ["16", "Contexto", "min_required_years", "Minimo de anios requeridos", "[0, 5]"],
            ["17", "Contexto", "experience_delta", "total_years - min_required", "[-5, 10+]"],
        ],
        caption="Composicion del vector de features (18 dimensiones)",
        nota="Los features de interaccion (indices 10-14) capturan la alineacion entre las fortalezas del candidato y las prioridades de la institucion. El delta de experiencia (indice 17) puede ser negativo si el candidato no cumple el minimo.",
    )

    # Codigo del feature extractor
    add_code_block(doc, '''# feature_extractor.py - Construccion del vector de 18 dimensiones
def extract_features(gemini_output, institutional_config):
    # PASO 1: Extraer datos del CV
    cv_hard = extract_hard_skills_from_gemini(gemini_output)
    cv_soft = extract_soft_skills_from_gemini(gemini_output)
    cv_edu  = extract_education_from_gemini(gemini_output)
    cv_exp  = extract_experience_from_gemini(gemini_output)
    cv_lang = extract_languages_from_gemini(gemini_output)

    # PASO 2: Calcular scores individuales [0-1]
    hard_result = calculate_hard_skills_score(cv_hard, required, preferred)
    soft_result = calculate_soft_skills_score(cv_soft, required_soft)
    edu_result  = calculate_education_score(cv_edu, required_level)
    exp_result  = calculate_experience_score_from_cv(cv_exp, min_years)
    lang_result = calculate_languages_score_from_cv(cv_lang, required_lang)

    # PASO 3: Construir vector [5 scores + 5 pesos + 5 interacciones + 3 contexto]
    feature_vector = [
        hard_result['score'], soft_result['score'],      # CV scores (5)
        exp_result['score'], edu_result['score'],
        lang_result['score'],
        w_hard, w_soft, w_exp, w_edu, w_lang,            # Pesos inst. (5)
        hard*w_hard, soft*w_soft, exp*w_exp,              # Interacciones (5)
        edu*w_edu, lang*w_lang,
        total_years, min_required, total_years - min_req  # Contexto (3)
    ]
    return {'feature_vector': feature_vector, 'cv_scores': cv_scores, ...}''',
        caption="Codigo fuente del Feature Extractor (feature_extractor.py)",
        nota="El extractor orquesta los 5 scorers y construye el vector numerico. Los pesos institucionales deben sumar 1.0 (validacion estricta con tolerancia de 0.01).",
    )

    # Diagrama de arquitectura
    add_figure_placeholder(doc,
        "Diagrama de arquitectura del Feature Engineering:\n"
        "Gemini JSON -> [Hard Skills Scorer] -> score 0-1\n"
        "           -> [Soft Skills Scorer] -> score 0-1\n"
        "           -> [Education Scorer]   -> score 0-1\n"
        "           -> [Experience Scorer]  -> score 0-1\n"
        "           -> [Languages Scorer]   -> score 0-1\n"
        "5 scores + 5 pesos inst. + 5 interacciones + 3 contexto = Vector[18]",
        "Arquitectura del pipeline de ingenieria de caracteristicas",
        "Cada scorer es un modulo independiente que recibe los datos del CV y los requisitos institucionales, y retorna un score normalizado entre 0 y 1.",
    )

    # --------------------------------------------------------
    # 3.4.1.3 Scorers por dimension
    # --------------------------------------------------------

    doc.add_heading("3.4.1.3. Scorers por dimension", level=3)

    add_paragraph_text(doc,
        "Cada dimension del perfil profesional es evaluada por un scorer especializado que aplica "
        "algoritmos especificos de su dominio. A continuacion se documenta la logica, formula y "
        "parametros de cada uno de los 5 scorers implementados."
    )

    # --- SCORER 1: HARD SKILLS ---
    doc.add_heading("a) Scorer de Habilidades Tecnicas (Hard Skills)", level=4)

    add_paragraph_text(doc,
        "El scorer de habilidades tecnicas es el mas complejo del sistema. Combina tres tecnicas "
        "de NLP para evaluar la correspondencia entre las habilidades del candidato y las requeridas: "
        "(1) coincidencia exacta normalizada, (2) similitud semantica con TF-IDF, y (3) evaluacion "
        "de habilidades preferidas y amplitud del perfil. Ademas, aplica una penalizacion si el "
        "candidato no alcanza un umbral minimo de coincidencia con las habilidades requeridas."
    )

    add_formula_block(doc,
        "S_hard = 0.50 * exact_match + 0.25 * semantic_similarity + 0.15 * preferred_match + 0.10 * breadth_bonus",
        "Formula del scorer de habilidades tecnicas",
        "exact_match: ratio Jaccard normalizado entre skills del CV y requeridos. semantic_similarity: coseno TF-IDF entre vectores de skills. preferred_match: ratio de habilidades preferidas encontradas. breadth_bonus: penalizacion/bonificacion por amplitud del perfil (min(1, total_skills/10)).",
    )

    add_formatted_table(doc,
        headers=["Componente", "Peso", "Algoritmo", "Descripcion"],
        rows=[
            ["Coincidencia exacta", "50%", "Jaccard + normalizacion", "Compara skills normalizados (lower, sinonimos: JS->javascript). ratio = |A interseccion B| / |B_requeridos|"],
            ["Similitud semantica", "25%", "TF-IDF + Coseno", "Vectoriza skills con TfidfVectorizer, calcula similitud coseno entre el corpus del CV y el corpus requerido"],
            ["Skills preferidos", "15%", "Ratio simple", "preferred_found / total_preferred. Evalua habilidades deseables pero no obligatorias"],
            ["Amplitud del perfil", "10%", "Normalizacion lineal", "min(1.0, total_candidate_skills / 10). Bonifica perfiles amplios"],
            ["Penalizacion", "-30%", "Umbral condicional", "Si ratio de coincidencia exacta < 50%, se aplica multiplicador de 0.70 al score total"],
        ],
        caption="Componentes del scorer de habilidades tecnicas",
        nota="La normalizacion de skills incluye conversion a minusculas, eliminacion de espacios extras y mapeo de sinonimos comunes (e.g., 'JS' -> 'javascript', 'TS' -> 'typescript', 'ML' -> 'machine learning').",
    )

    add_code_block(doc, '''# hard_skills_scorer.py - Algoritmo de scoring de habilidades tecnicas
SKILL_SYNONYMS = {
    'js': 'javascript', 'ts': 'typescript', 'py': 'python',
    'ml': 'machine learning', 'dl': 'deep learning',
    'react.js': 'react', 'reactjs': 'react', 'node.js': 'nodejs', ...
}

def calculate_hard_skills_score(cv_skills, required_skills, preferred_skills=[]):
    cv_norm = [normalize_skill(s) for s in cv_skills]
    req_norm = [normalize_skill(s) for s in required_skills]

    # 1. Coincidencia exacta (Jaccard normalizado)
    exact_matches = set(cv_norm) & set(req_norm)
    exact_ratio = len(exact_matches) / len(req_norm) if req_norm else 0.5

    # 2. Similitud semantica (TF-IDF)
    if cv_norm and req_norm:
        vectorizer = TfidfVectorizer()
        tfidf = vectorizer.fit_transform([' '.join(cv_norm), ' '.join(req_norm)])
        semantic_sim = cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0]
    else:
        semantic_sim = 0.0

    # 3. Skills preferidos
    pref_norm = [normalize_skill(s) for s in preferred_skills]
    pref_found = len(set(cv_norm) & set(pref_norm))
    pref_ratio = pref_found / len(pref_norm) if pref_norm else 0.5

    # 4. Amplitud
    breadth = min(1.0, len(cv_norm) / 10)

    # Score combinado
    score = 0.50*exact_ratio + 0.25*semantic_sim + 0.15*pref_ratio + 0.10*breadth

    # Penalizacion si coincidencia < 50%
    if exact_ratio < 0.5 and req_norm:
        score *= 0.70

    return {'score': round(min(1.0, max(0.0, score)), 3), ...}''',
        caption="Codigo fuente del scorer de habilidades tecnicas (hard_skills_scorer.py)",
        nota="Se utiliza TfidfVectorizer de scikit-learn para la similitud semantica y normalizacion de sinonimos para mejorar la precision de coincidencia.",
    )

    # --- SCORER 2: SOFT SKILLS ---
    doc.add_heading("b) Scorer de Habilidades Blandas (Soft Skills)", level=4)

    add_paragraph_text(doc,
        "El scorer de habilidades blandas utiliza un enfoque basado en categorias semanticas. "
        "Las habilidades blandas se agrupan en 4 categorias predefinidas: interpersonal, cognitivo, "
        "organizacional y personal. El score combina coincidencia exacta (70%) con coincidencia "
        "por categoria (30%), lo que permite reconocer habilidades equivalentes aunque no tengan "
        "el mismo nombre exacto."
    )

    add_formula_block(doc,
        "S_soft = 0.70 * exact_match_ratio + 0.30 * category_match_ratio",
        "Formula del scorer de habilidades blandas",
        "exact_match_ratio: ratio de habilidades blandas requeridas encontradas exactamente. category_match_ratio: proporcion de categorias requeridas cubiertas por al menos una habilidad del CV.",
    )

    add_formatted_table(doc,
        headers=["Categoria", "Habilidades incluidas"],
        rows=[
            ["Interpersonal", "Liderazgo, Trabajo en equipo, Comunicacion, Empatia, Negociacion, Colaboracion"],
            ["Cognitivo", "Pensamiento critico, Resolucion de problemas, Creatividad, Adaptabilidad, Aprendizaje continuo, Innovacion"],
            ["Organizacional", "Gestion del tiempo, Planificacion, Organizacion, Gestion de proyectos, Orientacion a resultados"],
            ["Personal", "Responsabilidad, Proactividad, Etica, Autodisciplina, Resiliencia, Compromiso"],
        ],
        caption="Categorias semanticas de habilidades blandas",
        nota="Cada habilidad blanda del CV y de los requisitos se clasifica en una de estas 4 categorias. Habilidades equivalentes en ingles y espanol son mapeadas automaticamente (e.g., 'Leadership' -> 'Liderazgo').",
    )

    # --- SCORER 3: EDUCACION ---
    doc.add_heading("c) Scorer de Formacion Academica (Education)", level=4)

    add_paragraph_text(doc,
        "El scorer de educacion evalua el nivel academico del candidato contra el nivel minimo "
        "requerido por la institucion. Se utiliza una jerarquia de 8 niveles educativos del "
        "sistema boliviano, donde cada nivel tiene un valor numerico asignado. Si el candidato "
        "cumple o supera el nivel requerido, obtiene un score de 1.0; si esta por debajo, el "
        "score es proporcional a la distancia entre su nivel y el requerido."
    )

    add_formatted_table(doc,
        headers=["Nivel Educativo", "Valor Numerico", "Descripcion"],
        rows=[
            ["Tecnico Medio", "0.25", "Formacion tecnica basica (bachillerato tecnico)"],
            ["Tecnico Superior", "0.45", "Titulo tecnico superior (2-3 anios)"],
            ["Licenciatura", "0.75", "Grado universitario de licenciatura (4-5 anios)"],
            ["Ingenieria", "0.75", "Titulo de ingenieria (5 anios, equivalente a licenciatura)"],
            ["Diplomado", "0.80", "Formacion de posgrado corta (6-12 meses)"],
            ["Especialidad", "0.85", "Especializacion de posgrado"],
            ["Maestria", "0.92", "Grado de maestria (2 anios adicionales)"],
            ["Doctorado", "1.00", "Grado maximo academico"],
        ],
        caption="Jerarquia de niveles educativos del sistema boliviano",
        nota="Los valores estan calibrados para el contexto del mercado laboral boliviano. Licenciatura e Ingenieria comparten el mismo valor (0.75) por ser equivalentes en el sistema CEUB.",
        fuente="Elaboracion propia basada en el sistema educativo boliviano (CEUB)"
    )

    add_formula_block(doc,
        "S_edu = { 1.0 si nivel_candidato >= nivel_requerido ;  valor_candidato / valor_requerido en caso contrario }",
        "Formula del scorer de formacion academica",
        "Se comparan los valores numericos de los niveles. Si el candidato iguala o supera el nivel requerido, el score es perfecto (1.0). De lo contrario, el score es proporcional.",
    )

    # --- SCORER 4: EXPERIENCIA ---
    doc.add_heading("d) Scorer de Experiencia Laboral (Experience)", level=4)

    add_paragraph_text(doc,
        "El scorer de experiencia utiliza una curva logaritmica para evaluar los anios de "
        "experiencia del candidato contra el minimo requerido. La funcion logaritmica se eligio "
        "porque refleja el retorno decreciente de la experiencia adicional: los primeros anios "
        "aportan mas valor que los anios posteriores. El scorer contempla 4 escenarios: sin "
        "experiencia, por debajo del minimo, entre el minimo y el ideal (5 anios), y por encima "
        "del ideal."
    )

    add_formatted_table(doc,
        headers=["Escenario", "Condicion", "Formula", "Score Resultante"],
        rows=[
            ["Sin experiencia", "total_years = 0", "0.0", "0.0"],
            ["Debajo del minimo", "total_years < min_required", "(years / min) * 0.5", "0.0 - 0.5"],
            ["Entre minimo e ideal", "min <= years < 5", "0.5 + 0.5 * log_norm", "0.5 - 1.0"],
            ["Encima del ideal", "years >= 5 (ideal)", "1.0", "1.0"],
        ],
        caption="Escenarios del scorer de experiencia laboral",
        nota="El valor ideal por defecto es 5 anios. La normalizacion logaritmica usa la formula: log_norm = (log(years+1) - log(min+1)) / (log(max+1) - log(min+1)), donde max es el ideal.",
    )

    add_formula_block(doc,
        "S_exp(y) = 0.5 + 0.5 * [ ln(y+1) - ln(min+1) ] / [ ln(ideal+1) - ln(min+1) ]    para min <= y < ideal",
        "Formula logaritmica del scorer de experiencia laboral",
        "y = anios de experiencia totales. min = minimo requerido. ideal = 5 anios (umbral de experiencia plena). La funcion logaritmica natural modela el retorno decreciente de la experiencia.",
    )

    add_figure_placeholder(doc,
        "Grafica de la curva logaritmica del scorer de experiencia:\n"
        "Eje X: Anios de experiencia (0-8), Eje Y: Score (0-1)\n"
        "Mostrar las 4 zonas: sin exp (0,0), debajo minimo (lineal 0-0.5), logaritmica (0.5-1.0), plateau (1.0)",
        "Curva de scoring de experiencia laboral",
        "La curva logaritmica garantiza que un candidato con el minimo de experiencia ya obtiene un score de 0.5, y el crecimiento es rapido en los primeros anios pero desacelera a medida que se acerca al ideal.",
    )

    # --- SCORER 5: IDIOMAS ---
    doc.add_heading("e) Scorer de Idiomas (Languages)", level=4)

    add_paragraph_text(doc,
        "El scorer de idiomas evalua la competencia linguistica del candidato utilizando el "
        "Marco Comun Europeo de Referencia (CEFR). Cada nivel de idioma tiene un valor numerico "
        "asignado que refleja la competencia comunicativa. El score final se calcula como el "
        "promedio de los scores de coincidencia para cada idioma requerido."
    )

    add_formatted_table(doc,
        headers=["Nivel CEFR", "Valor", "Nivel Descriptivo", "Valor"],
        rows=[
            ["A1", "0.35", "Basico", "0.30"],
            ["A2", "0.45", "Intermedio", "0.55"],
            ["B1", "0.60", "Avanzado", "0.85"],
            ["B2", "0.75", "Fluido", "0.80"],
            ["C1", "0.90", "Nativo", "1.00"],
            ["C2", "0.95", "Materno", "1.00"],
        ],
        caption="Niveles de idiomas y sus valores numericos (CEFR)",
        nota="El sistema acepta tanto niveles CEFR (A1-C2) como niveles descriptivos (Basico, Intermedio, Avanzado). 'Nativo' y 'Materno' se consideran equivalentes con valor maximo.",
        fuente="Elaboracion propia basada en el Marco Comun Europeo de Referencia para las Lenguas (CEFR)"
    )

    # --------------------------------------------------------
    # 3.4.1.4 Dataset sintetico estructurado
    # --------------------------------------------------------

    doc.add_heading("3.4.1.4. Dataset sintetico estructurado", level=3)

    add_paragraph_text(doc,
        "Dado que no se dispone de un dataset historico de evaluaciones de candidatos contra "
        "perfiles institucionales, se implemento un generador de dataset sintetico basado en "
        "reglas expertas. El generador crea 5000 ejemplos de entrenamiento con distribuciones "
        "estadisticas realistas y un sistema de 15 reglas heuristicas que definen la relacion "
        "entre las features de entrada y el score de matching esperado."
    )

    add_paragraph_text(doc,
        "La justificacion del enfoque sintetico radica en tres factores: (1) el sistema es nuevo "
        "y no existe informacion historica de matching candidato-institucion, (2) las reglas expertas "
        "capturan el conocimiento de dominio sobre como se evaluan perfiles profesionales en Bolivia, "
        "y (3) el modelo supervisado entrenado con estos datos puede aprender patrones no lineales "
        "y generalizaciones que van mas alla de las reglas individuales."
    )

    # Tabla de distribuciones
    add_formatted_table(doc,
        headers=["Feature", "Distribucion", "Parametros", "Justificacion"],
        rows=[
            ["Hard Skills Score", "Beta(5, 2)", "Sesgada hacia alto", "En el sector tecnologico, la mayoria tiene habilidades digitales basicas"],
            ["Soft Skills Score", "Beta(4, 3)", "Ligeramente alta", "Distribucion mas uniforme pero con sesgo positivo"],
            ["Experience Score", "Beta(3, 3)", "Uniforme (simetrica)", "La experiencia varia significativamente entre candidatos"],
            ["Education Score", "Categorico", "{0.25, 0.45, 0.75, 0.92, 1.0}", "Distribucion discreta segun niveles educativos"],
            ["Languages Score", "Beta(3, 4)", "Sesgada hacia bajo", "Idiomas extranjeros son dificiles en el contexto boliviano"],
            ["Institutional Weights", "Dirichlet(2,2,2,2,2)", "Suma = 1.0 garantizada", "Distribucion simetrica que genera pesos validos automaticamente"],
            ["Total Years (Exp.)", "Gamma(3, 1)", "Media ~3, max 10", "Refleja experiencia tipica de egresados recientes"],
            ["Min Required Years", "Uniform(0, 5)", "Rango [0, 5]", "Requisito minimo uniforme"],
        ],
        caption="Distribuciones estadisticas para la generacion del dataset sintetico",
        nota="Las distribuciones Beta y Gamma fueron seleccionadas por su capacidad de modelar sesgos realistas. La distribucion Dirichlet garantiza que los pesos siempre sumen 1.0.",
    )

    # Tabla de distribuciones objetivo de clases
    add_formatted_table(doc,
        headers=["Clase", "Proporcion Objetivo", "Descripcion"],
        rows=[
            ["APTO", "40%", "Alta correspondencia (score >= 0.70)"],
            ["CONSIDERADO", "35%", "Correspondencia media (0.50 <= score < 0.70)"],
            ["NO_APTO", "25%", "Baja correspondencia (score < 0.50)"],
        ],
        caption="Distribucion objetivo de clases en el dataset sintetico",
        nota="La distribucion refleja un escenario realista donde hay mas candidatos aptos que no aptos, ya que el sistema esta disenado para el sector de tecnologia y egresados con formacion relevante.",
    )

    # --------------------------------------------------------
    # 3.4.1.5 Reglas expertas (15 reglas)
    # --------------------------------------------------------

    doc.add_heading("3.4.1.5. Sistema de reglas expertas para etiquetado", level=3)

    add_paragraph_text(doc,
        "El corazon del generador sintetico es un sistema de 15 reglas heuristicas que definen "
        "como calcular el score 'verdadero' (label) para cada ejemplo. Estas reglas codifican "
        "el conocimiento de dominio sobre evaluacion de perfiles y son la logica que el modelo "
        "ML aprendera a replicar y generalizar."
    )

    add_formatted_table(doc,
        headers=["#", "Regla", "Condicion", "Efecto", "Magnitud"],
        rows=[
            ["1", "Score base ponderado", "Siempre", "Suma ponderada de scores x pesos", "Base"],
            ["2", "Penalizacion experiencia insuficiente", "years < min_required", "Penalizacion proporcional al deficit", "*(0.5 + 0.5*ratio)"],
            ["3", "Hard skills criticos bajos", "hard_skills < 0.5 AND peso > 0.3", "Penalizacion por habilidades tecnicas debiles", "*0.70 (-30%)"],
            ["4", "Soft skills criticos bajos", "soft_skills < 0.4 AND peso > 0.25", "Penalizacion por habilidades blandas debiles", "*0.75 (-25%)"],
            ["5", "Bonificacion educacion maxima", "education >= 0.92 (Maestria+)", "Bonificacion por formacion avanzada", "*1.08 (+8%)"],
            ["6", "Perfil muy completo", "Todos los scores >= 0.7", "Bonificacion por excelencia general", "*1.10 (+10%)"],
            ["7", "Educacion insuficiente con peso alto", "education < 0.75 AND peso > 0.20", "Penalizacion por deficit educativo", "*0.85 (-15%)"],
            ["8", "Experiencia muy superior", "delta > 2 anios", "Bonificacion por experiencia excedente", "*1.05 (+5%)"],
            ["9", "Idiomas insuficientes", "languages < 0.5 AND peso > 0.15", "Penalizacion por deficit de idiomas", "*0.80 (-20%)"],
            ["10", "Sinergia hard+soft altas", "hard > 0.75 AND soft > 0.75", "Bonificacion por equilibrio tecnico-humano", "*1.07 (+7%)"],
            ["11", "Debil en dimension prioritaria", "max_peso > 0.35 AND score < 0.4", "Penalizacion fuerte por fallar en lo critico", "*0.65 (-35%)"],
            ["12", "Sobresale en dimension prioritaria", "max_peso > 0.30 AND score > 0.85", "Bonificacion por destacar donde importa", "*1.06 (+6%)"],
            ["13", "Perfil desequilibrado", "std(scores) > 0.35", "Penalizacion por alta variabilidad", "*0.90 (-10%)"],
            ["14", "Experiencia + educacion alta", "exp > 0.7 AND edu >= 0.75", "Bonificacion por combinacion solida", "*1.05 (+5%)"],
            ["15", "Multiples dimensiones criticas", "2+ scores bajo 0.3 o 0.5", "Penalizacion extrema por multiples fallos", "*0.60 (-40%)"],
        ],
        caption="Sistema de 15 reglas expertas para etiquetado del dataset sintetico",
        nota="Las reglas se aplican secuencialmente. Finalmente se agrega ruido gaussiano (media=0, std=0.04) para simular variabilidad real y se clipea el score a [0, 1].",
    )

    add_code_block(doc, '''# synthetic_generator.py - Sistema de reglas expertas (extracto)
def apply_expert_rules(self, cv_scores, weights, context):
    # REGLA 1: Score base ponderado
    base_score = sum(cv_scores[dim] * weights[dim] for dim in dimensions)

    # REGLA 2: Penalizacion si no cumple minimo de experiencia
    if context['total_years'] < context['min_required']:
        deficit_ratio = total_years / min_required if min_required > 0 else 1.0
        base_score *= (0.5 + 0.5 * deficit_ratio)

    # REGLA 3: Hard skills criticos
    if cv_scores['hard_skills'] < 0.5 and weights['weight_hard'] > 0.3:
        base_score *= 0.70  # Penalizacion del 30%

    # ... (reglas 4-15) ...

    # REGLA 15: Multiples dimensiones criticas fallan
    critical_failures = sum([
        cv_scores['hard_skills'] < 0.3,
        cv_scores['soft_skills'] < 0.3,
        cv_scores['experience'] < 0.3,
        cv_scores['education'] < 0.5
    ])
    if critical_failures >= 2:
        base_score *= 0.60  # Penalizacion muy fuerte del 40%

    # Ruido gaussiano realista
    noise = np.random.normal(0, 0.04)
    final_score = np.clip(base_score + noise, 0, 1)
    return float(final_score)''',
        caption="Codigo fuente del sistema de reglas expertas (synthetic_generator.py)",
        nota="Las 15 reglas cubren escenarios de penalizacion (experiencia insuficiente, skills bajos, perfil desequilibrado) y bonificacion (educacion avanzada, perfil completo, sinergia). El ruido gaussiano (std=0.04) simula variabilidad del ~3-5%.",
    )

    # Estadisticas del dataset generado
    add_formatted_table(doc,
        headers=["Metrica", "Valor"],
        rows=[
            ["Total de ejemplos", "5,000"],
            ["Features de entrada", "18"],
            ["Variable objetivo", "match_score (continuo, [0-1])"],
            ["Media del match_score", "~0.55"],
            ["Desviacion estandar", "~0.20"],
            ["Distribucion APTO (>=0.70)", "~40%"],
            ["Distribucion CONSIDERADO (0.50-0.69)", "~35%"],
            ["Distribucion NO_APTO (<0.50)", "~25%"],
            ["Semilla aleatoria", "42 (reproducible)"],
            ["Formato de salida", "CSV (synthetic_dataset.csv)"],
        ],
        caption="Estadisticas del dataset sintetico generado",
        nota="El dataset se genera con seed=42 para garantizar reproducibilidad. Las distribuciones aproximadas dependen de la interaccion entre las reglas expertas y las distribuciones de entrada.",
    )

    add_figure_placeholder(doc,
        "Histograma de distribucion de match_score en el dataset sintetico:\n"
        "Ejecutar: python -m app.ml.data.synthetic_generator y capturar la salida con las estadisticas",
        "Distribucion del match_score en el dataset sintetico",
        "La distribucion muestra una concentracion de scores entre 0.4 y 0.8, con colas en ambos extremos representando candidatos muy aptos y no aptos.",
    )

    # ================================================================
    # 3.4.2. MODELADO
    # ================================================================

    doc.add_heading("3.4.2. Modelado", level=2)

    add_paragraph_text(doc,
        "Esta seccion documenta la seleccion, entrenamiento y evaluacion del modelo de Machine Learning. "
        "Se eligio Ridge Regression (regresion lineal con regularizacion L2) como algoritmo principal "
        "debido a su interpretabilidad, robustez ante multicolinealidad y capacidad de generalizacion "
        "a multiples configuraciones institucionales sin reentrenamiento."
    )

    # --------------------------------------------------------
    # 3.4.2.1 Informe de seleccion del algoritmo
    # --------------------------------------------------------

    doc.add_heading("3.4.2.1. Informe de seleccion del algoritmo", level=3)

    add_paragraph_text(doc,
        "La seleccion del algoritmo se baso en cinco criterios fundamentales del proyecto: "
        "interpretabilidad, generalizacion, eficiencia computacional, facilidad de despliegue "
        "y adecuacion al tipo de problema. A continuacion se presenta el analisis comparativo "
        "de los algoritmos candidatos."
    )

    add_formatted_table(doc,
        headers=["Criterio", "Ridge Regression", "Random Forest", "SVM (SVR)", "Red Neuronal (MLP)"],
        rows=[
            ["Interpretabilidad", "Alta - coeficientes explicitos", "Media - importancia de features", "Baja - kernel trick", "Muy baja - caja negra"],
            ["Generalizacion", "Alta - regularizacion L2", "Media - riesgo sobreajuste", "Alta con buen kernel", "Alta con suficientes datos"],
            ["Eficiencia", "Muy alta - O(n*p^2)", "Media - multiples arboles", "Baja con kernel RBF", "Baja - backpropagation"],
            ["Despliegue", "Simple - joblib ~100KB", "Medio - bosque grande", "Medio", "Complejo - framework DL"],
            ["Datos necesarios", "Pocos (< 5000)", "Moderados (5000+)", "Moderados", "Muchos (10000+)"],
            ["Multicolinealidad", "Robusta (L2)", "Inmune", "Sensible", "Sensible"],
            ["Reentrenamiento", "No necesario*", "Necesario", "Necesario", "Necesario"],
        ],
        caption="Analisis comparativo de algoritmos candidatos",
        nota="(*) Ridge Regression no requiere reentrenamiento cuando cambian las configuraciones institucionales, ya que estas son features de entrada (no parametros del modelo). Esto es una ventaja clave del enfoque basado en features de interaccion.",
    )

    add_paragraph_text(doc,
        "Se eligio Ridge Regression por las siguientes razones justificadas:"
    )

    doc.add_paragraph("Interpretabilidad: Los coeficientes del modelo revelan directamente la importancia relativa de cada feature. Esto permite explicar a los usuarios por que un candidato obtuvo un determinado score, identificando fortalezas y debilidades.", style="List Bullet")
    doc.add_paragraph("Regularizacion L2: La penalizacion L2 previene el sobreajuste al reducir la magnitud de los coeficientes. Esto es critico con un dataset sintetico de 5000 ejemplos donde el riesgo de sobreajuste es real.", style="List Bullet")
    doc.add_paragraph("Generalizacion multi-institucional: Al incluir los pesos institucionales como features de entrada (no como parametros), el modelo aprende a adaptarse a cualquier configuracion institucional sin necesidad de reentrenamiento.", style="List Bullet")
    doc.add_paragraph("Eficiencia computacional: El entrenamiento toma menos de 1 segundo y la prediccion es instantanea, lo que permite uso en tiempo real en la API.", style="List Bullet")
    doc.add_paragraph("Facilidad de despliegue: El modelo serializado ocupa menos de 100KB y se carga en memoria con joblib, sin dependencias adicionales de frameworks de deep learning.", style="List Bullet")

    # Formula de Ridge
    add_formula_block(doc,
        "J(w) = || y - Xw ||^2 + alpha * || w ||^2\n\n"
        "donde: w = vector de coeficientes, X = matriz de features, y = scores objetivo, alpha = parametro de regularizacion",
        "Funcion de costo de Ridge Regression (L2)",
        "El termino de regularizacion alpha * ||w||^2 penaliza coeficientes grandes, forzando al modelo a distribuir el poder predictivo entre multiples features en lugar de depender excesivamente de uno solo.",
        fuente="Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning."
    )

    # --------------------------------------------------------
    # 3.4.2.2 Modelo entrenado y serializado
    # --------------------------------------------------------

    doc.add_heading("3.4.2.2. Arquitectura y entrenamiento del modelo", level=3)

    add_paragraph_text(doc,
        "El modelo se implementa en la clase InstitutionalMatchModel que encapsula el estimador "
        "Ridge de scikit-learn junto con un StandardScaler para normalizacion de features. El "
        "pipeline de entrenamiento incluye: division train/test (80/20), busqueda de hiperparametros "
        "con Grid Search, validacion cruzada 5-fold, y evaluacion final en el conjunto de test."
    )

    # Tabla de hiperparametros
    add_formatted_table(doc,
        headers=["Hiperparametro", "Valores Explorados", "Valor Seleccionado", "Justificacion"],
        rows=[
            ["alpha", "[0.01, 0.1, 0.5, 1.0, 2.0, 5.0, 10.0]", "Seleccion por Grid Search", "Balance entre sesgo y varianza"],
            ["normalize", "True", "True (StandardScaler)", "Necesario porque features tienen escalas diferentes (scores [0,1] vs years [0,10])"],
            ["fit_intercept", "True", "True", "Permite sesgo base (intercepto)"],
            ["solver", "auto", "auto", "scikit-learn selecciona automaticamente el solver optimo"],
            ["random_state", "42", "42", "Reproducibilidad"],
            ["test_size", "0.20", "0.20", "1000 ejemplos de test (de 5000 totales)"],
            ["cv_folds", "5", "5", "Validacion cruzada estandar"],
        ],
        caption="Hiperparametros del modelo Ridge Regression",
        nota="Grid Search explora 7 valores de alpha en escala logaritmica. La validacion cruzada 5-fold se usa para seleccionar el mejor alpha antes de entrenar el modelo final con todo el training set.",
    )

    # Pipeline de entrenamiento
    add_code_block(doc, '''# train_model.py - Pipeline de entrenamiento completo
trainer = ModelTrainer(random_state=42)

# 1. Cargar dataset sintetico
X, y, feature_names = trainer.load_dataset('synthetic_dataset.csv')
# Dataset: 5000 ejemplos x 18 features

# 2. Split train/test (80/20)
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, shuffle=True
)  # Train: 4000 | Test: 1000

# 3. Grid Search de alpha (7 valores, 5-fold CV)
grid_search = GridSearchCV(
    Ridge(), param_grid={'alpha': [0.01, 0.1, 0.5, 1.0, 2.0, 5.0, 10.0]},
    cv=5, scoring='r2', n_jobs=-1
)
grid_search.fit(X_train, y_train)
best_alpha = grid_search.best_params_['alpha']

# 4. Entrenar modelo final con mejor alpha
model = InstitutionalMatchModel(alpha=best_alpha, normalize=True)
model.fit(X_train, y_train, feature_names=feature_names)

# 5. Evaluar en test set
y_pred = model.predict(X_test)  # Predicciones clipeadas a [0, 1]

# 6. Guardar modelo serializado
model.save('trained_models/ridge_v1.joblib')''',
        caption="Codigo fuente del pipeline de entrenamiento (train_model.py)",
        nota="El pipeline es completamente reproducible gracias a random_state=42. El modelo final se serializa con joblib incluyendo el scaler, coeficientes, feature_names y metricas de entrenamiento.",
    )

    # Arquitectura de la clase
    add_code_block(doc, '''# ridge_model.py - Clase InstitutionalMatchModel
class InstitutionalMatchModel:
    """Modelo Ridge con normalizacion, explicabilidad y serializacion"""

    def __init__(self, alpha=1.0, normalize=True):
        self.model = Ridge(alpha=alpha, fit_intercept=True,
                          solver='auto', random_state=42)
        self.scaler = StandardScaler() if normalize else None

    def fit(self, X, y, feature_names=None):
        if self.normalize:
            X_scaled = self.scaler.fit_transform(X)
        self.model.fit(X_scaled, y)
        self.is_trained = True

    def predict(self, X):
        X_scaled = self.scaler.transform(X) if self.normalize else X
        y_pred = self.model.predict(X_scaled)
        return np.clip(y_pred, 0, 1)  # Clipear a [0, 1]

    def predict_single(self, feature_vector):
        """Prediccion con explicabilidad"""
        score = self.predict(feature_vector.reshape(1, -1))[0]
        classification = self._classify_score(score)
        contributions = self._calculate_feature_contributions(feature_vector)
        return {
            'match_score': score,
            'classification': classification,  # APTO/CONSIDERADO/NO_APTO
            'feature_contributions': contributions,
            'top_strengths': top_3_positive,
            'top_weaknesses': top_3_negative
        }

    def save(self, filepath):
        joblib.dump({
            'model': self.model, 'scaler': self.scaler,
            'alpha': self.alpha, 'feature_names': self.feature_names,
            'training_metrics': self.training_metrics
        }, filepath)

    @classmethod
    def load(cls, filepath):
        model_data = joblib.load(filepath)
        instance = cls(alpha=model_data['alpha'])
        instance.model = model_data['model']
        instance.scaler = model_data['scaler']
        return instance''',
        caption="Codigo fuente del modelo Ridge con explicabilidad (ridge_model.py)",
        nota="La clase encapsula entrenamiento, prediccion, explicabilidad (contribuciones de features) y serializacion. El metodo predict_single retorna score, clasificacion y explicacion detallada.",
    )

    # Tabla de clasificacion
    add_formatted_table(doc,
        headers=["Clasificacion", "Umbral", "Interpretacion", "Color en UI"],
        rows=[
            ["APTO", "score >= 0.70", "Alta correspondencia entre perfil y requisitos", "Verde (#2ecc71)"],
            ["CONSIDERADO", "0.50 <= score < 0.70", "Correspondencia media, revisar detalles", "Naranja (#f39c12)"],
            ["NO_APTO", "score < 0.50", "Baja correspondencia, perfil no se ajusta", "Rojo (#e74c3c)"],
        ],
        caption="Umbrales de clasificacion del sistema",
        nota="Los umbrales son configurables por perfil institucional. Los valores por defecto (0.70 y 0.50) fueron establecidos en base a criterios de seleccion tipicos del mercado laboral boliviano.",
    )

    # --------------------------------------------------------
    # 3.4.2.3 Explicabilidad del modelo
    # --------------------------------------------------------

    doc.add_heading("3.4.2.3. Explicabilidad del modelo", level=3)

    add_paragraph_text(doc,
        "Una ventaja clave de Ridge Regression es su interpretabilidad intrinseca. Los coeficientes "
        "aprendidos por el modelo tienen un significado directo: cada coeficiente indica cuanto "
        "contribuye su feature correspondiente al score final. Esto permite generar explicaciones "
        "personalizadas para cada prediccion."
    )

    add_paragraph_text(doc,
        "La contribucion de cada feature se calcula como: contribucion_i = coeficiente_i * valor_feature_i. "
        "Las features con mayor contribucion positiva se identifican como fortalezas del candidato, "
        "mientras que las de menor contribucion se identifican como areas de mejora. Esta informacion "
        "se presenta al usuario en el frontend."
    )

    add_formatted_table(doc,
        headers=["Grupo de Features", "Features", "Interpretacion de Coeficientes"],
        rows=[
            ["CV Scores (0-4)", "5 scores individuales", "Coeficiente positivo alto = esa dimension es muy valorada por el modelo"],
            ["Pesos Inst. (5-9)", "5 pesos de configuracion", "El modelo aprende como los pesos modulan la importancia de cada dimension"],
            ["Interacciones (10-14)", "5 productos cruzados", "Capturan la sinergia: un score alto importa MAS si la institucion lo valora"],
            ["Contexto (15-17)", "3 features de experiencia", "El delta (anios - minimo) tiene coeficiente positivo fuerte"],
        ],
        caption="Interpretacion de los coeficientes del modelo por grupo de features",
        nota="Los features de interaccion (grupo 3) son los mas informativos porque combinan la calidad del candidato con la prioridad institucional. Un candidato con alto score en hard_skills obtiene mayor beneficio si la institucion asigna peso alto a esa dimension.",
    )

    add_figure_placeholder(doc,
        "Grafica de barras horizontales mostrando los coeficientes del modelo:\n"
        "Ejecutar: python -m app.ml.scripts.train_model y capturar la seccion 'COEFICIENTES DEL MODELO'\n"
        "O usar la visualizacion: backend/app/ml/visualizations/feature_importance.png",
        "Coeficientes del modelo Ridge (importancia de features)",
        "Los coeficientes positivos (verde) aumentan el score de matching, mientras que los negativos (rojo) lo reducen. Las barras mas largas indican mayor influencia en la prediccion.",
    )

    # --------------------------------------------------------
    # 3.4.2.4 Reporte de metricas
    # --------------------------------------------------------

    doc.add_heading("3.4.2.4. Reporte de metricas de evaluacion", level=3)

    add_paragraph_text(doc,
        "El modelo se evalua con metricas de regresion (R2, RMSE, MAE) y de clasificacion "
        "(Accuracy, Precision, Recall, F1-Score). Las metricas de regresion miden la precision "
        "numerica del score predicho, mientras que las de clasificacion evaluan si el modelo "
        "asigna correctamente la categoria APTO/CONSIDERADO/NO_APTO."
    )

    # Tabla de metricas de regresion
    add_formatted_table(doc,
        headers=["Metrica", "Formula", "Objetivo", "Interpretacion"],
        rows=[
            ["R2 Score", "1 - SS_res / SS_tot", ">= 0.75", "Proporcion de varianza explicada. 1.0 = perfecto, 0 = modelo media"],
            ["MSE", "(1/n) * sum((y_true - y_pred)^2)", "Minimizar", "Error cuadratico medio. Penaliza errores grandes"],
            ["RMSE", "sqrt(MSE)", "<= 0.17", "Raiz del MSE, misma unidad que y. Interpretable directamente"],
            ["MAE", "(1/n) * sum(|y_true - y_pred|)", "<= 0.12", "Error absoluto medio. Menos sensible a outliers que RMSE"],
            ["MAPE", "(100/n) * sum(|y_true - y_pred| / y_true)", "Minimizar", "Error porcentual. Expresa el error como porcentaje"],
        ],
        caption="Metricas de regresion para evaluacion del modelo",
        nota="Los objetivos de R2 >= 0.75 y RMSE <= 0.17 fueron establecidos como criterios de aceptacion del modelo. Un RMSE de 0.17 significa un error promedio de ~17 puntos porcentuales en el score.",
    )

    # Tabla de metricas de clasificacion
    add_formatted_table(doc,
        headers=["Metrica", "Formula", "Objetivo", "Aplicacion"],
        rows=[
            ["Accuracy", "Correctos / Total", ">= 0.80", "Proporcion general de clasificaciones correctas"],
            ["Precision (por clase)", "VP / (VP + FP)", "Alta", "De los clasificados como X, cuantos realmente son X"],
            ["Recall (por clase)", "VP / (VP + FN)", "Alta", "De los que realmente son X, cuantos se clasificaron como X"],
            ["F1-Score (por clase)", "2 * (P * R) / (P + R)", "Alta", "Media armonica de Precision y Recall"],
            ["F1-Score Weighted", "Promedio ponderado F1", ">= 0.80", "F1 global ponderado por soporte de cada clase"],
        ],
        caption="Metricas de clasificacion para evaluacion del modelo",
        nota="La clasificacion se deriva del score de regresion usando los umbrales: APTO >= 0.70, CONSIDERADO >= 0.50, NO_APTO < 0.50.",
    )

    add_figure_placeholder(doc,
        "Captura de la salida del pipeline de entrenamiento mostrando:\n"
        "- Grid Search results (mejor alpha)\n"
        "- Cross-validation R2 scores (5 folds)\n"
        "- Metricas en test set (R2, RMSE, MAE, Accuracy)\n"
        "Ejecutar: python -m app.ml.scripts.train_model",
        "Resultados del pipeline de entrenamiento del modelo",
        "La salida muestra el proceso completo: carga de datos, grid search, cross-validation, evaluacion en test y resumen final del modelo.",
    )

    # Comparacion con baselines
    doc.add_heading("3.4.2.5. Comparacion con baselines", level=3)

    add_paragraph_text(doc,
        "Para validar que el modelo Ridge aporta valor predictivo mas alla de aproximaciones "
        "triviales, se comparo contra tres baselines: (1) predictor de media (siempre predice "
        "el promedio del training set), (2) predictor de mediana, y (3) regresion lineal simple "
        "(sin regularizacion L2)."
    )

    add_formatted_table(doc,
        headers=["Modelo", "R2 Score", "RMSE", "MAE", "Accuracy (clasif.)"],
        rows=[
            ["Baseline Media", "0.0000", "~0.20", "~0.17", "~0.35"],
            ["Baseline Mediana", "~-0.01", "~0.20", "~0.17", "~0.35"],
            ["Regresion Lineal (sin L2)", "~R2", "~RMSE", "~MAE", "~Acc"],
            ["Ridge Regression (L2)", "R2*", "RMSE*", "MAE*", "Acc*"],
        ],
        caption="Comparacion del modelo Ridge contra baselines",
        nota="(*) Reemplazar con los valores reales ejecutando: python -m app.ml.scripts.train_model. Los baselines Media y Mediana tienen R2 cercano a 0 por definicion, lo que confirma que el modelo Ridge captura patrones reales.",
    )

    add_figure_placeholder(doc,
        "Insertar los valores reales de metricas del modelo entrenado.\n"
        "Ejecutar el script train_model.py y tomar las metricas de la seccion 'Metricas en Test'",
        "Verificacion de objetivos del modelo",
        "Se verifican 4 objetivos: R2 >= 0.75, RMSE <= 0.17, MAE <= 0.12, Accuracy >= 0.80.",
    )

    # --------------------------------------------------------
    # 3.4.2.6 Visualizaciones del modelo
    # --------------------------------------------------------

    doc.add_heading("3.4.2.6. Visualizaciones de evaluacion", level=3)

    add_paragraph_text(doc,
        "Se generan 8 visualizaciones para el analisis completo del modelo. Cada visualizacion "
        "se encuentra en el directorio backend/app/ml/visualizations/ y puede regenerarse ejecutando "
        "el modulo de evaluacion. A continuacion se describen las mas relevantes."
    )

    # Predicciones vs Reales
    add_figure_placeholder(doc,
        "Scatter plot Predicciones vs Valores Reales con colores por clase (APTO=verde, CONSIDERADO=naranja, NO_APTO=rojo).\n"
        "Archivo: backend/app/ml/visualizations/predictions_vs_actual.png\n"
        "O ejecutar: python -m app.ml.evaluation.visualizations",
        "Predicciones vs Valores Reales del modelo Ridge",
        "Los puntos cercanos a la diagonal roja indican predicciones precisas. La dispersion en torno a la diagonal muestra el error del modelo. Los colores permiten identificar si las clasificaciones derivadas son correctas.",
    )

    # Residuos
    add_figure_placeholder(doc,
        "Doble grafica: (1) Histograma de residuos, (2) Residuos vs Predicciones.\n"
        "Archivo: backend/app/ml/visualizations/residuals_distribution.png",
        "Distribucion de residuos del modelo",
        "Un histograma centrado en cero con forma gaussiana indica que el modelo no tiene sesgo sistematico. El plot de residuos vs predicciones verifica homocedasticidad (varianza constante del error).",
    )

    # Matriz de confusion
    add_figure_placeholder(doc,
        "Matriz de confusion 3x3 con valores absolutos y porcentajes.\n"
        "Archivo: backend/app/ml/visualizations/confusion_matrix.png",
        "Matriz de confusion del modelo (3 clases)",
        "La diagonal principal muestra las clasificaciones correctas. Los valores fuera de la diagonal son errores. Los porcentajes indican la proporcion de cada celda respecto a la fila (clase real).",
    )

    # Curva de aprendizaje
    add_figure_placeholder(doc,
        "Curva de aprendizaje: Score de entrenamiento y validacion vs tamano del training set.\n"
        "Archivo: backend/app/ml/visualizations/learning_curve.png",
        "Curva de aprendizaje del modelo Ridge",
        "Si las curvas convergen, el modelo tiene buen ajuste. Una brecha grande entre train y validation indica sobreajuste. Incluye diagnostico automatico (Overfitting/Underfitting/Buen ajuste).",
    )

    # Scores de Cross-Validation
    add_figure_placeholder(doc,
        "Barras de R2 score por fold de Cross-Validation (5 folds).\n"
        "Archivo: backend/app/ml/visualizations/cv_scores_distribution.png",
        "Scores de Cross-Validation (5-Fold)",
        "Las barras muestran la estabilidad del modelo a traves de los 5 folds. Barras de altura similar indican un modelo estable. Se incluyen lineas de media y desviacion estandar.",
    )

    # Heatmap de coeficientes
    add_figure_placeholder(doc,
        "Heatmap de coeficientes del modelo organizado por grupo de features.\n"
        "Archivo: backend/app/ml/visualizations/coefficients_heatmap.png",
        "Mapa de calor de coeficientes del modelo por grupo de features",
        "Organiza los 18 coeficientes en 4 filas (CV Scores, Pesos Inst., Interacciones, Contexto). Los colores verde/rojo indican el signo y magnitud de cada coeficiente.",
    )

    # Tabla resumen de visualizaciones
    add_formatted_table(doc,
        headers=["#", "Visualizacion", "Archivo", "Proposito"],
        rows=[
            ["1", "Importancia de Features", "feature_importance.png", "Coeficientes del modelo ordenados por magnitud"],
            ["2", "Predicciones vs Reales", "predictions_vs_actual.png", "Precision del modelo (scatter plot con diagonal)"],
            ["3", "Distribucion de Residuos", "residuals_distribution.png", "Verificar ausencia de sesgo y homocedasticidad"],
            ["4", "Matriz de Confusion", "confusion_matrix.png", "Errores de clasificacion por clase"],
            ["5", "Scores por Clase", "score_distribution_by_class.png", "Boxplots de scores predichos por clase real"],
            ["6", "Cross-Validation", "cv_scores_distribution.png", "Estabilidad del modelo (5-fold)"],
            ["7", "Curva de Aprendizaje", "learning_curve.png", "Diagnostico over/underfitting"],
            ["8", "Heatmap Coeficientes", "coefficients_heatmap.png", "Mapa de coeficientes por grupo de features"],
        ],
        caption="Visualizaciones generadas por el modulo de evaluacion",
        nota="Todas las visualizaciones se generan automaticamente con el modulo ModelVisualizer. DPI: 150. Estilo: seaborn whitegrid.",
    )

    # ================================================================
    # 3.4.3. INTEGRACION CON EL BACKEND
    # ================================================================

    doc.add_heading("3.4.3. Integracion del modelo con el backend", level=2)

    add_paragraph_text(doc,
        "El servicio MLIntegrationService actua como capa de integracion entre el modelo ML "
        "y la API REST de FastAPI. Implementa el patron Singleton para garantizar una unica "
        "instancia del modelo en memoria, con cache de perfiles institucionales (TTL: 5 minutos) "
        "y persistencia de evaluaciones en Supabase."
    )

    # Diagrama de flujo completo
    add_figure_placeholder(doc,
        "Diagrama de flujo del proceso completo de evaluacion:\n"
        "1. Usuario sube CV (PDF) -> Frontend\n"
        "2. Frontend envia PDF en base64 -> API /evaluate-cv\n"
        "3. pdfplumber extrae texto del PDF\n"
        "4. Gemini extrae JSON estructurado del texto\n"
        "5. FeatureExtractor genera vector de 18 dimensiones\n"
        "6. Score heuristico ponderado (suma ponderada)\n"
        "7. Clasificacion: APTO / CONSIDERADO / NO_APTO\n"
        "8. Guardar evaluacion en Supabase\n"
        "9. Retornar resultado con explicacion al Frontend",
        "Flujo completo de evaluacion de CV",
        "El flujo integra LLM (Gemini), Feature Engineering (5 scorers), scoring heuristico ponderado y persistencia en base de datos. El resultado incluye score, clasificacion, fortalezas y debilidades.",
    )

    add_code_block(doc, '''# ml_integration_service.py - Servicio de integracion (Singleton)
class MLIntegrationService:
    _instance = None
    _predictor = None
    _profile_cache = {}
    _cache_ttl = timedelta(minutes=5)

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def evaluate_cv(self, gemini_output, institutional_config):
        # Extraer features (18 dimensiones)
        extractor = FeatureExtractor()
        features = extractor.extract_features(gemini_output, institutional_config)

        # Calcular score heuristico (suma ponderada)
        heuristic_score = extractor.calculate_weighted_score(features)

        # Clasificar segun umbrales institucionales
        classification = extractor.classify(
            heuristic_score, institutional_config.get('thresholds')
        )

        return {
            'match_score': heuristic_score,
            'classification': classification,
            'cv_scores': features['cv_scores'],
            'top_strengths': top_3_scores,
            'top_weaknesses': bottom_3_scores
        }

    def get_recommendations(self, gemini_output, top_n=5):
        profiles = self.load_all_active_profiles()
        recommendations = []
        for profile in profiles:
            result = self.evaluate_cv(gemini_output, profile)
            recommendations.append({...result, 'institution_name': profile['name']})
        recommendations.sort(key=lambda x: x['match_score'], reverse=True)
        return recommendations[:top_n]''',
        caption="Codigo fuente del servicio de integracion ML (ml_integration_service.py)",
        nota="El servicio Singleton gestiona el ciclo completo: extraccion del CV, feature engineering, scoring, clasificacion, generacion de recomendaciones y persistencia. El cache de perfiles se invalida automaticamente cada 5 minutos o manualmente al crear/editar perfiles.",
    )

    # Tabla de endpoints de la API
    add_formatted_table(doc,
        headers=["Endpoint", "Metodo", "Descripcion", "Requiere"],
        rows=[
            ["/api/ml/evaluate-cv", "POST", "Evalua un CV contra un perfil institucional", "Auth + PDF base64 + profile_id"],
            ["/api/ml/recommendations", "POST", "Obtiene top N recomendaciones de instituciones", "Auth + PDF base64"],
            ["/api/ml/model-info", "GET", "Informacion del modelo cargado", "Auth (admin)"],
        ],
        caption="Endpoints de la API REST para el modulo ML",
        nota="Todos los endpoints requieren autenticacion JWT. El endpoint de evaluacion recibe el PDF codificado en base64 y retorna score, clasificacion y explicacion.",
    )

    # ================================================================
    # 3.4.4. RESUMEN Y CONTRIBUCION
    # ================================================================

    doc.add_heading("3.4.4. Resumen del aporte ingenieril", level=2)

    add_paragraph_text(doc,
        "El modulo de Machine Learning constituye el aporte ingenieril principal de este trabajo "
        "de grado. A continuacion se resume la contribucion tecnica en terminos de componentes "
        "desarrollados, decisiones de diseno y resultados obtenidos."
    )

    add_formatted_table(doc,
        headers=["Componente", "Archivos", "LOC Aprox.", "Contribucion"],
        rows=[
            ["LLM Extractor", "llm_extractor.py", "~80", "Transformacion de texto a JSON con Google Gemini"],
            ["Feature Engineering", "5 scorers + feature_extractor.py", "~600", "5 algoritmos de scoring + vector de 18 dimensiones"],
            ["Dataset Sintetico", "synthetic_generator.py", "~370", "15 reglas expertas + distribuciones estadisticas"],
            ["Modelo Ridge", "ridge_model.py + model_trainer.py", "~470", "Entrenamiento con Grid Search + CV + explicabilidad"],
            ["Predictor", "predictor.py", "~315", "Interfaz de prediccion con explicaciones detalladas"],
            ["Evaluacion", "metrics.py + visualizations.py", "~740", "8 visualizaciones + reporte completo de metricas"],
            ["Integracion", "ml_integration_service.py", "~540", "Singleton con cache + persistencia + recomendaciones"],
            ["Total", "15+ archivos", "~3,100+", "Pipeline completo de ML para intermediacion laboral"],
        ],
        caption="Resumen de componentes del modulo de Machine Learning",
        nota="LOC = Lineas de codigo (aproximadas). El modulo es auto-contenido, con modulos desacoplados que permiten la evolucion independiente de cada componente.",
    )

    # Tabla de decisiones de diseno
    add_formatted_table(doc,
        headers=["Decision", "Alternativa Descartada", "Justificacion"],
        rows=[
            ["Ridge Regression", "Random Forest, SVM, MLP", "Interpretabilidad + regularizacion + eficiencia"],
            ["Dataset sintetico (5000)", "Datos reales (no disponibles)", "Reglas expertas capturan dominio + reproducibilidad"],
            ["Features de interaccion", "Solo features base", "Capturan alineacion candidato-institucion"],
            ["Score heuristico ponderado", "Solo modelo ML", "Mayor granularidad y control sobre los resultados"],
            ["Google Gemini (LLM)", "spaCy NER manual", "Extraction mas rica y flexible de CVs heterogeneos"],
            ["StandardScaler", "Sin normalizacion", "Features de diferentes escalas (scores vs anios)"],
            ["Singleton + Cache", "Instancia nueva por request", "Eficiencia: modelo cargado 1 vez en memoria"],
            ["Clasificacion 3 clases", "Binaria (apto/no apto)", "Mayor granularidad para la toma de decisiones"],
        ],
        caption="Principales decisiones de diseno del modulo ML",
        nota="Cada decision fue tomada en base a los requisitos del proyecto, las restricciones de datos disponibles y los criterios de calidad definidos.",
    )


# ============================================================
# MAIN - GENERADOR
# ============================================================

def main():
    print("=" * 70)
    print("GENERANDO DOCUMENTACION - SECCION 3.4: ML")
    print("=" * 70)

    doc = Document()

    # Configurar margenes
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)

    # Generar seccion 3.4
    generate_section_3_4(doc)

    # Guardar documento
    output_path = os.path.join(os.path.dirname(__file__), "Seccion_3_4_ML.docx")
    doc.save(output_path)

    print(f"\nDocumento generado exitosamente: {output_path}")
    print(f"Tablas generadas: {tabla_counter['value']}")
    print(f"Figuras generadas: {figura_counter['value']}")
    print(f"Total elementos: {tabla_counter['value'] + figura_counter['value']}")


if __name__ == "__main__":
    main()
