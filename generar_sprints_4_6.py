"""
Generador de Documentación de Sprints 4-6 - Trabajo de Grado
Autor: Ivan Condori Choquehuanca
Genera un archivo .docx con la documentación completa de los Sprints 4-6
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

AUTOR = "Ivan Condori Choquehuanca"

SPRINTS = {
    4: {"nombre": "Gestión de Oferta Laboral", "inicio": "10/11/2025", "fin": "30/11/2025"},
    5: {"nombre": "Perfiles Institucionales", "inicio": "01/12/2025", "fin": "21/12/2025"},
    6: {"nombre": "Informes y Reportes", "inicio": "13/01/2026", "fin": "08/02/2026"},
}

# Contadores globales (continúan desde Sprints 1-3: 25 tablas, 26 figuras)
tabla_counter = {"value": 25}
figura_counter = {"value": 26}


# ============================================================
# FUNCIONES AUXILIARES (idénticas al script 1-3)
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_table(doc, headers, rows, sprint_num, caption, nota, fuente="Elaboración propia"):
    tabla_counter["value"] += 1
    num = tabla_counter["value"]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(table)

    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Tabla {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_figure_placeholder(doc, descripcion_captura, caption, nota, fuente="Elaboración propia"):
    figura_counter["value"] += 1
    num = figura_counter["value"]

    p_box = doc.add_paragraph()
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_box.add_run(f"[INSERTAR CAPTURA DE PANTALLA]\n{descripcion_captura}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 0, 0)
    run.italic = True

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_code_block(doc, code_text, caption, nota, fuente="Elaboración propia"):
    figura_counter["value"] += 1
    num = figura_counter["value"]

    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Consolas" w:hAnsi="Consolas" w:cs="Consolas"/>')
        rPr.append(rFonts)

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_heading_sprint(doc, sprint_num):
    h = doc.add_heading(f"3.5.{sprint_num}. Sprint {sprint_num}: {SPRINTS[sprint_num]['nombre']}", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)


def add_subsection(doc, sprint_num, sub_idx, title, level=2):
    prefix = f"3.5.{sprint_num}.{sub_idx}"
    h = doc.add_heading(f"{prefix}. {title}", level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)


# ============================================================
# SPRINT 4: GESTIÓN DE OFERTA LABORAL
# ============================================================

def generar_sprint4(doc):
    sprint_num = 4
    add_heading_sprint(doc, sprint_num)

    doc.add_paragraph(
        f"El Sprint 4 se desarrolló del {SPRINTS[4]['inicio']} al {SPRINTS[4]['fin']} "
        "y se centró en la implementación del módulo de Gestión de Ofertas Laborales, que permite "
        "a los administradores y operadores crear, editar, activar, desactivar y consultar ofertas "
        "de pasantías y empleos. Las ofertas se vinculan a perfiles institucionales y constituyen "
        "la base sobre la cual el sistema de recomendaciones del Sprint posterior opera. "
        "Se implementó soft delete, filtrado por tipo/sector/estado, paginación y estadísticas."
    )

    # --- 1. Sprint Backlog ---
    add_subsection(doc, 4, 1, "Sprint Backlog")
    doc.add_paragraph(
        "A continuación se presenta el Sprint Backlog con las historias de usuario "
        "y tareas planificadas para el módulo de Gestión de Ofertas Laborales."
    )

    backlog_headers = ["ID", "Historia de Usuario", "Tarea", "Prioridad", "Estado", "Responsable", "Est. (h)"]
    backlog_rows = [
        ["HU-17", "Como admin, quiero crear ofertas laborales para publicarlas", "Diseñar formulario de creación de oferta (OfertaFormView)", "Alta", "Completado", AUTOR, "12"],
        ["HU-17", "", "Implementar endpoint POST /api/admin/ofertas", "Alta", "Completado", AUTOR, "10"],
        ["HU-17", "", "Implementar OfertaService.create_oferta() con Singleton", "Alta", "Completado", AUTOR, "8"],
        ["HU-17", "", "Crear tabla 'ofertas_laborales' en Supabase", "Alta", "Completado", AUTOR, "4"],
        ["HU-17", "", "Implementar vinculación con perfil institucional (FK)", "Alta", "Completado", AUTOR, "6"],
        ["HU-18", "Como admin, quiero listar ofertas con filtros y paginación", "Implementar endpoint GET /api/admin/ofertas con filtros", "Alta", "Completado", AUTOR, "12"],
        ["HU-18", "", "Implementar filtrado por tipo (pasantía/empleo)", "Alta", "Completado", AUTOR, "4"],
        ["HU-18", "", "Implementar filtrado por sector y estado activo", "Media", "Completado", AUTOR, "4"],
        ["HU-18", "", "Implementar exclusión de ofertas expiradas por fecha_cierre", "Media", "Completado", AUTOR, "4"],
        ["HU-18", "", "Diseñar vista OfertasAdminView con tabla y filtros", "Alta", "Completado", AUTOR, "14"],
        ["HU-19", "Como admin, quiero editar ofertas existentes", "Implementar endpoint PUT /api/admin/ofertas/{id}", "Media", "Completado", AUTOR, "8"],
        ["HU-19", "", "Implementar modo edición en OfertaFormView", "Media", "Completado", AUTOR, "8"],
        ["HU-20", "Como admin, quiero desactivar/reactivar ofertas", "Implementar soft delete DELETE /api/admin/ofertas/{id}", "Media", "Completado", AUTOR, "6"],
        ["HU-20", "", "Implementar reactivación POST /api/admin/ofertas/{id}/activate", "Media", "Completado", AUTOR, "4"],
        ["HU-20", "", "Implementar botones activar/desactivar en la vista", "Media", "Completado", AUTOR, "4"],
        ["HU-21", "Como admin, quiero ver estadísticas de ofertas", "Implementar endpoint GET /api/admin/ofertas/stats/summary", "Baja", "Completado", AUTOR, "6"],
        ["HU-21", "", "Implementar tarjetas de estadísticas en OfertasAdminView", "Baja", "Completado", AUTOR, "6"],
        ["HU-22", "Como sistema, necesito enriquecer ofertas con datos institucionales", "Implementar _enrich_oferta() con JOIN a institutional_profiles", "Media", "Completado", AUTOR, "6"],
        ["HU-22", "", "Implementar filtrado de ofertas por rol de usuario", "Media", "Completado", AUTOR, "6"],
        ["HU-22", "", "Implementar API frontend (ofertas.api.js)", "Alta", "Completado", AUTOR, "4"],
    ]

    add_formatted_table(doc, backlog_headers, backlog_rows, 4,
                        "Sprint Backlog – Sprint 4: Gestión de Oferta Laboral",
                        "Se detallan las historias de usuario y tareas técnicas para el CRUD de ofertas laborales, "
                        "incluyendo filtrado, paginación, soft delete, estadísticas y vinculación con perfiles institucionales.")

    # --- 2. Bitácora de temas importantes ---
    add_subsection(doc, 4, 2, "Bitácora de Identificación de Temas Importantes")

    bitacora_headers = ["Fecha", "Tema Identificado", "Impacto", "Decisión Tomada"]
    bitacora_rows = [
        ["10/11/2025", "Tipos de oferta: pasantía vs empleo", "Alto", "Se definieron dos tipos: 'pasantia' y 'empleo'. Los estudiantes solo ven pasantías y los titulados solo ven empleos en recomendaciones."],
        ["12/11/2025", "Estrategia de eliminación: soft vs hard delete", "Alto", "Se implementó soft delete (is_active=false) para preservar integridad referencial con evaluaciones y recomendaciones existentes."],
        ["15/11/2025", "Vinculación con perfiles institucionales", "Alto", "Cada oferta se vincula opcionalmente a un perfil institucional mediante FK. El JOIN enriquece la oferta con institution_name y sector."],
        ["18/11/2025", "Gestión de vigencia temporal", "Medio", "Se implementaron campos fecha_inicio y fecha_cierre. Las ofertas expiradas se excluyen por defecto pero pueden incluirse con include_expired=true."],
        ["21/11/2025", "Modalidad de trabajo", "Medio", "Se definieron 3 modalidades: presencial, remoto, híbrido. Campo opcional que enriquece la información de la oferta."],
        ["25/11/2025", "Campos de la oferta laboral", "Medio", "Se definieron: titulo, descripcion, tipo, modalidad, ubicacion, requisitos_especificos (JSONB), cupos_disponibles, fecha_inicio, fecha_cierre."],
        ["28/11/2025", "Estadísticas del dashboard de ofertas", "Bajo", "Se implementaron 5 métricas: pasantías activas, empleos activos, ofertas expiradas, ofertas inactivas y total activas."],
    ]

    add_formatted_table(doc, bitacora_headers, bitacora_rows, 4,
                        "Bitácora de temas importantes – Sprint 4",
                        "Registro de decisiones técnicas durante el desarrollo del módulo de ofertas laborales, "
                        "incluyendo tipos, soft delete, vinculación institucional y gestión de vigencia.")

    # --- 3. Actualización tabla de tareas ---
    add_subsection(doc, 4, 3, "Actualización de la Tabla de Tareas")

    tareas_headers = ["Tarea", "Estado", "Horas Est.", "Horas Real", "Observaciones"]
    tareas_rows = [
        ["Crear tabla ofertas_laborales en Supabase", "Completado", "4", "3", "Incluye FK a institutional_profiles"],
        ["OfertaService (Singleton completo)", "Completado", "26", "28", "CRUD + estadísticas + enrich"],
        ["Endpoint POST /api/admin/ofertas", "Completado", "10", "10", "Creación con vinculación institucional"],
        ["Endpoint GET /api/admin/ofertas (listado)", "Completado", "12", "14", "Filtros + paginación + JOIN"],
        ["Endpoint PUT /api/admin/ofertas/{id}", "Completado", "8", "7", "Solo campos permitidos"],
        ["Endpoint DELETE (soft delete)", "Completado", "6", "5", "Marca is_active=false"],
        ["Endpoint POST /activate", "Completado", "4", "4", "Reactivación"],
        ["Endpoint GET /stats/summary", "Completado", "6", "7", "5 métricas agregadas"],
        ["OfertaFormView (crear/editar)", "Completado", "20", "22", "Dropdown de perfiles institucionales"],
        ["OfertasAdminView (listado)", "Completado", "14", "16", "Tabla + filtros + tarjetas stats"],
        ["ofertas.api.js (API frontend)", "Completado", "4", "3", "7 funciones de API"],
        ["Filtrado por rol de usuario", "Completado", "6", "6", "estudiante→pasantía, titulado→empleo"],
    ]

    add_formatted_table(doc, tareas_headers, tareas_rows, 4,
                        "Actualización de tareas – Sprint 4",
                        "Total estimado: 120h, Total real: 125h. Desviación: +4.2%. "
                        "El listado con filtros y la vista frontend requirieron ajustes adicionales.")

    # --- 4. Diagrama de caso de uso expandido ---
    add_subsection(doc, 4, 4, "Diagrama de Caso de Uso Expandido del Módulo")

    add_figure_placeholder(doc,
        "Tomar captura del diagrama de caso de uso del módulo de Gestión de Ofertas.\n"
        "Actores: Administrador/Operador, Sistema.\n"
        "Casos de uso: Crear Oferta, Listar Ofertas (con filtros), Editar Oferta,\n"
        "Desactivar Oferta (soft delete), Reactivar Oferta, Ver Estadísticas de Ofertas,\n"
        "Vincular a Perfil Institucional, Filtrar por Tipo/Sector/Estado,\n"
        "Enriquecer Oferta con Datos Institucionales (sistema).",
        "Diagrama de caso de uso expandido – Módulo de Gestión de Ofertas Laborales",
        "El diagrama muestra las operaciones CRUD sobre ofertas laborales, incluyendo "
        "la vinculación con perfiles institucionales y el filtrado por tipo, sector y estado.")

    cu_headers = ["Campo", "Descripción"]
    cu_rows = [
        ["Caso de Uso", "CU-17: Crear Oferta Laboral"],
        ["Actor Principal", "Administrador / Operador"],
        ["Precondiciones", "El usuario tiene rol administrador u operador. Opcionalmente existe un perfil institucional para vincular."],
        ["Postcondiciones", "Se crea un registro en ofertas_laborales con is_active=true, vinculado opcionalmente a un perfil institucional."],
        ["Flujo Principal", "1. El admin accede al formulario de nueva oferta\n"
                           "2. Completa los campos: título (requerido), tipo (requerido: pasantía/empleo)\n"
                           "3. Opcionalmente: descripción, modalidad, ubicación, fechas, cupos\n"
                           "4. Selecciona un perfil institucional del dropdown (opcional)\n"
                           "5. El sistema valida los campos requeridos\n"
                           "6. Se inserta el registro con created_by = user_id del admin\n"
                           "7. Se enriquece la respuesta con datos del perfil institucional (JOIN)\n"
                           "8. Se redirige al listado de ofertas"],
        ["Flujo Alternativo", "2a. Título vacío → Error 400: campo requerido\n"
                              "2b. Tipo inválido → Error 400: debe ser pasantia o empleo\n"
                              "4a. Perfil institucional no existe → Se crea sin vinculación"],
    ]

    add_formatted_table(doc, cu_headers, cu_rows, 4,
                        "Caso de uso expandido CU-17: Crear Oferta Laboral",
                        "Descripción del flujo para la creación de una oferta laboral con vinculación "
                        "opcional a un perfil institucional y asignación automática del creador.")

    # --- 5. Diseño de colección ---
    add_subsection(doc, 4, 5, "Diseño de la Colección y Documentos")

    doc.add_paragraph(
        "El Sprint 4 introduce la tabla 'ofertas_laborales' que almacena las ofertas "
        "de pasantías y empleos publicadas por los administradores."
    )

    col_headers = ["Campo", "Tipo", "Restricciones", "Descripción"]
    col_rows = [
        ["id", "UUID", "PRIMARY KEY, auto-generado", "Identificador único de la oferta"],
        ["institutional_profile_id", "UUID", "FK → institutional_profiles(id), NULLABLE", "Perfil institucional vinculado (opcional)"],
        ["titulo", "TEXT", "NOT NULL", "Título de la oferta laboral"],
        ["descripcion", "TEXT", "NULLABLE", "Descripción detallada de la oferta"],
        ["tipo", "TEXT", "NOT NULL", "Tipo: 'pasantia' o 'empleo'"],
        ["modalidad", "TEXT", "NULLABLE", "Modalidad: 'presencial', 'remoto', 'hibrido'"],
        ["ubicacion", "TEXT", "NULLABLE", "Ubicación geográfica del puesto"],
        ["requisitos_especificos", "JSONB", "DEFAULT '{}'", "Requisitos adicionales en formato flexible"],
        ["is_active", "BOOLEAN", "DEFAULT true", "Estado activo/inactivo (soft delete)"],
        ["fecha_inicio", "DATE", "NULLABLE", "Fecha de inicio de vigencia"],
        ["fecha_cierre", "DATE", "NULLABLE", "Fecha de cierre de la oferta"],
        ["cupos_disponibles", "INTEGER", "DEFAULT 1", "Número de vacantes disponibles"],
        ["created_by", "UUID", "FK → usuarios(id)", "Administrador que creó la oferta"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de creación"],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de última actualización"],
    ]

    add_formatted_table(doc, col_headers, col_rows, 4,
                        "Estructura de la tabla 'ofertas_laborales'",
                        "Tabla que almacena las ofertas de pasantías y empleos. Utiliza soft delete (is_active) "
                        "para preservar integridad referencial. Se vincula opcionalmente a institutional_profiles "
                        "para heredar sector y nombre de institución mediante JOIN.")

    # --- 6. Mockups ---
    add_subsection(doc, 4, 6, "Diseño e Interfaz del Módulo (Mockups)")

    add_figure_placeholder(doc,
        "Tomar captura de OfertasAdminView.vue.\n"
        "Debe mostrar: tarjetas de estadísticas (pasantías activas, empleos activos,\n"
        "expiradas, inactivas), filtros (tipo, estado, incluir expiradas),\n"
        "tabla con columnas (Oferta, Tipo, Institución, Estado, Vigencia, Acciones),\n"
        "paginación inferior.",
        "Vista administrativa de gestión de ofertas laborales",
        "Panel principal de gestión de ofertas que muestra estadísticas resumidas en tarjetas, "
        "filtros por tipo y estado, tabla paginada con acciones de edición y activación/desactivación.")

    add_figure_placeholder(doc,
        "Tomar captura de OfertaFormView.vue en modo CREACIÓN.\n"
        "Debe mostrar: campos título, descripción, tipo (dropdown pasantía/empleo),\n"
        "modalidad, ubicación, perfil institucional (dropdown), fecha inicio,\n"
        "fecha cierre, cupos disponibles, botón Guardar.",
        "Formulario de creación de oferta laboral",
        "Formulario para crear una nueva oferta laboral con selección de tipo, modalidad, "
        "vinculación a perfil institucional y configuración de vigencia temporal.")

    add_figure_placeholder(doc,
        "Tomar captura de OfertaFormView.vue en modo EDICIÓN.\n"
        "Debe mostrar: los mismos campos prellenados con datos existentes,\n"
        "toggle de estado activo/inactivo visible solo en edición.",
        "Formulario de edición de oferta laboral existente",
        "Formulario de edición con los datos precargados de la oferta seleccionada. "
        "En modo edición se habilita el toggle de estado activo/inactivo.")

    # --- 7. Código fuente ---
    add_subsection(doc, 4, 7, "Código Fuente del Módulo (Incremento del Producto)")

    code_create = '''@router.post("", response_model=OfertaLaboralResponse, status_code=201)
async def create_oferta(
    data: OfertaLaboralCreate,
    admin_user: dict = Depends(verify_operator_access)
):
    oferta_service = get_oferta_service()
    oferta = oferta_service.create_oferta(
        data.model_dump(),
        created_by=admin_user['user_id']
    )
    return OfertaLaboralResponse(
        id=oferta['id'],
        institutional_profile_id=oferta.get('institutional_profile_id'),
        institution_name=oferta.get('institution_name'),
        sector=oferta.get('sector'),
        titulo=oferta['titulo'],
        tipo=oferta['tipo'],
        modalidad=oferta.get('modalidad'),
        is_active=oferta['is_active'],
        fecha_inicio=oferta.get('fecha_inicio'),
        fecha_cierre=oferta.get('fecha_cierre'),
        cupos_disponibles=oferta.get('cupos_disponibles', 1),
        created_by=oferta.get('created_by'), ...)'''

    add_code_block(doc, code_create,
                   "Endpoint de creación de oferta laboral (ofertas.py)",
                   "Endpoint POST /api/admin/ofertas que crea una nueva oferta vinculada opcionalmente "
                   "a un perfil institucional. Solo accesible para administradores y operadores.")

    code_service = '''class OfertaService:
    _instance = None  # Singleton

    def create_oferta(self, data: Dict, created_by: str) -> Dict:
        oferta_data = {
            'institutional_profile_id': data.get('institutional_profile_id'),
            'titulo': data['titulo'],
            'descripcion': data.get('descripcion'),
            'tipo': data['tipo'],  # 'pasantia' o 'empleo'
            'modalidad': data.get('modalidad'),
            'ubicacion': data.get('ubicacion'),
            'requisitos_especificos': data.get('requisitos_especificos', {}),
            'is_active': True,
            'fecha_inicio': data.get('fecha_inicio'),
            'fecha_cierre': data.get('fecha_cierre'),
            'cupos_disponibles': data.get('cupos_disponibles', 1),
            'created_by': created_by
        }
        response = supabase.table("ofertas_laborales")\\
            .insert(oferta_data).execute()
        return self._enrich_oferta(response.data[0])

    def list_ofertas(self, tipo=None, is_active=None, sector=None,
                     include_expired=False, page=1, page_size=20):
        query = supabase.table("ofertas_laborales")\\
            .select("*, institutional_profiles(institution_name, sector)",
                    count="exact")
        if tipo: query = query.eq("tipo", tipo)
        if is_active is not None: query = query.eq("is_active", is_active)
        if not include_expired:
            today = date.today().isoformat()
            query = query.or_(f"fecha_cierre.is.null,fecha_cierre.gte.{today}")
        query = query.order("created_at", desc=True)
        offset = (page - 1) * page_size
        query = query.range(offset, offset + page_size - 1)
        response = query.execute()
        return {'ofertas': [self._enrich_oferta(o) for o in response.data],
                'total': response.count, ...}

    def delete_oferta(self, oferta_id: str) -> bool:
        """Soft delete: marca is_active=False"""
        response = supabase.table("ofertas_laborales")\\
            .update({'is_active': False,
                     'updated_at': datetime.utcnow().isoformat()})\\
            .eq("id", oferta_id).execute()
        return bool(response.data)

    def _enrich_oferta(self, oferta: Dict) -> Dict:
        inst_profile = oferta.pop('institutional_profiles', None) or {}
        return {**oferta,
                'institution_name': inst_profile.get('institution_name'),
                'sector': inst_profile.get('sector')}'''

    add_code_block(doc, code_service,
                   "OfertaService: servicio CRUD de ofertas laborales (oferta_service.py)",
                   "Servicio Singleton que gestiona el ciclo de vida de las ofertas laborales, incluyendo "
                   "creación con vinculación institucional, listado con filtros, soft delete y enriquecimiento "
                   "de datos mediante JOIN con la tabla institutional_profiles.")

    code_stats = '''def get_statistics(self) -> Dict:
    pasantias = supabase.table("ofertas_laborales")\\
        .select("id", count="exact")\\
        .eq("tipo", "pasantia").eq("is_active", True).execute()
    empleos = supabase.table("ofertas_laborales")\\
        .select("id", count="exact")\\
        .eq("tipo", "empleo").eq("is_active", True).execute()
    today = date.today().isoformat()
    expiradas = supabase.table("ofertas_laborales")\\
        .select("id", count="exact")\\
        .lt("fecha_cierre", today).execute()
    inactivas = supabase.table("ofertas_laborales")\\
        .select("id", count="exact")\\
        .eq("is_active", False).execute()
    return {
        'pasantias_activas': pasantias.count or 0,
        'empleos_activos': empleos.count or 0,
        'ofertas_expiradas': expiradas.count or 0,
        'ofertas_inactivas': inactivas.count or 0,
        'total_activas': (pasantias.count or 0) + (empleos.count or 0)
    }'''

    add_code_block(doc, code_stats,
                   "Estadísticas de ofertas laborales (oferta_service.py)",
                   "Método que calcula métricas agregadas de ofertas laborales: pasantías activas, "
                   "empleos activos, ofertas expiradas por fecha de cierre y ofertas desactivadas.")

    # --- 8. Pruebas unitarias ---
    add_subsection(doc, 4, 8, "Resultado de Pruebas Unitarias")

    pruebas_headers = ["ID Prueba", "Descripción", "Entrada", "Resultado Esperado", "Estado"]
    pruebas_rows = [
        ["PU-4.01", "Crear oferta de pasantía", "titulo, tipo='pasantia', datos válidos", "Oferta creada + status 201 + is_active=true", "Aprobado"],
        ["PU-4.02", "Crear oferta de empleo con perfil institucional", "tipo='empleo' + institutional_profile_id", "Oferta con institution_name y sector enriched", "Aprobado"],
        ["PU-4.03", "Crear oferta sin título (requerido)", "tipo válido, titulo vacío", "Error 400: campo requerido", "Aprobado"],
        ["PU-4.04", "Listar ofertas sin filtros", "GET /api/admin/ofertas", "Lista paginada con ofertas activas no expiradas", "Aprobado"],
        ["PU-4.05", "Listar ofertas filtradas por tipo", "tipo='pasantia'", "Solo ofertas de tipo pasantía", "Aprobado"],
        ["PU-4.06", "Listar ofertas incluyendo expiradas", "include_expired=true", "Incluye ofertas con fecha_cierre pasada", "Aprobado"],
        ["PU-4.07", "Editar oferta existente", "PUT con nuevo título y descripción", "Campos actualizados + updated_at cambiado", "Aprobado"],
        ["PU-4.08", "Desactivar oferta (soft delete)", "DELETE /api/admin/ofertas/{id}", "is_active=false, registro preservado", "Aprobado"],
        ["PU-4.09", "Reactivar oferta desactivada", "POST /api/admin/ofertas/{id}/activate", "is_active=true", "Aprobado"],
        ["PU-4.10", "Obtener estadísticas", "GET /api/admin/ofertas/stats/summary", "JSON con 5 métricas numéricas", "Aprobado"],
        ["PU-4.11", "Acceso sin rol admin/operador", "Token con rol='estudiante'", "Error 403: Acceso denegado", "Aprobado"],
        ["PU-4.12", "Oferta no encontrada", "GET con ID inexistente", "Error 404: Oferta no encontrada", "Aprobado"],
    ]

    add_formatted_table(doc, pruebas_headers, pruebas_rows, 4,
                        "Resultados de pruebas unitarias – Sprint 4",
                        "Se ejecutaron 12 pruebas unitarias cubriendo CRUD completo, filtrado, soft delete, "
                        "estadísticas y control de acceso. Todas las pruebas pasaron exitosamente.")

    add_figure_placeholder(doc,
        "Tomar captura de la terminal ejecutando las pruebas del Sprint 4.\n"
        "Puede ser pytest o pruebas manuales con Postman/Thunder Client\n"
        "mostrando los endpoints respondiendo correctamente.",
        "Ejecución de pruebas unitarias del Sprint 4",
        "Captura de la terminal mostrando la ejecución exitosa de las pruebas unitarias "
        "del módulo de Gestión de Ofertas Laborales.")

    # --- 9. Bitácora de cambios ---
    add_subsection(doc, 4, 9, "Bitácora de Cambios al Sprint Backlog")

    cambios_headers = ["Fecha", "Cambio", "Motivo", "Impacto"]
    cambios_rows = [
        ["12/11/2025", "Se añadió soft delete en lugar de hard delete", "Preservar integridad con evaluaciones y recomendaciones existentes", "Medio: +6h (endpoint activate + lógica adicional)"],
        ["18/11/2025", "Se añadió endpoint de estadísticas /stats/summary", "El dashboard admin requiere métricas resumidas de ofertas", "Bajo: +12h (endpoint + tarjetas frontend)"],
        ["22/11/2025", "Se añadió campo 'modalidad' (presencial/remoto/híbrido)", "Requisito identificado durante revisión con tutor", "Bajo: +4h (campo en BD + formulario)"],
        ["26/11/2025", "Se implementó enriquecimiento con JOIN", "Las ofertas necesitan mostrar el nombre de la institución sin consulta extra", "Medio: +6h (_enrich_oferta + JOIN en queries)"],
    ]

    add_formatted_table(doc, cambios_headers, cambios_rows, 4,
                        "Bitácora de cambios al Sprint Backlog – Sprint 4",
                        "Los cambios principales se debieron a la estrategia de soft delete y a requisitos "
                        "adicionales de estadísticas y modalidad de trabajo.")

    doc.add_page_break()


# ============================================================
# SPRINT 5: PERFILES INSTITUCIONALES
# ============================================================

def generar_sprint5(doc):
    sprint_num = 5
    add_heading_sprint(doc, sprint_num)

    doc.add_paragraph(
        f"El Sprint 5 se desarrolló del {SPRINTS[5]['inicio']} al {SPRINTS[5]['fin']} "
        "y se centró en la implementación del módulo de Perfiles Institucionales, que permite "
        "a los administradores crear y gestionar los perfiles de evaluación que definen los pesos, "
        "requisitos y umbrales de clasificación para cada institución/empresa. Estos perfiles "
        "constituyen la configuración central del sistema de matching ML, ya que determinan qué "
        "habilidades, nivel educativo, experiencia e idiomas se priorizan al evaluar candidatos."
    )

    # --- 1. Sprint Backlog ---
    add_subsection(doc, 5, 1, "Sprint Backlog")

    backlog_headers = ["ID", "Historia de Usuario", "Tarea", "Prioridad", "Estado", "Responsable", "Est. (h)"]
    backlog_rows = [
        ["HU-23", "Como admin, quiero crear perfiles institucionales de evaluación", "Diseñar ProfileFormView con secciones (info, pesos, umbrales, requisitos)", "Alta", "Completado", AUTOR, "16"],
        ["HU-23", "", "Implementar endpoint POST /api/admin/institutional-profiles", "Alta", "Completado", AUTOR, "12"],
        ["HU-23", "", "Implementar validación de pesos (deben sumar 1.0)", "Alta", "Completado", AUTOR, "4"],
        ["HU-23", "", "Implementar validación de umbrales (apto > considerado)", "Alta", "Completado", AUTOR, "3"],
        ["HU-23", "", "Crear tabla 'institutional_profiles' en Supabase", "Alta", "Completado", AUTOR, "4"],
        ["HU-23", "", "Implementar validación de nombre único de institución", "Media", "Completado", AUTOR, "3"],
        ["HU-24", "Como admin, quiero listar y buscar perfiles institucionales", "Implementar endpoint GET /api/admin/institutional-profiles", "Alta", "Completado", AUTOR, "8"],
        ["HU-24", "", "Implementar filtrado por sector y estado activo/inactivo", "Media", "Completado", AUTOR, "4"],
        ["HU-24", "", "Diseñar ProfilesAdminView con tarjetas y búsqueda", "Alta", "Completado", AUTOR, "14"],
        ["HU-25", "Como admin, quiero editar perfiles institucionales", "Implementar endpoint PUT /api/admin/institutional-profiles/{id}", "Media", "Completado", AUTOR, "10"],
        ["HU-25", "", "Implementar modo edición en ProfileFormView", "Media", "Completado", AUTOR, "8"],
        ["HU-25", "", "Implementar invalidación de caché ML al modificar perfil", "Alta", "Completado", AUTOR, "6"],
        ["HU-26", "Como admin, quiero desactivar/reactivar perfiles", "Implementar soft delete DELETE /api/admin/institutional-profiles/{id}", "Media", "Completado", AUTOR, "6"],
        ["HU-26", "", "Implementar reactivación POST /{id}/activate", "Media", "Completado", AUTOR, "4"],
        ["HU-27", "Como admin, quiero consultar sectores disponibles", "Implementar endpoint GET /api/admin/sectors", "Baja", "Completado", AUTOR, "4"],
        ["HU-27", "", "Implementar sectores predefinidos + opción 'Otro'", "Baja", "Completado", AUTOR, "4"],
        ["HU-28", "Como sistema, necesito esquemas de validación robustos", "Implementar InstitutionalWeights (validación suma=1.0)", "Alta", "Completado", AUTOR, "6"],
        ["HU-28", "", "Implementar InstitutionalRequirements (skills, educación, exp)", "Alta", "Completado", AUTOR, "6"],
        ["HU-28", "", "Implementar InstitutionalThresholds (validación apto>considerado)", "Alta", "Completado", AUTOR, "4"],
        ["HU-28", "", "Implementar API frontend (profiles.api.js)", "Alta", "Completado", AUTOR, "4"],
    ]

    add_formatted_table(doc, backlog_headers, backlog_rows, 5,
                        "Sprint Backlog – Sprint 5: Perfiles Institucionales",
                        "Se detallan las historias de usuario y tareas para el CRUD de perfiles institucionales, "
                        "incluyendo validación de pesos, umbrales, requisitos e invalidación de caché ML.")

    # --- 2. Bitácora de temas importantes ---
    add_subsection(doc, 5, 2, "Bitácora de Identificación de Temas Importantes")

    bitacora_headers = ["Fecha", "Tema Identificado", "Impacto", "Decisión Tomada"]
    bitacora_rows = [
        ["01/12/2025", "Validación de pesos institucionales", "Alto", "Los pesos de las 5 dimensiones (hard_skills, soft_skills, experience, education, languages) deben sumar exactamente 1.0 (tolerancia 0.01). Se implementó validador Pydantic."],
        ["03/12/2025", "Estructura de requisitos como JSONB", "Alto", "Se almacenan requisitos en JSONB: required_skills[], preferred_skills[], min_experience_years, required_education_level, required_languages[]. Permite flexibilidad sin migraciones."],
        ["06/12/2025", "Umbrales de clasificación", "Alto", "Cada perfil define sus propios umbrales: 'apto' (default 0.70) y 'considerado' (default 0.50). Se valida que apto > considerado para consistencia."],
        ["09/12/2025", "Invalidación de caché ML", "Alto", "Al crear, editar o eliminar un perfil institucional, se invalida la caché del MLIntegrationService (TTL 5 min) para que las evaluaciones usen datos actualizados."],
        ["12/12/2025", "Sectores predefinidos", "Medio", "Se definieron 11 sectores: Tecnología, Finanzas, Salud, Educación, Construcción, Comercio, Servicios, Industria, Gobierno, ONG, Otro. Con opción personalizada."],
        ["15/12/2025", "Unicidad del nombre de institución", "Medio", "Se valida que no existan dos perfiles con el mismo institution_name para evitar duplicados y confusión."],
        ["18/12/2025", "Autocompletado de skills y idiomas", "Bajo", "El formulario incluye listas de autocompletado con skills comunes (Python, React, SQL...) e idiomas frecuentes para agilizar la creación."],
    ]

    add_formatted_table(doc, bitacora_headers, bitacora_rows, 5,
                        "Bitácora de temas importantes – Sprint 5",
                        "Registro de decisiones técnicas para los perfiles institucionales, incluyendo validaciones "
                        "de pesos/umbrales, estructura JSONB y la invalidación de caché ML.")

    # --- 3. Actualización tabla de tareas ---
    add_subsection(doc, 5, 3, "Actualización de la Tabla de Tareas")

    tareas_headers = ["Tarea", "Estado", "Horas Est.", "Horas Real", "Observaciones"]
    tareas_rows = [
        ["Crear tabla institutional_profiles en Supabase", "Completado", "4", "4", "JSONB para weights, requirements, thresholds"],
        ["Endpoint POST (crear perfil)", "Completado", "12", "14", "Validación nombre único + pesos + umbrales"],
        ["Endpoint GET (listar perfiles)", "Completado", "8", "8", "Filtro por sector y activo/inactivo"],
        ["Endpoint GET /{id} (detalle)", "Completado", "4", "3", "—"],
        ["Endpoint PUT (editar perfil)", "Completado", "10", "12", "Validación nombre único al cambiar"],
        ["Endpoint DELETE (soft delete)", "Completado", "6", "5", "is_active=false"],
        ["Endpoint POST /{id}/activate", "Completado", "4", "4", "Validación de estado previo"],
        ["Endpoint GET /sectors", "Completado", "4", "3", "Lista sectores únicos activos"],
        ["Invalidación caché MLIntegrationService", "Completado", "6", "7", "invalidate_cache() en create/update/delete"],
        ["Schemas Pydantic (Weights, Requirements, Thresholds)", "Completado", "16", "18", "Validadores custom suma=1.0 y apto>considerado"],
        ["ProfileFormView (666 líneas)", "Completado", "24", "28", "4 secciones + autocompletado + normalización"],
        ["ProfilesAdminView (352 líneas)", "Completado", "14", "15", "Tarjetas + búsqueda + filtro sector"],
        ["profiles.api.js", "Completado", "4", "3", "7 funciones de API"],
    ]

    add_formatted_table(doc, tareas_headers, tareas_rows, 5,
                        "Actualización de tareas – Sprint 5",
                        "Total estimado: 116h, Total real: 124h. Desviación: +6.9%. "
                        "El formulario de creación/edición fue el componente más complejo del sprint.")

    # --- 4. Diagrama de caso de uso expandido ---
    add_subsection(doc, 5, 4, "Diagrama de Caso de Uso Expandido del Módulo")

    add_figure_placeholder(doc,
        "Tomar captura del diagrama de caso de uso del módulo de Perfiles Institucionales.\n"
        "Actores: Administrador/Operador, MLIntegrationService (sistema).\n"
        "Casos de uso: Crear Perfil Institucional, Listar Perfiles, Editar Perfil,\n"
        "Desactivar Perfil, Reactivar Perfil, Consultar Sectores,\n"
        "Validar Pesos (sistema), Validar Umbrales (sistema),\n"
        "Invalidar Caché ML (sistema).",
        "Diagrama de caso de uso expandido – Módulo de Perfiles Institucionales",
        "El diagrama muestra la gestión de perfiles de evaluación institucional, incluyendo "
        "validaciones de pesos y umbrales, y la invalidación automática de caché ML.")

    cu_headers = ["Campo", "Descripción"]
    cu_rows = [
        ["Caso de Uso", "CU-23: Crear Perfil Institucional"],
        ["Actor Principal", "Administrador / Operador"],
        ["Precondiciones", "El usuario tiene rol administrador u operador. No existe un perfil con el mismo nombre."],
        ["Postcondiciones", "Se crea un registro en institutional_profiles con pesos validados, umbrales configurados y caché ML invalidada."],
        ["Flujo Principal", "1. El admin accede al formulario de nuevo perfil institucional\n"
                           "2. Completa información básica: nombre (único), sector, descripción\n"
                           "3. Configura pesos de evaluación (5 dimensiones, deben sumar 100%)\n"
                           "4. Configura umbrales: apto (default 70%) y considerado (default 50%)\n"
                           "5. Define requisitos: skills requeridos/preferidos, educación, experiencia, idiomas\n"
                           "6. El sistema valida: nombre único, pesos=1.0, apto>considerado\n"
                           "7. Se inserta en BD con created_by y timestamps\n"
                           "8. Se invalida la caché del MLIntegrationService\n"
                           "9. Se redirige al listado de perfiles"],
        ["Flujo Alternativo", "2a. Nombre duplicado → Error 400: 'Ya existe un perfil con ese nombre'\n"
                              "3a. Pesos no suman 1.0 → Error 400: validación de pesos\n"
                              "4a. Apto <= considerado → Error 400: validación de umbrales"],
    ]

    add_formatted_table(doc, cu_headers, cu_rows, 5,
                        "Caso de uso expandido CU-23: Crear Perfil Institucional",
                        "Descripción del flujo para crear un perfil de evaluación con validación de pesos, "
                        "umbrales y requisitos. La invalidación de caché ML garantiza consistencia.")

    # --- 5. Diseño de colección ---
    add_subsection(doc, 5, 5, "Diseño de la Colección y Documentos")

    col_headers = ["Campo", "Tipo", "Restricciones", "Descripción"]
    col_rows = [
        ["id", "UUID", "PRIMARY KEY, auto-generado", "Identificador único del perfil"],
        ["institution_name", "TEXT", "UNIQUE, NOT NULL", "Nombre de la institución/empresa"],
        ["sector", "TEXT", "NOT NULL", "Sector: Tecnología, Finanzas, Salud, etc."],
        ["description", "TEXT", "NULLABLE", "Descripción del perfil institucional"],
        ["weights", "JSONB", "NOT NULL", "Pesos de evaluación (5 dimensiones, deben sumar 1.0)"],
        ["requirements", "JSONB", "NOT NULL", "Requisitos: skills, educación, experiencia, idiomas"],
        ["thresholds", "JSONB", "DEFAULT '{\"apto\":0.70,\"considerado\":0.50}'", "Umbrales de clasificación"],
        ["is_active", "BOOLEAN", "DEFAULT true", "Estado activo/inactivo (soft delete)"],
        ["created_by", "UUID", "FK → usuarios(id), NULLABLE", "Admin que creó el perfil"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de creación"],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de última actualización"],
    ]

    add_formatted_table(doc, col_headers, col_rows, 5,
                        "Estructura de la tabla 'institutional_profiles'",
                        "Tabla que almacena la configuración de evaluación por institución. Los campos JSONB "
                        "permiten configuraciones flexibles de pesos, requisitos y umbrales sin migraciones de esquema.")

    # Estructura weights
    weights_headers = ["Campo JSONB", "Tipo", "Rango", "Descripción"]
    weights_rows = [
        ["weights.hard_skills", "float", "0.0 - 1.0", "Peso de habilidades técnicas en la evaluación"],
        ["weights.soft_skills", "float", "0.0 - 1.0", "Peso de habilidades blandas"],
        ["weights.experience", "float", "0.0 - 1.0", "Peso de experiencia laboral"],
        ["weights.education", "float", "0.0 - 1.0", "Peso de formación académica"],
        ["weights.languages", "float", "0.0 - 1.0", "Peso de idiomas"],
        ["Σ weights", "float", "= 1.0", "La suma de los 5 pesos debe ser exactamente 1.0"],
    ]

    add_formatted_table(doc, weights_headers, weights_rows, 5,
                        "Estructura del campo JSONB 'weights' (pesos de evaluación)",
                        "Los pesos determinan la importancia relativa de cada dimensión en el cálculo del "
                        "match score. Se valida que la suma sea exactamente 1.0 con tolerancia de 0.01.")

    # Estructura requirements
    req_headers = ["Campo JSONB", "Tipo", "Ejemplo", "Descripción"]
    req_rows = [
        ["requirements.required_skills", "Array[str]", '["Python", "SQL"]', "Habilidades técnicas obligatorias"],
        ["requirements.preferred_skills", "Array[str]", '["Docker", "AWS"]', "Habilidades técnicas deseables"],
        ["requirements.required_soft_skills", "Array[str]", '["Liderazgo"]', "Habilidades blandas requeridas"],
        ["requirements.min_experience_years", "float", "1.5", "Años mínimos de experiencia"],
        ["requirements.required_education_level", "str", '"Licenciatura"', "Nivel educativo mínimo"],
        ["requirements.required_languages", "Array[str]", '["Inglés"]', "Idiomas requeridos"],
    ]

    add_formatted_table(doc, req_headers, req_rows, 5,
                        "Estructura del campo JSONB 'requirements' (requisitos)",
                        "Los requisitos definen los criterios mínimos que el FeatureExtractor utiliza para "
                        "calcular los scores individuales de cada dimensión del CV.")

    # --- 6. Mockups ---
    add_subsection(doc, 5, 6, "Diseño e Interfaz del Módulo (Mockups)")

    add_figure_placeholder(doc,
        "Tomar captura de ProfilesAdminView.vue.\n"
        "Debe mostrar: barra de búsqueda por nombre/sector, filtro por sector\n"
        "(dropdown), toggle 'Mostrar inactivos', tarjetas de perfiles con:\n"
        "nombre institución, sector, descripción, umbrales (apto/considerado %),\n"
        "desglose de pesos, botones Editar/Desactivar.",
        "Vista administrativa de perfiles institucionales",
        "Panel de gestión de perfiles institucionales con búsqueda, filtrado por sector, "
        "y tarjetas que muestran la configuración de pesos y umbrales de cada perfil.")

    add_figure_placeholder(doc,
        "Tomar captura de ProfileFormView.vue sección INFORMACIÓN BÁSICA.\n"
        "Debe mostrar: campo nombre institución, selector de sector\n"
        "(dropdown con 11 opciones + Otro), campo descripción.",
        "Formulario de perfil institucional – Sección información básica",
        "Primera sección del formulario que recoge el nombre único de la institución, "
        "el sector de actividad y una descripción opcional.")

    add_figure_placeholder(doc,
        "Tomar captura de ProfileFormView.vue sección PESOS DE EVALUACIÓN.\n"
        "Debe mostrar: 5 sliders o inputs numéricos (hard_skills, soft_skills,\n"
        "experience, education, languages) con porcentajes que suman 100%,\n"
        "botón 'Normalizar' para ajustar automáticamente.",
        "Formulario de perfil institucional – Sección pesos de evaluación",
        "Configuración de los pesos ponderados para las 5 dimensiones de evaluación. "
        "Incluye botón de normalización automática para ajustar al 100%.")

    add_figure_placeholder(doc,
        "Tomar captura de ProfileFormView.vue sección REQUISITOS.\n"
        "Debe mostrar: campos de skills requeridos/preferidos (con autocompletado),\n"
        "nivel educativo, años de experiencia, idiomas requeridos,\n"
        "botones agregar/eliminar para listas dinámicas.",
        "Formulario de perfil institucional – Sección requisitos",
        "Configuración de requisitos específicos: habilidades técnicas y blandas requeridas/preferidas, "
        "nivel educativo mínimo, años de experiencia y idiomas. Con autocompletado de skills comunes.")

    # --- 7. Código fuente ---
    add_subsection(doc, 5, 7, "Código Fuente del Módulo (Incremento del Producto)")

    code_create_profile = '''@router.post("/institutional-profiles",
             response_model=InstitutionalProfileResponse,
             status_code=201)
async def create_institutional_profile(
    profile: InstitutionalProfileCreate,
    admin_user: dict = Depends(verify_operator_access),
    ml_service: MLIntegrationService = Depends(get_ml_service_dependency)
):
    # Verificar nombre único
    existing = supabase.table("institutional_profiles")\\
        .select("id").eq("institution_name", profile.institution_name)\\
        .execute()
    if existing.data:
        raise HTTPException(400, f"Ya existe: {profile.institution_name}")

    # Preparar datos con validaciones Pydantic
    data = {
        'institution_name': profile.institution_name,
        'sector': profile.sector,
        'description': profile.description,
        'weights': profile.weights.model_dump(),       # Suma = 1.0
        'requirements': profile.requirements.model_dump(),
        'thresholds': profile.thresholds.model_dump(), # apto > considerado
        'is_active': True,
        'created_by': admin_user['user_id']
    }
    response = supabase.table("institutional_profiles")\\
        .insert(data).execute()

    # Invalidar caché ML para usar datos actualizados
    ml_service.invalidate_cache()
    return InstitutionalProfileResponse(**response.data[0])'''

    add_code_block(doc, code_create_profile,
                   "Endpoint de creación de perfil institucional (institutional_profiles.py)",
                   "Endpoint POST que crea un perfil con validación de nombre único, pesos que suman 1.0, "
                   "umbrales consistentes e invalidación de caché del servicio ML.")

    code_schemas = '''class InstitutionalWeights(BaseModel):
    hard_skills: float = Field(ge=0, le=1, default=0.30)
    soft_skills: float = Field(ge=0, le=1, default=0.20)
    experience: float = Field(ge=0, le=1, default=0.25)
    education: float = Field(ge=0, le=1, default=0.15)
    languages: float = Field(ge=0, le=1, default=0.10)

    @model_validator(mode='after')
    def validate_weights_sum(self):
        total = (self.hard_skills + self.soft_skills +
                 self.experience + self.education + self.languages)
        if abs(total - 1.0) > 0.01:
            raise ValueError(f'Los pesos deben sumar 1.0 (actual: {total})')
        return self

class InstitutionalThresholds(BaseModel):
    apto: float = Field(default=0.70, ge=0, le=1)
    considerado: float = Field(default=0.50, ge=0, le=1)

    @model_validator(mode='after')
    def validate_order(self):
        if self.apto <= self.considerado:
            raise ValueError('Umbral apto debe ser > considerado')
        return self

class InstitutionalRequirements(BaseModel):
    min_experience_years: float = 0.0
    required_skills: List[str] = []
    preferred_skills: List[str] = []
    required_education_level: str = "Licenciatura"
    required_languages: List[str] = []'''

    add_code_block(doc, code_schemas,
                   "Schemas Pydantic para perfiles institucionales (ml_schemas.py)",
                   "Modelos de validación que garantizan la consistencia de los datos: pesos que suman 1.0, "
                   "umbrales donde apto > considerado, y requisitos con valores por defecto sensatos.")

    code_cache = '''@router.put("/institutional-profiles/{profile_id}")
async def update_institutional_profile(
    profile_id: str,
    profile: InstitutionalProfileUpdate,
    admin_user: dict = Depends(verify_operator_access),
    ml_service: MLIntegrationService = Depends(get_ml_service_dependency)
):
    # Verificar existencia
    existing = supabase.table("institutional_profiles")\\
        .select("*").eq("id", profile_id).execute()
    if not existing.data:
        raise HTTPException(404, f"Perfil no encontrado: {profile_id}")

    # Verificar nombre único si cambia
    if profile.institution_name:
        name_check = supabase.table("institutional_profiles")\\
            .select("id").eq("institution_name", profile.institution_name)\\
            .neq("id", profile_id).execute()
        if name_check.data:
            raise HTTPException(400, "Nombre ya existe en otro perfil")

    # Actualizar campos proporcionados
    update_data = {'updated_at': datetime.utcnow().isoformat()}
    if profile.weights: update_data['weights'] = profile.weights.model_dump()
    if profile.requirements: update_data['requirements'] = ...
    if profile.thresholds: update_data['thresholds'] = ...

    response = supabase.table("institutional_profiles")\\
        .update(update_data).eq("id", profile_id).execute()

    # CRÍTICO: Invalidar caché ML
    ml_service.invalidate_cache()
    return InstitutionalProfileResponse(**response.data[0])'''

    add_code_block(doc, code_cache,
                   "Endpoint de edición con invalidación de caché ML (institutional_profiles.py)",
                   "Endpoint PUT que actualiza un perfil institucional con validación de nombre único "
                   "e invalida la caché del MLIntegrationService para que las evaluaciones posteriores "
                   "utilicen la configuración actualizada.")

    # --- 8. Pruebas unitarias ---
    add_subsection(doc, 5, 8, "Resultado de Pruebas Unitarias")

    pruebas_headers = ["ID Prueba", "Descripción", "Entrada", "Resultado Esperado", "Estado"]
    pruebas_rows = [
        ["PU-5.01", "Crear perfil con datos válidos", "nombre, sector, pesos=1.0, umbrales válidos", "Perfil creado + status 201", "Aprobado"],
        ["PU-5.02", "Crear perfil con nombre duplicado", "institution_name existente", "Error 400: 'Ya existe un perfil con ese nombre'", "Aprobado"],
        ["PU-5.03", "Crear perfil con pesos que no suman 1.0", "weights total = 0.85", "Error 400: 'Los pesos deben sumar 1.0'", "Aprobado"],
        ["PU-5.04", "Crear perfil con apto <= considerado", "apto=0.50, considerado=0.70", "Error 400: 'apto debe ser > considerado'", "Aprobado"],
        ["PU-5.05", "Listar perfiles activos", "include_inactive=false", "Solo perfiles con is_active=true", "Aprobado"],
        ["PU-5.06", "Listar perfiles con inactivos", "include_inactive=true", "Todos los perfiles", "Aprobado"],
        ["PU-5.07", "Filtrar por sector", "sector='Tecnología'", "Solo perfiles del sector seleccionado", "Aprobado"],
        ["PU-5.08", "Editar perfil: cambiar pesos", "PUT con nuevos pesos válidos", "Pesos actualizados + caché invalidada", "Aprobado"],
        ["PU-5.09", "Editar perfil: nombre duplicado", "Cambiar a nombre existente", "Error 400", "Aprobado"],
        ["PU-5.10", "Desactivar perfil (soft delete)", "DELETE /{id}", "is_active=false + caché invalidada", "Aprobado"],
        ["PU-5.11", "Reactivar perfil ya activo", "POST /activate en perfil activo", "Error 400: 'El perfil ya está activo'", "Aprobado"],
        ["PU-5.12", "Listar sectores", "GET /sectors", "Lista de sectores únicos ordenados", "Aprobado"],
        ["PU-5.13", "Acceso sin rol admin", "Token rol='estudiante'", "Error 403: Acceso denegado", "Aprobado"],
    ]

    add_formatted_table(doc, pruebas_headers, pruebas_rows, 5,
                        "Resultados de pruebas unitarias – Sprint 5",
                        "Se ejecutaron 13 pruebas unitarias cubriendo CRUD, validaciones de pesos/umbrales/nombre, "
                        "filtrado, soft delete e invalidación de caché. Todas las pruebas pasaron exitosamente.")

    add_figure_placeholder(doc,
        "Tomar captura de la terminal ejecutando las pruebas del Sprint 5.\n"
        "Puede ser pytest o pruebas con Postman mostrando respuestas exitosas\n"
        "de los endpoints de perfiles institucionales.",
        "Ejecución de pruebas unitarias del Sprint 5",
        "Captura de la terminal mostrando la ejecución exitosa de las pruebas unitarias "
        "del módulo de Perfiles Institucionales.")

    # --- 9. Bitácora de cambios ---
    add_subsection(doc, 5, 9, "Bitácora de Cambios al Sprint Backlog")

    cambios_headers = ["Fecha", "Cambio", "Motivo", "Impacto"]
    cambios_rows = [
        ["03/12/2025", "Se añadió validación de nombre único de institución", "Evitar duplicados que confundan el sistema de matching", "Bajo: +3h en validación backend"],
        ["09/12/2025", "Se añadió invalidación de caché ML en todos los CRUD", "Sin invalidar, las evaluaciones usaban datos obsoletos del perfil por hasta 5 minutos", "Alto: +6h (impacto transversal en create, update, delete)"],
        ["12/12/2025", "Se añadieron 11 sectores predefinidos + opción 'Otro'", "Estandarizar la clasificación de sectores para mejores filtros", "Bajo: +4h (frontend dropdown + endpoint sectors)"],
        ["18/12/2025", "Se aumentó complejidad del formulario (4 secciones)", "El formulario original de 2 secciones era insuficiente; se añadieron secciones de umbrales y requisitos detallados", "Alto: +12h en ProfileFormView (666 líneas finales)"],
    ]

    add_formatted_table(doc, cambios_headers, cambios_rows, 5,
                        "Bitácora de cambios al Sprint Backlog – Sprint 5",
                        "Los cambios principales se debieron a la necesidad de invalidación de caché ML "
                        "y a la complejidad del formulario de configuración de perfiles institucionales.")

    doc.add_page_break()


# ============================================================
# SPRINT 6: INFORMES Y REPORTES
# ============================================================

def generar_sprint6(doc):
    sprint_num = 6
    add_heading_sprint(doc, sprint_num)

    doc.add_paragraph(
        f"El Sprint 6 se desarrolló del {SPRINTS[6]['inicio']} al {SPRINTS[6]['fin']} "
        "y se centró en la implementación del módulo de Informes y Reportes, que proporciona "
        "al administrador un dashboard analítico con visualizaciones de datos del sistema. "
        "El módulo incluye resumen de usuarios, crecimiento temporal, distribución por rol, "
        "nube de habilidades más frecuentes, estadísticas de completitud de perfiles y "
        "exportación del reporte a PDF. Adicionalmente se integraron las recomendaciones "
        "personalizadas que cierran el ciclo del sistema de intermediación."
    )

    # --- 1. Sprint Backlog ---
    add_subsection(doc, 6, 1, "Sprint Backlog")

    backlog_headers = ["ID", "Historia de Usuario", "Tarea", "Prioridad", "Estado", "Responsable", "Est. (h)"]
    backlog_rows = [
        ["HU-29", "Como admin, quiero ver un resumen de usuarios registrados", "Implementar endpoint GET /api/analytics/users-summary", "Alta", "Completado", AUTOR, "8"],
        ["HU-29", "", "Implementar tarjetas de resumen en ReportsView", "Alta", "Completado", AUTOR, "6"],
        ["HU-30", "Como admin, quiero ver el crecimiento de usuarios en el tiempo", "Implementar endpoint GET /api/analytics/user-growth", "Alta", "Completado", AUTOR, "10"],
        ["HU-30", "", "Implementar UserGrowthChart (Chart.js line chart)", "Alta", "Completado", AUTOR, "8"],
        ["HU-30", "", "Implementar agrupación dinámica (día/mes según rango)", "Media", "Completado", AUTOR, "4"],
        ["HU-31", "Como admin, quiero ver la distribución de usuarios por rol", "Implementar conteo por rol en users-summary", "Media", "Completado", AUTOR, "4"],
        ["HU-31", "", "Implementar UserDistributionChart (Chart.js doughnut)", "Media", "Completado", AUTOR, "6"],
        ["HU-32", "Como admin, quiero ver las habilidades más frecuentes", "Implementar endpoint GET /api/analytics/skills-cloud", "Alta", "Completado", AUTOR, "8"],
        ["HU-32", "", "Implementar SkillsBarChart (horizontal bars) x2", "Alta", "Completado", AUTOR, "8"],
        ["HU-32", "", "Top 15 hard skills + Top 15 soft skills", "Media", "Completado", AUTOR, "4"],
        ["HU-33", "Como admin, quiero ver estadísticas de completitud de perfiles", "Implementar endpoint GET /api/analytics/profile-completion", "Media", "Completado", AUTOR, "6"],
        ["HU-33", "", "Implementar tarjetas de métricas de completitud", "Media", "Completado", AUTOR, "4"],
        ["HU-34", "Como admin, quiero filtrar reportes por rango de fechas", "Implementar selector de rango (30 días, este año, personalizado)", "Alta", "Completado", AUTOR, "8"],
        ["HU-34", "", "Implementar apply_date_filter() reutilizable", "Media", "Completado", AUTOR, "4"],
        ["HU-35", "Como admin, quiero exportar el reporte a PDF", "Integrar html2pdf.js para exportación", "Baja", "Completado", AUTOR, "10"],
        ["HU-35", "", "Implementar layout con break-before-page para PDF", "Baja", "Completado", AUTOR, "4"],
        ["HU-36", "Como sistema, necesito integrar recomendaciones personalizadas", "Implementar RecommendationService (Singleton)", "Alta", "Completado", AUTOR, "20"],
        ["HU-36", "", "Implementar endpoint GET /api/recommendations", "Alta", "Completado", AUTOR, "12"],
        ["HU-36", "", "Implementar filtrado por rol (estudiante→pasantía, titulado→empleo)", "Alta", "Completado", AUTOR, "6"],
        ["HU-36", "", "Implementar historial y marcado de vistas", "Media", "Completado", AUTOR, "8"],
    ]

    add_formatted_table(doc, backlog_headers, backlog_rows, 6,
                        "Sprint Backlog – Sprint 6: Informes y Reportes",
                        "Se detallan las historias de usuario para el dashboard analítico, incluyendo resumen de usuarios, "
                        "crecimiento, habilidades frecuentes, completitud de perfiles, exportación PDF y recomendaciones.")

    # --- 2. Bitácora de temas importantes ---
    add_subsection(doc, 6, 2, "Bitácora de Identificación de Temas Importantes")

    bitacora_headers = ["Fecha", "Tema Identificado", "Impacto", "Decisión Tomada"]
    bitacora_rows = [
        ["13/01/2026", "Librería de gráficos para Vue 3", "Alto", "Se eligió Chart.js con vue-chartjs por su compatibilidad con Vue 3, variedad de gráficos (line, doughnut, bar) y facilidad de personalización."],
        ["15/01/2026", "Agrupación temporal dinámica", "Medio", "Si el rango de fechas es < 60 días, se agrupa por día; caso contrario por mes. Esto evita gráficos congestionados con muchos puntos."],
        ["18/01/2026", "Filtrado de fechas reutilizable", "Medio", "Se implementó la función apply_date_filter() que se reutiliza en los 4 endpoints de analytics, recibiendo start_date y end_date como query params."],
        ["20/01/2026", "Métricas absolutas vs en período", "Medio", "Total de usuarios y distribución por rol son siempre absolutos. Nuevos usuarios y crecimiento respetan el filtro de fechas. Default: últimos 30 días."],
        ["23/01/2026", "Exportación a PDF", "Medio", "Se integró html2pdf.js que captura el contenedor HTML del dashboard y lo exporta como PDF. Se usó break-before-page CSS para separar secciones."],
        ["28/01/2026", "Normalización de skills para nube", "Bajo", "Los skills se normalizan a minúsculas (strip + lower) antes de contarlos con Counter para evitar duplicados como 'Python' vs 'python'."],
        ["02/02/2026", "Recomendaciones basadas en perfil guardado", "Alto", "Las recomendaciones usan el perfil profesional persistido (Sprint 2) en lugar de requerir subir el CV cada vez. Requisito: perfil >=70% completo."],
    ]

    add_formatted_table(doc, bitacora_headers, bitacora_rows, 6,
                        "Bitácora de temas importantes – Sprint 6",
                        "Registro de decisiones técnicas para el dashboard analítico, incluyendo librería de gráficos, "
                        "agrupación temporal, exportación PDF y sistema de recomendaciones personalizadas.")

    # --- 3. Actualización tabla de tareas ---
    add_subsection(doc, 6, 3, "Actualización de la Tabla de Tareas")

    tareas_headers = ["Tarea", "Estado", "Horas Est.", "Horas Real", "Observaciones"]
    tareas_rows = [
        ["Endpoint GET /analytics/users-summary", "Completado", "8", "8", "Total + por rol + nuevos en período"],
        ["Endpoint GET /analytics/user-growth", "Completado", "10", "12", "Agrupación dinámica día/mes"],
        ["Endpoint GET /analytics/skills-cloud", "Completado", "8", "9", "Counter + normalización + top 15"],
        ["Endpoint GET /analytics/profile-completion", "Completado", "6", "6", "Candidatos, completados, con CV"],
        ["apply_date_filter() reutilizable", "Completado", "4", "3", "Helper para todos los endpoints"],
        ["ReportsView.vue (321 líneas)", "Completado", "18", "20", "Layout responsive + selectores + cards"],
        ["UserGrowthChart.vue", "Completado", "8", "7", "Line chart con fill + smooth curve"],
        ["UserDistributionChart.vue", "Completado", "6", "5", "Doughnut chart con tooltips %"],
        ["SkillsBarChart.vue x2", "Completado", "8", "8", "Horizontal bars para hard y soft skills"],
        ["Exportación PDF (html2pdf.js)", "Completado", "10", "12", "Layout + break-before-page"],
        ["RecommendationService (Singleton)", "Completado", "20", "24", "Score all ofertas + save + history"],
        ["Endpoint GET /api/recommendations", "Completado", "12", "14", "Filtros + rol auto + top_n"],
        ["Endpoints historial + vistas + elegibilidad", "Completado", "14", "13", "4 endpoints adicionales"],
        ["Selector rango de fechas (frontend)", "Completado", "8", "7", "30 días, este año, personalizado"],
    ]

    add_formatted_table(doc, tareas_headers, tareas_rows, 6,
                        "Actualización de tareas – Sprint 6",
                        "Total estimado: 140h, Total real: 148h. Desviación: +5.7%. "
                        "El RecommendationService y la exportación PDF requirieron tiempo adicional.")

    # --- 4. Diagrama de caso de uso expandido ---
    add_subsection(doc, 6, 4, "Diagrama de Caso de Uso Expandido del Módulo")

    add_figure_placeholder(doc,
        "Tomar captura del diagrama de caso de uso del módulo de Informes y Reportes.\n"
        "Actores: Administrador/Operador, Usuario Autenticado, Sistema.\n"
        "Casos de uso (Admin): Ver Resumen de Usuarios, Ver Crecimiento Temporal,\n"
        "Ver Distribución por Rol, Ver Habilidades Frecuentes, Ver Completitud de Perfiles,\n"
        "Filtrar por Rango de Fechas, Exportar Reporte a PDF.\n"
        "Casos de uso (Usuario): Obtener Recomendaciones Personalizadas,\n"
        "Ver Historial de Recomendaciones, Marcar Recomendación como Vista.\n"
        "Sistema: Calcular Estadísticas, Generar Recomendaciones ML.",
        "Diagrama de caso de uso expandido – Módulo de Informes y Reportes",
        "El diagrama muestra las funcionalidades analíticas del administrador y el sistema de "
        "recomendaciones personalizadas para usuarios, cerrando el ciclo del sistema de intermediación.")

    cu_headers = ["Campo", "Descripción"]
    cu_rows = [
        ["Caso de Uso", "CU-29: Consultar Dashboard de Informes"],
        ["Actor Principal", "Administrador / Operador"],
        ["Precondiciones", "El usuario tiene rol administrador u operador."],
        ["Postcondiciones", "Se visualizan las estadísticas del sistema con gráficos interactivos."],
        ["Flujo Principal", "1. El admin accede a la vista de Reportes\n"
                           "2. Selecciona un rango de fechas (default: últimos 30 días)\n"
                           "3. El sistema consulta 4 endpoints de analytics en paralelo\n"
                           "4. Se muestran tarjetas de resumen: total usuarios, activos, completados\n"
                           "5. Se renderiza gráfico de crecimiento (línea temporal)\n"
                           "6. Se renderiza gráfico de distribución por rol (doughnut)\n"
                           "7. Se muestran barras horizontales de top 15 hard y soft skills\n"
                           "8. Opcionalmente, el admin exporta el reporte a PDF"],
        ["Flujo Alternativo", "2a. Rango personalizado inválido → Se usa default 30 días\n"
                              "8a. Exportación PDF falla → Se muestra mensaje de error"],
    ]

    add_formatted_table(doc, cu_headers, cu_rows, 6,
                        "Caso de uso expandido CU-29: Consultar Dashboard de Informes",
                        "Descripción del flujo del dashboard analítico con consulta paralela de endpoints, "
                        "visualización con Chart.js y exportación a PDF con html2pdf.js.")

    # --- 5. Diseño de colección ---
    add_subsection(doc, 6, 5, "Diseño de la Colección y Documentos")

    doc.add_paragraph(
        "El Sprint 6 introduce la tabla 'recomendaciones' para almacenar las recomendaciones "
        "generadas para cada usuario. Los endpoints de analytics consultan tablas existentes "
        "(usuarios, perfiles_profesionales) con agregaciones."
    )

    col_headers = ["Campo", "Tipo", "Restricciones", "Descripción"]
    col_rows = [
        ["id", "UUID", "PRIMARY KEY, auto-generado", "Identificador único de la recomendación"],
        ["usuario_id", "UUID", "FK → usuarios(id), NOT NULL", "Usuario que recibe la recomendación"],
        ["oferta_id", "UUID", "FK → ofertas_laborales(id), NOT NULL", "Oferta laboral recomendada"],
        ["match_score", "NUMERIC", "NOT NULL", "Score de compatibilidad (0.0 a 1.0)"],
        ["clasificacion", "TEXT", "NOT NULL", "Clasificación: APTO, CONSIDERADO, NO_APTO"],
        ["scores_detalle", "JSONB", "NOT NULL", "Scores individuales por dimensión"],
        ["fortalezas", "JSONB", "DEFAULT '[]'", "Lista de fortalezas principales"],
        ["debilidades", "JSONB", "DEFAULT '[]'", "Lista de debilidades principales"],
        ["fue_vista", "BOOLEAN", "DEFAULT false", "Si el usuario vio la recomendación"],
        ["vista_at", "TIMESTAMPTZ", "NULLABLE", "Fecha en que se marcó como vista"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de generación"],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de última actualización"],
    ]

    add_formatted_table(doc, col_headers, col_rows, 6,
                        "Estructura de la tabla 'recomendaciones'",
                        "Tabla que almacena las recomendaciones personalizadas generadas por el sistema ML. "
                        "Incluye scores detallados, fortalezas/debilidades y seguimiento de visualización.")

    # Endpoints analytics
    analytics_headers = ["Endpoint", "Método", "Parámetros", "Respuesta"]
    analytics_rows = [
        ["/api/analytics/users-summary", "GET", "start_date?, end_date?", "total_users, roles{}, new_users_in_period"],
        ["/api/analytics/user-growth", "GET", "start_date?, end_date?", "labels[], data[], grouped_by (day|month)"],
        ["/api/analytics/skills-cloud", "GET", "start_date?, end_date?", "hard_skills[{name,count}], soft_skills[{name,count}]"],
        ["/api/analytics/profile-completion", "GET", "start_date?, end_date?", "total_candidates, profiles_completed, completion_rate"],
    ]

    add_formatted_table(doc, analytics_headers, analytics_rows, 6,
                        "Endpoints de Analytics para el Dashboard de Informes",
                        "Los 4 endpoints de analytics soportan filtrado por rango de fechas. "
                        "Por defecto retornan datos de los últimos 30 días si no se especifica rango.")

    # --- 6. Mockups ---
    add_subsection(doc, 6, 6, "Diseño e Interfaz del Módulo (Mockups)")

    add_figure_placeholder(doc,
        "Tomar captura de ReportsView.vue - PARTE SUPERIOR.\n"
        "Debe mostrar: selector de rango de fechas (Últimos 30 días / Este año /\n"
        "Personalizado), tarjetas de resumen (Total Usuarios, Perfiles Activos,\n"
        "Perfiles Completados, Estudiantes vs Titulados), botón Exportar PDF.",
        "Dashboard de informes – Resumen y selector de fechas",
        "Parte superior del dashboard analítico con selector de rango temporal, "
        "tarjetas de métricas resumidas y botón de exportación a PDF.")

    add_figure_placeholder(doc,
        "Tomar captura de ReportsView.vue - GRÁFICOS.\n"
        "Debe mostrar: gráfico de línea (crecimiento de usuarios),\n"
        "gráfico doughnut (distribución por rol), gráficos de barras\n"
        "horizontales (top hard skills y top soft skills).",
        "Dashboard de informes – Gráficos de visualización",
        "Sección de gráficos del dashboard: línea temporal de crecimiento, doughnut de distribución "
        "por rol (estudiante/titulado/admin) y barras horizontales de habilidades más frecuentes.")

    add_figure_placeholder(doc,
        "Tomar captura del PDF EXPORTADO del reporte.\n"
        "Debe mostrar: el reporte generado como PDF con las mismas\n"
        "secciones del dashboard, con separaciones de página.",
        "Reporte exportado en formato PDF",
        "Vista del archivo PDF generado por html2pdf.js con las secciones del dashboard "
        "separadas por saltos de página para una presentación profesional.")

    # --- 7. Código fuente ---
    add_subsection(doc, 6, 7, "Código Fuente del Módulo (Incremento del Producto)")

    code_summary = '''@router.get("/users-summary")
async def get_users_summary(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    current_user: dict = Depends(verify_operator_access)
):
    # Total usuarios (siempre absoluto)
    total = supabase.table("usuarios")\\
        .select("*", count="exact", head=True).execute()

    # Distribución por rol (absoluta)
    estudiantes = supabase.table("usuarios")\\
        .select("*", count="exact", head=True)\\
        .eq("rol", "estudiante").execute()
    titulados = supabase.table("usuarios")\\
        .select("*", count="exact", head=True)\\
        .eq("rol", "titulado").execute()
    admins = supabase.table("usuarios")\\
        .select("*", count="exact", head=True)\\
        .eq("rol", "administrador").execute()

    # Nuevos usuarios en período (con filtro de fechas)
    recent = supabase.table("usuarios")\\
        .select("*", count="exact", head=True)
    recent = apply_date_filter(recent, "created_at", start_date, end_date)

    return {
        "total_users": total.count,
        "roles": {"estudiante": estudiantes.count,
                  "titulado": titulados.count,
                  "administrador": admins.count},
        "new_users_in_period": recent.execute().count
    }'''

    add_code_block(doc, code_summary,
                   "Endpoint de resumen de usuarios (analytics.py)",
                   "Endpoint GET /api/analytics/users-summary que retorna métricas absolutas (total, distribución por rol) "
                   "y métricas filtradas por período (nuevos usuarios). Default: últimos 30 días.")

    code_skills = '''@router.get("/skills-cloud")
async def get_skills_cloud(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    current_user: dict = Depends(verify_operator_access)
):
    query = supabase.table("perfiles_profesionales")\\
        .select("hard_skills, soft_skills")
    query = apply_date_filter(query, "updated_at", start_date, end_date)
    response = query.execute()

    hard_counter = Counter()
    soft_counter = Counter()

    for profile in response.data:
        for s in (profile.get("hard_skills") or []):
            if s: hard_counter[s.strip().lower()] += 1
        for s in (profile.get("soft_skills") or []):
            if s: soft_counter[s.strip().lower()] += 1

    return {
        "hard_skills": [{"name": k, "count": v}
                        for k, v in hard_counter.most_common(15)],
        "soft_skills": [{"name": k, "count": v}
                        for k, v in soft_counter.most_common(15)]
    }'''

    add_code_block(doc, code_skills,
                   "Endpoint de habilidades frecuentes (analytics.py)",
                   "Endpoint GET /api/analytics/skills-cloud que agrega las habilidades de perfiles actualizados "
                   "en el período, normalizando a minúsculas y retornando top 15 con frecuencia.")

    code_recommendation = '''class RecommendationService:
    _instance = None  # Singleton

    def get_recommendations_for_user(self, user_id, user_role,
                                      top_n=10, tipo=None, sector=None,
                                      recalcular=False):
        # 1. Obtener perfil del usuario
        profile = profile_service.get_profile(user_id)
        completeness = profile_service.calculate_completeness(profile)
        if completeness['score'] < 0.70:
            raise ValueError("Perfil incompleto (< 70%)")

        # 2. Obtener ofertas activas filtradas por rol
        oferta_service = get_oferta_service()
        ofertas = oferta_service.get_ofertas_for_user_role(user_role, top_n=50)
        # estudiante → pasantias, titulado → empleos

        # 3. Evaluar cada oferta contra el perfil del usuario
        gemini_output = profile_service\\
            .get_profile_for_recommendations(user_id)
        resultados = []
        for oferta in ofertas:
            result = self._evaluate_oferta(gemini_output, oferta)
            resultados.append({
                'oferta': oferta,
                'match_score': result['match_score'],
                'clasificacion': result['classification'],
                'fortalezas': self._extract_fortalezas(result),
                'debilidades': self._extract_debilidades(result), ...
            })

        # 4. Ordenar por score y retornar top N
        resultados.sort(key=lambda x: x['match_score'], reverse=True)
        top_results = resultados[:top_n]

        # 5. Guardar recomendaciones en BD
        self._save_recommendations(user_id, top_results)
        return top_results'''

    add_code_block(doc, code_recommendation,
                   "RecommendationService: generación de recomendaciones personalizadas",
                   "Servicio Singleton que genera recomendaciones evaluando el perfil del usuario contra "
                   "todas las ofertas activas filtradas por rol. Utiliza el FeatureExtractor y el modelo Ridge "
                   "para calcular match_score y clasificar cada oferta.")

    # --- 8. Pruebas unitarias ---
    add_subsection(doc, 6, 8, "Resultado de Pruebas Unitarias")

    pruebas_headers = ["ID Prueba", "Descripción", "Entrada", "Resultado Esperado", "Estado"]
    pruebas_rows = [
        ["PU-6.01", "Users summary sin filtro de fechas", "GET /analytics/users-summary", "total_users, roles, new_users (últimos 30 días)", "Aprobado"],
        ["PU-6.02", "Users summary con rango personalizado", "start_date=2025-10-01, end_date=2025-12-31", "new_users_in_period solo del rango", "Aprobado"],
        ["PU-6.03", "User growth agrupado por día (<60 días)", "Rango de 30 días", "grouped_by='day', labels diarios", "Aprobado"],
        ["PU-6.04", "User growth agrupado por mes (>=60 días)", "Rango de 6 meses", "grouped_by='month', labels mensuales", "Aprobado"],
        ["PU-6.05", "Skills cloud: top 15 hard skills", "GET /analytics/skills-cloud", "Array de {name, count} ordenado por frecuencia", "Aprobado"],
        ["PU-6.06", "Skills cloud: normalización", "Skills 'Python' y 'python' en BD", "Contados como uno solo: 'python'", "Aprobado"],
        ["PU-6.07", "Profile completion stats", "GET /analytics/profile-completion", "total_candidates, profiles_completed, completion_rate", "Aprobado"],
        ["PU-6.08", "Acceso a analytics sin rol admin", "Token con rol='estudiante'", "Error 403: Acceso denegado", "Aprobado"],
        ["PU-6.09", "Recomendaciones con perfil completo", "GET /api/recommendations (perfil >=70%)", "Lista de recomendaciones ordenadas por score", "Aprobado"],
        ["PU-6.10", "Recomendaciones con perfil incompleto", "Perfil con score < 70%", "Error: 'Perfil incompleto'", "Aprobado"],
        ["PU-6.11", "Filtrado por rol: estudiante → pasantías", "user_role='estudiante'", "Solo ofertas tipo='pasantia'", "Aprobado"],
        ["PU-6.12", "Marcar recomendación como vista", "POST /recommendations/{id}/viewed", "fue_vista=true, vista_at actualizado", "Aprobado"],
        ["PU-6.13", "Historial de recomendaciones", "GET /api/recommendations/history", "Lista paginada con scores y fechas", "Aprobado"],
    ]

    add_formatted_table(doc, pruebas_headers, pruebas_rows, 6,
                        "Resultados de pruebas unitarias – Sprint 6",
                        "Se ejecutaron 13 pruebas unitarias cubriendo los 4 endpoints de analytics, control de acceso, "
                        "recomendaciones personalizadas y filtrado por rol. Todas las pruebas pasaron exitosamente.")

    add_figure_placeholder(doc,
        "Tomar captura de la terminal ejecutando las pruebas del Sprint 6.\n"
        "Puede ser pytest o pruebas con Postman mostrando los endpoints\n"
        "de analytics y recomendaciones respondiendo correctamente.",
        "Ejecución de pruebas unitarias del Sprint 6",
        "Captura de la terminal mostrando la ejecución exitosa de las pruebas unitarias "
        "del módulo de Informes, Reportes y Recomendaciones.")

    # --- 9. Bitácora de cambios ---
    add_subsection(doc, 6, 9, "Bitácora de Cambios al Sprint Backlog")

    cambios_headers = ["Fecha", "Cambio", "Motivo", "Impacto"]
    cambios_rows = [
        ["15/01/2026", "Se añadió agrupación dinámica día/mes en user-growth", "Gráfico de línea con datos mensuales de muchos meses se congestionaba", "Bajo: +4h (lógica condicional en endpoint)"],
        ["20/01/2026", "Se incorporó sistema de recomendaciones al sprint", "El sistema de recomendaciones personalizadas era necesario para cerrar el ciclo del sistema", "Alto: +46h (RecommendationService + endpoints + historial)"],
        ["25/01/2026", "Se añadió exportación a PDF con html2pdf.js", "Requisito identificado durante revisión: reportes exportables para presentaciones", "Medio: +14h (integración + layout CSS)"],
        ["02/02/2026", "Se incrementaron pruebas unitarias de 8 a 13", "Se añadieron pruebas para recomendaciones que no estaban inicialmente planificadas", "Bajo: +6h"],
    ]

    add_formatted_table(doc, cambios_headers, cambios_rows, 6,
                        "Bitácora de cambios al Sprint Backlog – Sprint 6",
                        "El cambio más significativo fue la incorporación del sistema de recomendaciones "
                        "personalizadas, que añadió complejidad sustancial al sprint pero cerró el ciclo "
                        "completo del sistema de intermediación laboral.")


# ============================================================
# GENERADOR PRINCIPAL
# ============================================================

def generar_documento():
    doc = Document()

    # Configurar estilos base
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.54)

    # Título
    title = doc.add_heading("3.5. Desarrollo de Sprints (Continuación)", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    doc.add_paragraph(
        "Continuación de la documentación de sprints del proyecto. "
        "Este documento cubre los Sprints 4, 5 y 6 que completan el desarrollo del sistema."
    )
    doc.add_paragraph(
        f"Desarrollador: {AUTOR}"
    )
    doc.add_paragraph(
        f"Período: {SPRINTS[4]['inicio']} – {SPRINTS[6]['fin']}"
    )

    doc.add_page_break()

    # Generar sprints
    generar_sprint4(doc)
    generar_sprint5(doc)
    generar_sprint6(doc)

    # Guardar
    output_path = os.path.join(os.path.dirname(__file__), "Documentacion_Sprints_4_6.docx")
    doc.save(output_path)
    print(f"Documento generado exitosamente: {output_path}")
    print(f"Total tablas en este documento: {tabla_counter['value'] - 25}")
    print(f"Total figuras en este documento: {figura_counter['value'] - 26}")
    print(f"Tablas acumuladas (1-6): {tabla_counter['value']}")
    print(f"Figuras acumuladas (1-6): {figura_counter['value']}")
    return output_path


if __name__ == "__main__":
    generar_documento()
