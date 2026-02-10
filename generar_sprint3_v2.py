"""
Generador de Documentacion Sprint 3 v2 - Evaluacion de Perfiles
Enfocado en la funcionalidad del sistema, casos de uso y vistas del usuario.
Autor: Ivan Condori Choquehuanca
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONFIGURACION
# ============================================================

AUTOR = "Ivan Condori Choquehuanca"
SISTEMA = "Sistema de Intermediacion Laboral con PLN"
SPRINT = {"num": 3, "nombre": "Evaluacion de Perfiles", "inicio": "20/10/2025", "fin": "09/11/2025"}

# Contadores globales (empiezan desde donde termino el Sprint 2)
# Sprint 1: Tablas 1-8, Figuras 1-8
# Sprint 2: Tablas 9-16, Figuras 9-18
tabla_counter = {"value": 16}
figura_counter = {"value": 18}


# ============================================================
# FUNCIONES AUXILIARES
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_table(doc, headers, rows, caption, nota, fuente="Elaboracion propia"):
    tabla_counter["value"] += 1
    num = tabla_counter["value"]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(table)

    # Header
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")

    # Data
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    # Caption
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Tabla {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    # Nota
    nota_p = doc.add_paragraph()
    nota_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    # Fuente
    fuente_p = doc.add_paragraph()
    fuente_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_figure_placeholder(doc, descripcion_captura, caption, nota, fuente="Elaboracion propia"):
    figura_counter["value"] += 1
    num = figura_counter["value"]

    p_box = doc.add_paragraph()
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_box.add_run(f"[INSERTAR CAPTURA DE PANTALLA]\n{descripcion_captura}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 0, 0)
    run.italic = True

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_code_block(doc, code_text, caption, nota, fuente="Elaboracion propia"):
    figura_counter["value"] += 1
    num = figura_counter["value"]

    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Consolas" w:hAnsi="Consolas" w:cs="Consolas"/>')
        rPr.append(rFonts)

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_paragraph_text(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


# ============================================================
# GENERACION DEL DOCUMENTO
# ============================================================

def generate():
    doc = Document()

    # --------------------------------------------------------
    # TITULO PRINCIPAL
    # --------------------------------------------------------
    h1 = doc.add_heading("3.5.3. Sprint 3: Evaluacion de Perfiles", level=1)

    intro = (
        "El Sprint 3 se desarrollo del 20/10/2025 al 09/11/2025 y se centro en la implementacion "
        "del modulo de Evaluacion de Perfiles, el cual constituye el nucleo funcional del sistema de "
        "intermediacion laboral. Este modulo permite a los usuarios (Estudiantes y Titulados) evaluar "
        "su perfil profesional contra los requisitos de instituciones especificas, consultar un historial "
        "de evaluaciones realizadas, y recibir recomendaciones automaticas de ofertas laborales "
        "(pasantias o empleos) ordenadas por grado de compatibilidad. El sistema utiliza el pipeline de "
        "Machine Learning documentado en la Seccion 3.4 para calcular un Match Score sobre cinco "
        "dimensiones (Habilidades Tecnicas, Habilidades Blandas, Educacion, Experiencia e Idiomas), "
        "clasificar la candidatura (Apto, Considerado, No Apto) y generar insights de fortalezas y "
        "debilidades."
    )
    add_paragraph_text(doc, intro)

    # --------------------------------------------------------
    # 3.5.3.1 SPRINT BACKLOG
    # --------------------------------------------------------
    doc.add_heading("3.5.3.1. Sprint Backlog", level=2)

    sprint_backlog_rows = [
        ["HU-09", "Evaluar mi CV contra una institucion", "Implementar vista de evaluacion individual con seleccion de perfil institucional y carga de CV", "Alta", "Completada", AUTOR, "20"],
        ["HU-09", "Evaluar mi CV contra una institucion", "Integrar endpoint POST /api/ml/evaluate-cv con pipeline ML completo", "Alta", "Completada", AUTOR, "16"],
        ["HU-09", "Evaluar mi CV contra una institucion", "Disenar componente de resultado con score circular, clasificacion, fortalezas y debilidades", "Alta", "Completada", AUTOR, "12"],
        ["HU-10", "Ver mi historial de evaluaciones", "Implementar vista de historial con lista paginada y modal de detalles", "Media", "Completada", AUTOR, "10"],
        ["HU-10", "Ver mi historial de evaluaciones", "Implementar endpoint GET /api/ml/evaluations/history con paginacion", "Media", "Completada", AUTOR, "8"],
        ["HU-11", "Recibir recomendaciones de ofertas", "Implementar vista de recomendaciones con sidebar de estadisticas y lista de tarjetas expandibles", "Alta", "Completada", AUTOR, "20"],
        ["HU-11", "Recibir recomendaciones de ofertas", "Implementar servicio de recomendaciones con filtrado por tipo de usuario (estudiante->pasantias, titulado->empleos)", "Alta", "Completada", AUTOR, "16"],
        ["HU-11", "Recibir recomendaciones de ofertas", "Implementar verificacion de elegibilidad (perfil >= 70% completo)", "Alta", "Completada", AUTOR, "6"],
        ["HU-11", "Recibir recomendaciones de ofertas", "Disenar componente RecommendationCard con scores detallados, fortalezas/debilidades y datos de oferta", "Alta", "Completada", AUTOR, "14"],
        ["HU-12", "Ver desglose detallado de mi evaluacion", "Implementar componente ScoreChart con barras de progreso por dimension", "Media", "Completada", AUTOR, "8"],
        ["HU-12", "Ver desglose detallado de mi evaluacion", "Implementar componente EligibilityWarning para perfiles incompletos", "Media", "Completada", AUTOR, "4"],
        ["T-09", "Tarea tecnica", "Implementar store de evaluacion (evaluation.store.js) con Pinia", "Media", "Completada", AUTOR, "6"],
        ["T-10", "Tarea tecnica", "Implementar composable useRecommendations con logica de elegibilidad y estado", "Media", "Completada", AUTOR, "8"],
        ["T-11", "Tarea tecnica", "Implementar API clients (evaluation.api.js, recommendations.api.js)", "Media", "Completada", AUTOR, "6"],
        ["T-12", "Tarea tecnica", "Implementar cache de recomendaciones en base de datos con opcion de recalculo", "Baja", "Completada", AUTOR, "10"],
    ]

    add_formatted_table(doc,
        headers=["ID", "Historia de Usuario", "Tarea", "Prioridad", "Estado", "Responsable", "Est. (h)"],
        rows=sprint_backlog_rows,
        caption=f"Sprint Backlog - Sprint 3: Evaluacion de Perfiles",
        nota="Se detallan las historias de usuario y tareas tecnicas para las tres funcionalidades principales del modulo: evaluacion individual de CV, historial de evaluaciones y recomendaciones automaticas de ofertas.",
    )

    # --------------------------------------------------------
    # 3.5.3.2 BITACORA DE TEMAS IMPORTANTES
    # --------------------------------------------------------
    doc.add_heading("3.5.3.2. Bitacora de Identificacion de Temas Importantes", level=2)

    bitacora_rows = [
        ["22/10/2025", "Separacion de evaluacion individual vs. recomendaciones", "Alto",
         "Se decidio separar en dos flujos: (1) evaluacion individual donde el usuario elige una institucion especifica, y (2) recomendaciones automaticas donde el sistema evalua contra todas las ofertas activas."],
        ["25/10/2025", "Filtrado de ofertas por tipo de usuario", "Alto",
         "Los estudiantes solo ven pasantias y los titulados solo ven empleos, evitando recomendaciones irrelevantes."],
        ["28/10/2025", "Umbral de elegibilidad para recomendaciones", "Alto",
         "Se establecio que el perfil debe estar al menos 70% completo para acceder a recomendaciones. Se implemento EligibilityWarning con indicacion de campos faltantes."],
        ["30/10/2025", "Clasificacion por umbrales porcentuales", "Alto",
         "Se definieron tres categorias: Apto (>=70%), Considerado (>=50% y <70%), No Apto (<50%). Cada una con color y mensaje diferenciado."],
        ["02/11/2025", "Cache de recomendaciones", "Medio",
         "Se implemento cache en base de datos para evitar recalculos innecesarios. El usuario puede forzar recalculo con el boton 'Actualizar'."],
        ["05/11/2025", "Sidebar de estadisticas en recomendaciones", "Medio",
         "Se agrego sidebar colapsable con gauge del mejor match, contadores (total, nuevas, aptas, consideradas) y resumen del perfil del usuario."],
        ["07/11/2025", "Vista de evaluacion sin autenticacion", "Medio",
         "La evaluacion individual permite uso sin autenticacion como funcionalidad de demostracion. El historial si requiere autenticacion."],
    ]

    add_formatted_table(doc,
        headers=["Fecha", "Tema Identificado", "Impacto", "Decision Tomada"],
        rows=bitacora_rows,
        caption=f"Bitacora de temas importantes - Sprint 3",
        nota="Registro de decisiones de diseno funcional del modulo, enfocadas en la experiencia del usuario y la logica de negocio del sistema de matching.",
    )

    # --------------------------------------------------------
    # 3.5.3.3 ACTUALIZACION DE TAREAS
    # --------------------------------------------------------
    doc.add_heading("3.5.3.3. Actualizacion de la Tabla de Tareas", level=2)

    tareas_rows = [
        ["Vista de Evaluacion Individual (EvaluationView)", "Completada", "20", "22", "Requirio ajustes en el selector de perfil institucional"],
        ["Componente de Resultado (EvaluationResult)", "Completada", "12", "14", "Se agrego circulo de score animado y seccion colapsable de datos extraidos"],
        ["Endpoint de evaluacion POST /api/ml/evaluate-cv", "Completada", "16", "18", "Integracion completa del pipeline ML con extraccion Gemini"],
        ["Vista de Historial (HistoryView)", "Completada", "10", "10", "Paginacion y modal de detalles"],
        ["Endpoint de historial GET /api/ml/evaluations/history", "Completada", "8", "7", "Sin desviacion significativa"],
        ["Vista de Recomendaciones (MisRecomendacionesView)", "Completada", "20", "24", "Mayor complejidad por sidebar colapsable y estados multiples"],
        ["Servicio de recomendaciones (recommendation_service.py)", "Completada", "16", "18", "Logica de filtrado por tipo de usuario y cache"],
        ["Componente RecommendationCard expandible", "Completada", "14", "16", "Seccion expandible con scores detallados y datos de oferta"],
        ["Verificacion de elegibilidad", "Completada", "6", "5", "Endpoint y componente EligibilityWarning"],
        ["Componente ScoreChart (barras de progreso)", "Completada", "8", "7", "Barras coloreadas por umbral de score"],
        ["Sidebar de estadisticas (RecommendationsSidebar)", "Completada", "8", "10", "Gauge chart del mejor match y contadores"],
        ["Store de evaluacion y composable useRecommendations", "Completada", "14", "13", "Gestion de estado con Pinia y composable"],
        ["API clients (evaluation.api.js, recommendations.api.js)", "Completada", "6", "5", "Integracion con axios client compartido"],
        ["Cache de recomendaciones en base de datos", "Completada", "10", "9", "Persistencia y opcion de recalculo"],
        ["Pruebas unitarias del modulo", "Completada", "6", "7", "14 pruebas cubriendo scorers, endpoint y flujo completo"],
    ]

    add_formatted_table(doc,
        headers=["Tarea", "Estado", "Horas Est.", "Horas Real", "Observaciones"],
        rows=tareas_rows,
        caption=f"Actualizacion de tareas - Sprint 3",
        nota="Total estimado: 174h, Total real: 185h. Desviacion: +6.3%. La mayor desviacion se dio en la vista de recomendaciones por la complejidad del sidebar colapsable y el manejo de multiples estados (carga, elegibilidad, error, vacio, con datos).",
    )

    # --------------------------------------------------------
    # 3.5.3.4 DIAGRAMA DE CASO DE USO EXPANDIDO
    # --------------------------------------------------------
    doc.add_heading("3.5.3.4. Diagrama de Caso de Uso Expandido del Modulo", level=2)

    # Descripcion textual del diagrama
    desc_caso_uso = (
        "Este diagrama de caso de uso representa el modulo de Evaluacion de Perfiles, un sistema "
        "integral donde los actores Estudiante y Titulado interactuan con tres funcionalidades principales. "
        "En la primera, \"Evaluar CV contra Institucion\", el usuario selecciona un perfil institucional de "
        "los disponibles en el sistema y carga su CV en formato PDF, lo que detona la funcion central del "
        "actor Sistema Machine Learning: ejecutar el proceso \"Calcular Match Score\". Este calculo integra "
        "obligatoriamente (<<include>>) la puntuacion de cinco dimensiones clave (Habilidades Tecnicas, "
        "Habilidades Blandas, Educacion, Experiencia e Idiomas) que, tras pasar por una etapa de \"Aplicar "
        "Pesos Institucionales\" para ponderar su relevancia segun los requisitos de cada institucion, "
        "derivan en dos resultados simultaneos: una clasificacion automatica de la candidatura basada en "
        "umbrales porcentuales (No Apto <50%, Considerado >=50%, Apto >=70%) y la generacion de insights, "
        "que devuelve al usuario un desglose transparente de sus puntuaciones por dimension, evidenciando "
        "sus fortalezas y debilidades especificas frente a los requisitos institucionales.\n\n"
        "En la segunda funcionalidad, \"Consultar Historial\", el usuario autenticado accede a una lista "
        "paginada de todas sus evaluaciones previas, pudiendo ver los detalles de cada una en un modal "
        "que muestra la institucion evaluada, la fecha, el score obtenido, la clasificacion y los scores "
        "por dimension.\n\n"
        "La tercera funcionalidad, \"Consultar Recomendaciones\", es la mas compleja. Cuando el usuario "
        "accede a esta seccion, el sistema primero verifica su elegibilidad (<<include>> \"Verificar "
        "Elegibilidad\"), comprobando que su perfil este al menos 70% completo. Si no es elegible, se "
        "muestra una advertencia con los campos faltantes y un enlace para completar el perfil. Si es "
        "elegible, el sistema obtiene todas las ofertas activas filtradas por tipo de usuario "
        "(Estudiante -> solo pasantias, Titulado -> solo empleos) y ejecuta el proceso \"Calcular Match "
        "Score\" contra cada oferta. Los resultados se presentan como una lista ordenada de tarjetas con "
        "score de compatibilidad, clasificacion, fortalezas y debilidades, que el usuario puede expandir "
        "para ver el desglose detallado de puntuaciones por dimension y los datos completos de la oferta."
    )
    add_paragraph_text(doc, desc_caso_uso)

    # Placeholder del diagrama
    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura del diagrama de caso de uso del modulo de Evaluacion de Perfiles.\n"
            "Actores: Estudiante, Titulado, Sistema ML, Administrador.\n"
            "Casos de uso principales:\n"
            "- Evaluar CV contra Institucion (incluye: Seleccionar Perfil Institucional, Cargar CV, Calcular Match Score)\n"
            "- Consultar Historial de Evaluaciones (incluye: Ver Detalle de Evaluacion)\n"
            "- Consultar Recomendaciones (incluye: Verificar Elegibilidad, Calcular Match Score para cada oferta)\n"
            "Subcasos incluidos en Calcular Match Score:\n"
            "- Puntuar Hard Skills, Puntuar Soft Skills, Puntuar Educacion, Puntuar Experiencia, Puntuar Idiomas\n"
            "- Aplicar Pesos Institucionales\n"
            "- Clasificar Candidatura (Apto/Considerado/No Apto)\n"
            "- Generar Insights (Fortalezas y Debilidades)"
        ),
        caption="Diagrama de caso de uso expandido - Modulo de Evaluacion de Perfiles",
        nota="El diagrama muestra tres flujos principales (evaluacion individual, historial, recomendaciones) con el proceso central de matching ML compartido. Los actores Estudiante y Titulado difieren en el tipo de ofertas que reciben.",
    )

    # Tabla de caso de uso expandido CU-12
    cu12_rows = [
        ["Caso de uso", "CU-12: Evaluar CV contra Perfil Institucional"],
        ["Actores", "Estudiante, Titulado, Sistema ML"],
        ["Precondiciones", "Existen perfiles institucionales activos en el sistema. El usuario tiene un CV en formato PDF."],
        ["Descripcion", "El usuario selecciona una institucion de las disponibles, carga su CV en PDF, y el sistema ejecuta el pipeline de ML para calcular un score de compatibilidad con desglose por dimension."],
        ["Flujo principal",
         "1. El sistema muestra los perfiles institucionales activos como tarjetas seleccionables (nombre, sector, descripcion).\n"
         "2. El usuario selecciona un perfil institucional.\n"
         "3. El usuario arrastra o selecciona su CV en formato PDF en la zona de carga.\n"
         "4. El usuario presiona el boton 'Evaluar CV'.\n"
         "5. El sistema muestra un spinner con el texto 'Analizando...'.\n"
         "6. El Sistema ML extrae los datos del CV con Google Gemini (hard skills, soft skills, educacion, experiencia, idiomas).\n"
         "7. El Sistema ML calcula los scores de las 5 dimensiones contra los requisitos institucionales.\n"
         "8. El Sistema ML aplica los pesos institucionales y calcula el Match Score total.\n"
         "9. El Sistema ML clasifica la candidatura: Apto (>=70%), Considerado (>=50%), No Apto (<50%).\n"
         "10. El sistema muestra el resultado con: (a) circulo animado con el porcentaje de match, (b) badge de clasificacion con color diferenciado, (c) mensaje descriptivo, (d) nombre de la institucion evaluada.\n"
         "11. El sistema muestra el desglose de puntuacion por dimension con barras de progreso coloreadas.\n"
         "12. El sistema muestra las top 3 fortalezas y top 3 areas de mejora del candidato.\n"
         "13. El sistema muestra los datos extraidos del CV en una seccion colapsable.\n"
         "14. Si el usuario esta autenticado, la evaluacion se guarda en su historial."
        ],
        ["Flujos alternativos",
         "FA1: El usuario no selecciona institucion -> el boton 'Evaluar CV' permanece deshabilitado.\n"
         "FA2: El archivo no es PDF -> se muestra alerta 'Por favor selecciona un archivo PDF valido'.\n"
         "FA3: Error en el pipeline ML -> se muestra mensaje de error en banner rojo con opcion de reintentar.\n"
         "FA4: No hay perfiles institucionales disponibles -> se muestra mensaje 'No hay perfiles institucionales disponibles'."
        ],
        ["Postcondiciones", "El usuario visualiza su score de compatibilidad, clasificacion, fortalezas y debilidades. Si esta autenticado, la evaluacion queda registrada en su historial."],
    ]

    add_formatted_table(doc,
        headers=["Campo", "Descripcion"],
        rows=cu12_rows,
        caption="Caso de uso expandido CU-12: Evaluar CV contra Perfil Institucional",
        nota="Descripcion del flujo completo de evaluacion individual donde el usuario elige una institucion y carga su CV para obtener un analisis de compatibilidad.",
    )

    # Caso de uso CU-13: Consultar Recomendaciones
    cu13_rows = [
        ["Caso de uso", "CU-13: Consultar Recomendaciones de Ofertas Laborales"],
        ["Actores", "Estudiante, Titulado, Sistema ML"],
        ["Precondiciones", "El usuario esta autenticado. Existen ofertas laborales activas en el sistema. El usuario tiene un perfil profesional."],
        ["Descripcion", "El sistema evalua automaticamente el perfil del usuario contra todas las ofertas activas (filtradas por tipo de usuario) y presenta una lista ordenada de recomendaciones con scores de compatibilidad."],
        ["Flujo principal",
         "1. El usuario accede a la seccion 'Recomendaciones' desde el menu lateral.\n"
         "2. El sistema verifica la elegibilidad del usuario (perfil >= 70% completo).\n"
         "3. El sistema obtiene las ofertas activas filtradas por tipo: si es Estudiante, solo pasantias; si es Titulado, solo empleos.\n"
         "4. El Sistema ML ejecuta el Calcular Match Score del perfil del usuario contra cada oferta disponible.\n"
         "5. El sistema ordena los resultados por Match Score de mayor a menor.\n"
         "6. El sistema muestra el sidebar izquierdo con: (a) gauge circular del mejor match, (b) titulo de la mejor oferta, (c) contadores de total, nuevas, aptas y consideradas, (d) top skills del perfil, (e) anios de experiencia, (f) boton de actualizar.\n"
         "7. El sistema muestra la lista de tarjetas de recomendacion, cada una con: (a) badge de clasificacion (dorado=Apto, azul=Considerado, rojo=No Apto), (b) badge 'Nueva' si no ha sido vista, (c) titulo de la oferta, (d) informacion de institucion, sector y modalidad, (e) circulo de Match Score con porcentaje, (f) descripcion breve, (g) primeras 2 fortalezas y debilidades.\n"
         "8. El usuario hace clic en 'Ver detalles' en una tarjeta.\n"
         "9. La tarjeta se expande mostrando: (a) barras de progreso de los 5 scores detallados (Habilidades Tecnicas, Habilidades Blandas, Educacion, Experiencia, Idiomas), (b) datos completos de la oferta (tipo, ubicacion, cupos, fecha de cierre), (c) nota informativa sobre el proceso de postulacion.\n"
         "10. El sistema marca la recomendacion como 'vista' automaticamente."
        ],
        ["Flujos alternativos",
         "FA1: Perfil incompleto (<70%) -> se muestra componente EligibilityWarning con: motivo, accion requerida, lista de campos faltantes, y boton 'Completar Perfil' que dirige a la carga de CV.\n"
         "FA2: No hay ofertas disponibles -> se muestra estado vacio con icono, mensaje explicativo, posibles razones, y botones 'Revisar mi Perfil' / 'Reintentar'.\n"
         "FA3: Error de conexion -> se muestra banner rojo con mensaje de error y boton 'Reintentar'.\n"
         "FA4: El usuario presiona 'Actualizar' -> se recalculan las recomendaciones forzando un nuevo procesamiento ML."
        ],
        ["Postcondiciones", "El usuario visualiza sus recomendaciones ordenadas por compatibilidad, con desglose detallado de scores y fortalezas/debilidades para cada oferta."],
    ]

    add_formatted_table(doc,
        headers=["Campo", "Descripcion"],
        rows=cu13_rows,
        caption="Caso de uso expandido CU-13: Consultar Recomendaciones de Ofertas Laborales",
        nota="Descripcion del flujo de recomendaciones automaticas donde el sistema evalua el perfil del usuario contra todas las ofertas activas y presenta resultados ordenados.",
    )

    # Caso de uso CU-14: Consultar Historial
    cu14_rows = [
        ["Caso de uso", "CU-14: Consultar Historial de Evaluaciones"],
        ["Actores", "Estudiante, Titulado"],
        ["Precondiciones", "El usuario esta autenticado. El usuario ha realizado al menos una evaluacion previa."],
        ["Descripcion", "El usuario accede a una lista paginada de todas sus evaluaciones anteriores, con la posibilidad de ver los detalles de cada una."],
        ["Flujo principal",
         "1. El usuario accede a la seccion 'Historial' desde el menu lateral.\n"
         "2. El sistema muestra el contador total de evaluaciones y la pagina actual.\n"
         "3. El sistema muestra una lista de tarjetas de evaluacion, cada una con: (a) nombre de la institucion, (b) fecha y hora de la evaluacion, (c) badge de clasificacion (Apto/Considerado/No Apto) con color, (d) porcentaje de match en texto grande, (e) barra de progreso coloreada.\n"
         "4. El usuario hace clic en una tarjeta de evaluacion.\n"
         "5. El sistema muestra un modal con los detalles: (a) nombre de la institucion, (b) fecha de evaluacion, (c) badge de clasificacion grande, (d) score en formato de porcentaje prominente, (e) barras de progreso para cada una de las 5 dimensiones.\n"
         "6. El usuario cierra el modal haciendo clic en la X o fuera del modal.\n"
         "7. Si hay mas de 10 evaluaciones, el usuario puede navegar entre paginas."
        ],
        ["Flujos alternativos",
         "FA1: No autenticado -> se muestra mensaje con icono de candado, texto 'Inicia sesion para ver tu historial' y boton 'Iniciar Sesion'.\n"
         "FA2: Sin evaluaciones previas -> se muestra estado vacio con mensaje 'No tienes evaluaciones aun' y boton 'Evaluar mi CV'."
        ],
        ["Postcondiciones", "El usuario puede revisar el detalle de sus evaluaciones pasadas para monitorear su progreso."],
    ]

    add_formatted_table(doc,
        headers=["Campo", "Descripcion"],
        rows=cu14_rows,
        caption="Caso de uso expandido CU-14: Consultar Historial de Evaluaciones",
        nota="Descripcion del flujo de consulta de historial donde el usuario revisa evaluaciones previas con detalles por dimension.",
    )

    # --------------------------------------------------------
    # 3.5.3.5 DISENO DE COLECCION Y DOCUMENTOS
    # --------------------------------------------------------
    doc.add_heading("3.5.3.5. Diseno de la Coleccion y Documentos", level=2)

    add_paragraph_text(doc,
        "El Sprint 3 introduce tres estructuras de datos principales: la tabla de evaluaciones "
        "individuales, la tabla de recomendaciones generadas, y la tabla de estadisticas de recomendaciones. "
        "Estas tablas se integran con las tablas existentes de perfiles profesionales (Sprint 2) y "
        "perfiles institucionales (administracion)."
    )

    # Tabla evaluaciones
    eval_rows = [
        ["id", "UUID", "PK, auto-generado", "Identificador unico de la evaluacion"],
        ["usuario_id", "UUID", "FK -> auth.users, nullable", "Usuario que realizo la evaluacion (null si anonimo)"],
        ["institutional_profile_id", "UUID", "FK -> perfiles_institucionales", "Perfil institucional contra el que se evaluo"],
        ["institution_name", "VARCHAR(255)", "NOT NULL", "Nombre de la institucion (denormalizado para consulta rapida)"],
        ["match_score", "DECIMAL(5,4)", "NOT NULL, 0-1", "Score de compatibilidad total calculado por el pipeline ML"],
        ["classification", "VARCHAR(20)", "NOT NULL", "Clasificacion: APTO, CONSIDERADO, NO_APTO"],
        ["cv_scores", "JSONB", "NOT NULL", "Scores detallados por dimension: {hard_skills, soft_skills, education, experience, languages}"],
        ["top_strengths", "TEXT[]", "nullable", "Array con las top 3 fortalezas identificadas"],
        ["top_weaknesses", "TEXT[]", "nullable", "Array con las top 3 areas de mejora identificadas"],
        ["gemini_extraction", "JSONB", "nullable", "Datos completos extraidos del CV por Google Gemini"],
        ["evaluated_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha y hora de la evaluacion"],
    ]

    add_formatted_table(doc,
        headers=["Campo", "Tipo", "Restricciones", "Descripcion"],
        rows=eval_rows,
        caption="Estructura de la tabla 'evaluaciones'",
        nota="Almacena cada evaluacion individual de CV contra un perfil institucional. Los scores detallados se guardan en JSONB para flexibilidad, y las fortalezas/debilidades como arrays de texto.",
    )

    # Tabla recomendaciones
    rec_rows = [
        ["id", "UUID", "PK, auto-generado", "Identificador unico de la recomendacion"],
        ["usuario_id", "UUID", "FK -> auth.users, NOT NULL", "Usuario al que pertenece la recomendacion"],
        ["oferta_id", "UUID", "FK -> ofertas_laborales, NOT NULL", "Oferta laboral recomendada"],
        ["match_score", "DECIMAL(5,4)", "NOT NULL, 0-1", "Score de compatibilidad calculado"],
        ["clasificacion", "VARCHAR(20)", "NOT NULL", "APTO, CONSIDERADO o NO_APTO"],
        ["scores_detalle", "JSONB", "NOT NULL", "Scores de las 5 dimensiones individuales"],
        ["fortalezas", "TEXT[]", "nullable", "Top 3 fortalezas para esta oferta"],
        ["debilidades", "TEXT[]", "nullable", "Top 3 debilidades para esta oferta"],
        ["fue_vista", "BOOLEAN", "DEFAULT false", "Indica si el usuario ya expandio/vio la recomendacion"],
        ["vista_at", "TIMESTAMPTZ", "nullable", "Fecha y hora en que fue vista"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT NOW()", "Fecha de generacion de la recomendacion"],
    ]

    add_formatted_table(doc,
        headers=["Campo", "Tipo", "Restricciones", "Descripcion"],
        rows=rec_rows,
        caption="Estructura de la tabla 'recomendaciones'",
        nota="Almacena las recomendaciones generadas para cada usuario. El campo 'fue_vista' permite diferenciar recomendaciones nuevas de las ya consultadas. Las recomendaciones se regeneran cuando el usuario presiona 'Actualizar'.",
    )

    # --------------------------------------------------------
    # 3.5.3.6 DISENO E INTERFAZ DEL MODULO
    # --------------------------------------------------------
    doc.add_heading("3.5.3.6. Diseno e Interfaz del Modulo (Vistas del Sistema)", level=2)

    add_paragraph_text(doc,
        "A continuacion se presentan las interfaces de usuario del modulo de Evaluacion de Perfiles, "
        "mostrando como se visualizan las tres funcionalidades principales en el sistema: evaluacion "
        "individual, resultado de evaluacion, historial y recomendaciones."
    )

    # --- Vista 1: Evaluacion Individual ---
    add_paragraph_text(doc, "a) Vista de Evaluacion Individual de CV", bold=True)
    add_paragraph_text(doc,
        "La vista de evaluacion individual (EvaluationView) presenta un flujo en dos pasos. "
        "En el primer paso, el usuario selecciona una institucion de las disponibles, mostradas como "
        "tarjetas con radio buttons que resaltan en azul al seleccionarse, indicando el nombre de la "
        "institucion, su sector y una descripcion breve. Debajo, una zona de drag-and-drop permite "
        "cargar el CV en PDF, mostrando un icono de nube con el texto 'Haz clic o arrastra tu CV'. "
        "Al seleccionar un archivo valido, la zona cambia a verde con un icono de verificacion y el "
        "nombre del archivo. El boton 'Evaluar CV' en la esquina inferior derecha se activa solo cuando "
        "ambos campos estan completos. Durante el procesamiento, el boton muestra un spinner con el "
        "texto 'Analizando...'."
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura de la vista de Evaluacion Individual (EvaluationView.vue).\n"
            "Ruta: /evaluation\n"
            "Debe mostrar: titulo 'Evaluacion de CV', seccion 1 con tarjetas de instituciones\n"
            "seleccionables, seccion 2 con zona drag-and-drop para CV, boton 'Evaluar CV'."
        ),
        caption="Interfaz de evaluacion individual de CV - Paso 1: Seleccion",
        nota="Vista que permite al usuario seleccionar una institucion y cargar su CV en PDF. Las tarjetas institucionales muestran nombre, sector y descripcion. La zona de carga acepta solo archivos PDF.",
    )

    # --- Vista 2: Resultado de Evaluacion ---
    add_paragraph_text(doc, "b) Vista de Resultado de Evaluacion", bold=True)
    add_paragraph_text(doc,
        "El resultado de evaluacion (EvaluationResult) se muestra en el segundo paso del flujo. "
        "La tarjeta principal presenta un borde izquierdo coloreado segun la clasificacion "
        "(verde=Apto, amarillo=Considerado, rojo=No Apto). A la izquierda aparece un circulo SVG "
        "animado que se llena progresivamente mostrando el porcentaje de match. A la derecha se "
        "muestra un badge con la clasificacion, el texto 'Resultado de Evaluacion', un mensaje "
        "descriptivo (ej: 'Tu perfil cumple con los requisitos de esta institucion') y el nombre "
        "de la institucion evaluada. Debajo, el componente ScoreChart muestra barras de progreso "
        "horizontales para cada una de las 5 dimensiones. Luego, una grilla de dos columnas presenta "
        "las fortalezas (icono de check verde con items precedidos por '+') y las areas de mejora "
        "(icono de alerta rojo con items precedidos por '-'). Finalmente, una seccion colapsable "
        "'Datos Extraidos del CV' muestra las habilidades tecnicas (badges azules), habilidades "
        "blandas (badges morados), educacion y experiencia extraidos por Gemini."
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura del Resultado de Evaluacion (EvaluationResult.vue).\n"
            "Ruta: /evaluation (despues de evaluar)\n"
            "Debe mostrar: tarjeta principal con circulo de score animado, badge de clasificacion,\n"
            "barras de progreso por dimension, fortalezas y debilidades, seccion colapsable de datos extraidos."
        ),
        caption="Resultado detallado de evaluacion de CV",
        nota="Vista de resultados que presenta el match score en un circulo animado, la clasificacion con color diferenciado, barras de progreso por dimension, fortalezas, debilidades y datos extraidos del CV.",
    )

    # --- Vista 3: Historial ---
    add_paragraph_text(doc, "c) Vista de Historial de Evaluaciones", bold=True)
    add_paragraph_text(doc,
        "La vista de historial (HistoryView) muestra una lista paginada de evaluaciones previas. "
        "Un encabezado indica el total de evaluaciones y la pagina actual. Cada evaluacion se presenta "
        "como una tarjeta con hover shadow que muestra a la izquierda el nombre de la institucion y la "
        "fecha, y a la derecha un badge de clasificacion coloreado y el porcentaje de match en texto "
        "grande (verde >=70%, amarillo >=50%, rojo <50%), seguido de una barra de progreso. Al hacer "
        "clic en una tarjeta, se abre un modal centrado con: la institucion, fecha, un badge de "
        "clasificacion grande, el porcentaje en formato prominente (texto de 5xl), y barras de progreso "
        "para las 5 dimensiones. La paginacion en la parte inferior permite navegar entre paginas con "
        "botones 'Anterior' y 'Siguiente'."
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura de la vista de Historial (HistoryView.vue).\n"
            "Ruta: /history (autenticado, con evaluaciones previas)\n"
            "Debe mostrar: encabezado con total y pagina, lista de tarjetas de evaluaciones\n"
            "con institucion, fecha, badge de clasificacion, porcentaje y barra de progreso."
        ),
        caption="Historial de evaluaciones del usuario",
        nota="Vista que lista las evaluaciones previas del usuario con paginacion. Cada tarjeta muestra institucion, fecha, clasificacion y porcentaje. Al hacer clic se abre un modal con detalles y scores por dimension.",
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura del modal de detalles de evaluacion en HistoryView.\n"
            "Debe mostrar: nombre de institucion, fecha, badge de clasificacion grande,\n"
            "porcentaje prominente, y barras de progreso por dimension."
        ),
        caption="Modal de detalle de evaluacion en el historial",
        nota="Modal que se abre al seleccionar una evaluacion del historial, mostrando la clasificacion prominente, el porcentaje de match y el desglose de scores por las 5 dimensiones.",
    )

    # --- Vista 4: Recomendaciones ---
    add_paragraph_text(doc, "d) Vista de Recomendaciones de Ofertas", bold=True)
    add_paragraph_text(doc,
        "La vista de recomendaciones (MisRecomendacionesView) es la mas completa del modulo. "
        "Utiliza el componente AppLayout con un sidebar colapsable a la izquierda (RecommendationsSidebar). "
        "El sidebar expandido muestra: un gauge circular con el porcentaje del mejor match, el titulo "
        "de la mejor oferta, un cuadro de estadisticas con fondo azul navy que contiene 4 contadores "
        "(Total, Nuevas, Aptas, Consideradas) con badges de colores, un resumen del perfil del usuario "
        "(top 4 skills y anios de experiencia), un boton 'Actualizar' para recalcular las recomendaciones, "
        "y un enlace 'Ver mi Perfil'. El sidebar se puede colapsar a 60px de ancho mostrando solo "
        "indicadores resumidos."
    )

    add_paragraph_text(doc,
        "El area principal muestra un encabezado con el titulo 'Mis Recomendaciones' y un subtitulo "
        "que varia segun el rol ('Pasantias recomendadas' para estudiantes, 'Empleos recomendados' "
        "para titulados). Debajo se presenta la lista de tarjetas de recomendacion (RecommendationCard). "
        "Cada tarjeta en su estado colapsado muestra: una fila de badges con la clasificacion "
        "(dorado=Apto, azul navy=Considerado, rojo=No Apto) y un badge 'Nueva' si no ha sido vista; "
        "el titulo de la oferta; metadatos con iconos (institucion, sector, modalidad); un circulo "
        "prominente con el porcentaje de match coloreado segun umbral; los primeros 100 caracteres de "
        "descripcion; y las primeras 2 fortalezas (badges dorados) y debilidades (badges neutros)."
    )

    add_paragraph_text(doc,
        "Al expandir una tarjeta (clic en 'Ver detalles'), se despliega una seccion adicional con "
        "borde superior que muestra en una grilla de dos columnas: a la izquierda, el 'Detalle de "
        "Scores' con barras de progreso (ProgressBar) para cada dimension (Habilidades Tecnicas, "
        "Habilidades Blandas, Educacion, Experiencia, Idiomas); a la derecha, los 'Detalles de la "
        "Oferta' con tipo, ubicacion, cupos disponibles y fecha de cierre. En la parte inferior "
        "aparece una nota informativa en fondo azul claro: 'Para postular a esta oferta, contacta "
        "a la Unidad de Vinculacion de la EMI. Este sistema solo genera recomendaciones de "
        "correspondencia.'"
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura de la vista de Recomendaciones (MisRecomendacionesView.vue).\n"
            "Ruta: /mis-recomendaciones (autenticado, perfil completo)\n"
            "Debe mostrar: sidebar izquierdo con gauge de mejor match y estadisticas,\n"
            "lista de tarjetas de recomendacion con badges, scores y descripcion."
        ),
        caption="Vista de recomendaciones de ofertas laborales",
        nota="Vista principal del modulo de recomendaciones con sidebar de estadisticas y lista de tarjetas ordenadas por match score. El sidebar muestra el mejor match, contadores y resumen del perfil.",
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura de una RecommendationCard expandida.\n"
            "Debe mostrar: la tarjeta con badge de clasificacion, titulo, score circular,\n"
            "y la seccion expandida con barras de progreso por dimension y datos de la oferta."
        ),
        caption="Tarjeta de recomendacion expandida con detalle de scores",
        nota="Al expandir una tarjeta de recomendacion, el usuario puede ver las barras de progreso de cada dimension (Habilidades Tecnicas, Blandas, Educacion, Experiencia, Idiomas) y los datos completos de la oferta.",
    )

    # --- Vista 5: Elegibilidad ---
    add_paragraph_text(doc, "e) Advertencia de Elegibilidad (Perfil Incompleto)", bold=True)
    add_paragraph_text(doc,
        "Cuando un usuario intenta acceder a las recomendaciones sin tener su perfil completo "
        "al 70%, el sistema muestra el componente EligibilityWarning en lugar de la lista de "
        "recomendaciones. Este componente muestra una tarjeta con fondo dorado que contiene: "
        "un icono de advertencia, el motivo por el cual no es elegible (ej: 'Tu perfil no esta "
        "suficientemente completo'), la accion requerida, la lista de campos faltantes (ej: "
        "'Habilidades tecnicas', 'Idiomas'), y un boton 'Completar Perfil' que redirige a la "
        "vista de carga de CV (/digitalizacion/subir-cv)."
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura del componente EligibilityWarning.\n"
            "Ruta: /mis-recomendaciones (con perfil incompleto <70%)\n"
            "Debe mostrar: tarjeta dorada con icono de advertencia, motivo, campos faltantes\n"
            "y boton 'Completar Perfil'."
        ),
        caption="Advertencia de elegibilidad para perfiles incompletos",
        nota="Componente que se muestra cuando el perfil del usuario no alcanza el umbral del 70% de completitud necesario para recibir recomendaciones. Indica especificamente que campos faltan.",
    )

    # --------------------------------------------------------
    # 3.5.3.7 CODIGO FUENTE DEL MODULO
    # --------------------------------------------------------
    doc.add_heading("3.5.3.7. Codigo Fuente del Modulo (Incremento del Producto)", level=2)

    add_paragraph_text(doc,
        "A continuacion se presentan los fragmentos clave del codigo fuente que implementan "
        "las funcionalidades del modulo de Evaluacion de Perfiles. El detalle del pipeline de "
        "Machine Learning (scorers, feature engineering, modelo Ridge) se documenta en la "
        "Seccion 3.4."
    )

    # Codigo 1: Endpoint de evaluacion
    add_code_block(doc, """@router.post("/evaluate-cv", response_model=CVEvaluationResponse)
async def evaluate_cv(request: CVEvaluationRequest,
    current_user = Depends(get_current_user_optional),
    ml_service = Depends(verify_ml_model_loaded)):

    # 1. Extraer CV con Gemini
    gemini_output = ml_service.extract_cv_with_gemini(request.cv_file)

    # 2. Cargar perfil institucional
    profile = ml_service.load_institutional_profile(
        request.institutional_profile_id)

    # 3. Evaluar CV (Feature Engineering + Scoring)
    evaluation = ml_service.evaluate_cv(gemini_output, profile)

    # 4. Guardar si usuario autenticado
    if current_user:
        ml_service.save_evaluation(
            user_id=current_user['user_id'],
            profile_id=request.institutional_profile_id,
            evaluation_result=evaluation)

    return CVEvaluationResponse(
        match_score=evaluation['match_score'],
        classification=evaluation['classification'],
        cv_scores=evaluation['cv_scores'],
        top_strengths=evaluation['top_strengths'],
        top_weaknesses=evaluation['top_weaknesses'])""",
        caption="Endpoint de evaluacion de CV (ml_predictions.py)",
        nota="Endpoint POST /api/ml/evaluate-cv que orquesta el pipeline completo: extraccion con Gemini, scoring por dimension, clasificacion y persistencia opcional.",
    )

    # Codigo 2: Servicio de recomendaciones
    add_code_block(doc, """class RecommendationService:
    def get_recommendations(self, user_id, user_role, top_n=20,
                            recalcular=False):
        # 1. Verificar elegibilidad
        profile = self.get_user_profile(user_id)
        if profile['completeness_score'] < 0.70:
            raise EligibilityError("Perfil incompleto")

        # 2. Filtrar ofertas por tipo de usuario
        tipo_oferta = 'pasantia' if user_role == 'estudiante' else 'empleo'
        ofertas = self.get_active_offers(tipo=tipo_oferta)

        # 3. Evaluar perfil contra cada oferta
        recomendaciones = []
        for oferta in ofertas:
            inst_profile = self.get_or_create_profile(oferta)
            result = self.ml_service.evaluate_cv(
                profile['gemini_extraction'], inst_profile)
            recomendaciones.append({
                'oferta': oferta,
                'match_score': result['match_score'],
                'clasificacion': result['classification'],
                'scores_detalle': result['cv_scores'],
                'fortalezas': result['top_strengths'][:3],
                'debilidades': result['top_weaknesses'][:3]
            })

        # 4. Ordenar por match_score descendente
        recomendaciones.sort(key=lambda x: x['match_score'], reverse=True)
        return recomendaciones[:top_n]""",
        caption="Servicio de recomendaciones (recommendation_service.py)",
        nota="Servicio que obtiene las recomendaciones para un usuario: verifica elegibilidad, filtra ofertas por tipo (estudiante->pasantias, titulado->empleos), ejecuta el scoring ML contra cada oferta y ordena los resultados.",
    )

    # Codigo 3: Composable useRecommendations
    add_code_block(doc, """// useRecommendations.js - Composable de recomendaciones
export function useRecommendations() {
    const recommendations = ref([])
    const isEligible = ref(true)
    const eligibilityReason = ref('')
    const missingFields = ref([])

    // Computed - Estadisticas
    const bestMatch = computed(() =>
        recommendations.value.reduce((best, r) =>
            r.match_score > (best?.match_score || 0) ? r : best, null))

    const newCount = computed(() =>
        recommendations.value.filter(r => !r.fue_vista).length)

    const aptoCount = computed(() =>
        recommendations.value.filter(
            r => r.clasificacion === 'APTO').length)

    // Cargar recomendaciones
    async function loadRecommendations(forceRefresh = false) {
        await checkEligibility()
        if (!isEligible.value) return

        const { data } = await getMyRecommendations({
            top_n: 20, recalcular: forceRefresh })
        recommendations.value = data.recomendaciones
    }

    // Expandir tarjeta y marcar como vista
    function toggleExpand(recId) {
        if (expandedId.value !== recId) {
            markRecommendationViewed(recId)
        }
        expandedId.value = expandedId.value === recId ? null : recId
    }

    return { recommendations, bestMatch, newCount, aptoCount,
             isEligible, loadRecommendations, toggleExpand }
}""",
        caption="Composable useRecommendations (useRecommendations.js)",
        nota="Composable de Vue 3 que encapsula la logica de recomendaciones: verificacion de elegibilidad, carga de datos, estadisticas computadas (mejor match, conteo de nuevas y aptas), y manejo de expansion de tarjetas.",
    )

    # Codigo 4: RecommendationCard (template simplificado)
    add_code_block(doc, """<!-- RecommendationCard.vue - Tarjeta de recomendacion -->
<template>
  <div class="bg-white rounded-xl shadow-md p-6">
    <!-- Estado colapsado -->
    <div class="flex items-start justify-between">
      <!-- Badges: clasificacion + Nueva -->
      <div class="flex gap-2">
        <Badge :variant="classificationVariant">
          {{ recommendation.clasificacion }}
        </Badge>
        <Badge v-if="!recommendation.fue_vista" variant="gold">Nueva</Badge>
      </div>

      <!-- Circulo de Match Score -->
      <div class="w-16 h-16 rounded-full flex items-center justify-center"
           :class="scoreColorClass">
        <span class="text-xl font-bold">{{ matchPercent }}%</span>
      </div>
    </div>

    <!-- Titulo y metadatos -->
    <h3 class="text-lg font-semibold">{{ recommendation.oferta.titulo }}</h3>
    <p class="text-sm text-slate-500">
      {{ recommendation.oferta.institution_name }} | {{ sector }}
    </p>

    <!-- Fortalezas y debilidades (primeras 2) -->
    <div class="flex flex-wrap gap-1">
      <Badge v-for="f in fortalezas.slice(0,2)" variant="gold">{{ f }}</Badge>
      <Badge v-for="d in debilidades.slice(0,2)" variant="neutral">{{ d }}</Badge>
    </div>

    <!-- Seccion expandida: scores detallados + datos oferta -->
    <div v-if="isExpanded" class="pt-4 border-t mt-4">
      <div class="grid md:grid-cols-2 gap-6">
        <div>  <!-- Barras de progreso por dimension -->
          <ProgressBar v-for="(score, key) in scores_detalle"
            :label="formatLabel(key)" :value="score * 100" />
        </div>
        <div>  <!-- Datos de la oferta -->
          <p>Tipo: {{ oferta.tipo }}</p>
          <p>Ubicacion: {{ oferta.ubicacion }}</p>
          <p>Cupos: {{ oferta.cupos_disponibles }}</p>
        </div>
      </div>
    </div>
  </div>
</template>""",
        caption="Componente RecommendationCard (RecommendationCard.vue)",
        nota="Tarjeta de recomendacion con estado colapsado (badges, titulo, score circular, fortalezas) y estado expandido (barras de progreso por dimension y datos de la oferta). El template se muestra simplificado.",
    )

    # --------------------------------------------------------
    # 3.5.3.8 RESULTADO DE PRUEBAS UNITARIAS
    # --------------------------------------------------------
    doc.add_heading("3.5.3.8. Resultado de Pruebas Unitarias", level=2)

    pruebas_rows = [
        ["PU-S3-01", "Evaluacion individual de CV contra perfil institucional", "CV de prueba + perfil activo", "Score 0-1, clasificacion, scores por dimension, fortalezas/debilidades", "Aprobada"],
        ["PU-S3-02", "Evaluacion sin perfiles institucionales disponibles", "Request sin perfiles activos", "Mensaje 'No hay perfiles institucionales disponibles'", "Aprobada"],
        ["PU-S3-03", "Evaluacion con archivo no PDF", "Archivo .docx", "Rechazo con mensaje de error", "Aprobada"],
        ["PU-S3-04", "Clasificacion APTO (score >= 0.70)", "CV con alto match", "Clasificacion = APTO, badge verde", "Aprobada"],
        ["PU-S3-05", "Clasificacion CONSIDERADO (score 0.50-0.69)", "CV con match medio", "Clasificacion = CONSIDERADO, badge amarillo", "Aprobada"],
        ["PU-S3-06", "Clasificacion NO_APTO (score < 0.50)", "CV con bajo match", "Clasificacion = NO_APTO, badge rojo", "Aprobada"],
        ["PU-S3-07", "Historial de evaluaciones (paginado)", "Usuario con 15 evaluaciones", "Primera pagina con 10 items, segunda con 5", "Aprobada"],
        ["PU-S3-08", "Verificacion de elegibilidad (perfil completo)", "Perfil al 85% completo", "eligible = true", "Aprobada"],
        ["PU-S3-09", "Verificacion de elegibilidad (perfil incompleto)", "Perfil al 40% completo", "eligible = false, lista de campos faltantes", "Aprobada"],
        ["PU-S3-10", "Recomendaciones para estudiante (solo pasantias)", "Usuario rol 'estudiante'", "Solo ofertas tipo 'pasantia' en resultados", "Aprobada"],
        ["PU-S3-11", "Recomendaciones para titulado (solo empleos)", "Usuario rol 'titulado'", "Solo ofertas tipo 'empleo' en resultados", "Aprobada"],
        ["PU-S3-12", "Orden de recomendaciones por match_score", "5 ofertas con scores variados", "Lista ordenada de mayor a menor score", "Aprobada"],
        ["PU-S3-13", "Marcar recomendacion como vista", "Recomendacion nueva (fue_vista=false)", "fue_vista = true, vista_at con timestamp", "Aprobada"],
        ["PU-S3-14", "Recalculo forzado de recomendaciones", "Request con recalcular=true", "Recomendaciones regeneradas (no cache)", "Aprobada"],
    ]

    add_formatted_table(doc,
        headers=["ID Prueba", "Descripcion", "Entrada", "Resultado Esperado", "Estado"],
        rows=pruebas_rows,
        caption="Resultados de pruebas unitarias - Sprint 3",
        nota="Se ejecutaron 14 pruebas unitarias cubriendo los tres flujos principales: evaluacion individual (6 pruebas), historial (1 prueba), y recomendaciones (7 pruebas). Todas las pruebas pasaron exitosamente.",
    )

    add_figure_placeholder(doc,
        descripcion_captura=(
            "Tomar captura de la terminal ejecutando las pruebas del Sprint 3.\n"
            "Comando: python backend/tests/test_feature_engineering.py\n"
            "Debe verse: resultados de cada prueba y 'TODAS LAS PRUEBAS PASARON EXITOSAMENTE'"
        ),
        caption="Ejecucion de pruebas unitarias del Sprint 3",
        nota="Captura de la terminal mostrando la ejecucion exitosa de las pruebas unitarias del modulo de Evaluacion de Perfiles.",
    )

    # --------------------------------------------------------
    # 3.5.3.9 BITACORA DE CAMBIOS
    # --------------------------------------------------------
    doc.add_heading("3.5.3.9. Bitacora de Cambios al Sprint Backlog", level=2)

    cambios_rows = [
        ["25/10/2025", "Se agrego funcionalidad de recomendaciones automaticas", "El alcance inicial solo contemplaba evaluacion individual. Se agrego el flujo de recomendaciones automaticas contra todas las ofertas activas.", "Se agregaron 4 tareas adicionales (+48h estimadas)"],
        ["28/10/2025", "Se agrego verificacion de elegibilidad", "Usuarios con perfiles muy incompletos obtenian recomendaciones de baja calidad. Se decidio establecer umbral minimo del 70%.", "Se agrego componente EligibilityWarning y endpoint de verificacion (+10h)"],
        ["02/11/2025", "Se agrego sidebar colapsable en recomendaciones", "Se necesitaba mostrar estadisticas de recomendaciones y resumen del perfil sin ocupar todo el espacio principal.", "Se creo RecommendationsSidebar con estado expandido/colapsado (+18h)"],
        ["05/11/2025", "Se habilito evaluacion sin autenticacion", "Para demostraciones y pruebas, se permitio evaluar un CV sin estar logueado. El historial si requiere login.", "Se ajusto el guard de la ruta /evaluation a requiresAuth: false"],
        ["08/11/2025", "Se agrego cache de recomendaciones", "El recalculo de recomendaciones era lento con muchas ofertas. Se implemento cache en BD con opcion de recalculo forzado.", "Se agrego logica de cache y boton 'Actualizar' (+10h)"],
    ]

    add_formatted_table(doc,
        headers=["Fecha", "Cambio", "Motivo", "Impacto"],
        rows=cambios_rows,
        caption="Bitacora de cambios al Sprint Backlog - Sprint 3",
        nota="Los cambios principales se debieron a la ampliacion del alcance con recomendaciones automaticas y la necesidad de controles de elegibilidad para mantener la calidad de los resultados.",
    )

    # --------------------------------------------------------
    # GUARDAR
    # --------------------------------------------------------
    output_path = os.path.join(os.path.dirname(__file__), "sprint 3.docx")
    doc.save(output_path)
    print(f"Documento generado exitosamente: {output_path}")
    print(f"Total tablas: {tabla_counter['value']} (desde tabla 17)")
    print(f"Total figuras: {figura_counter['value']} (desde figura 19)")


if __name__ == "__main__":
    generate()
