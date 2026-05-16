import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_manual():
    doc = Document()

    # Define custom styles or just use built-in Heading 1, 2, Normal
    # Title Page
    doc.add_heading('MANUAL DE USUARIO', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    doc.add_heading('SISTEMA DE EVALUACION DE PERFILES PROFESIONALES APLICANDO PROCESAMIENTO DE LENGUAJE NATURAL Y MACHINE LEARNING', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n\n')
    doc.add_paragraph('VICERRECTORADO DE GRADO DE LA EMI').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC Placeholder
    doc.add_heading('ÍNDICE', 1)
    doc.add_paragraph('[Espacio reservado para Índice Automático]')
    doc.add_page_break()

    # CAPÍTULO 1
    doc.add_heading('CAPÍTULO 1', 1)
    doc.add_heading('INTRODUCCIÓN', 1)

    doc.add_heading('DESCRIPCIÓN DEL SISTEMA', 2)
    doc.add_paragraph('El Sistema de Evaluación de Perfiles Profesionales aplicando Procesamiento de Lenguaje Natural (NLP) y Machine Learning es una plataforma integral diseñada para facilitar la intermediación laboral. Conecta perfiles profesionales con convocatorias laborales mediante un análisis inteligente de currículums (CVs) y la correspondencia con los requerimientos de las instituciones.')

    doc.add_heading('ROLES DE USUARIO', 2)
    doc.add_paragraph('El sistema incorpora un mecanismo de control de acceso basado en roles para administrar las funcionalidades disponibles para cada usuario:')
    doc.add_paragraph('• Administrador: Posee acceso completo a todos los módulos del sistema (Gestión de Usuarios, Digitalización, Perfiles Institucionales, Convocatorias, Evaluación de Perfiles y Reportes).', style='List Bullet')
    doc.add_paragraph('• Operador: Posee acceso parcial, permitiendo gestionar usuarios, perfiles institucionales, convocatorias y reportes, pero sin acceso al módulo de Evaluación de Perfiles.', style='List Bullet')
    doc.add_paragraph('• Estudiante / Titulado: Posee acceso al módulo de Digitalización de Perfiles y a la Evaluación de Perfiles. (Nota: los Estudiantes tienen acceso enfocado a Pasantías, mientras que los Titulados a Convocatorias Laborales).', style='List Bullet')
    
    doc.add_page_break()

    # CAPÍTULO 2
    doc.add_heading('CAPÍTULO 2', 1)
    doc.add_heading('INSTALACIÓN Y CONFIGURACIÓN', 1)

    doc.add_heading('REQUISITOS PREVIOS', 2)
    doc.add_paragraph('Antes de la instalación, verifique que el entorno cumpla con las especificaciones mínimas de hardware y software.')
    
    doc.add_heading('Software', 3)
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Componente'
    hdr_cells[1].text = 'Tecnología'
    row1 = table.add_row().cells
    row1[0].text = 'Backend'
    row1[1].text = 'Python 3.9+, FastAPI, spaCy, Gemini API'
    row2 = table.add_row().cells
    row2[0].text = 'Base de Datos'
    row2[1].text = 'PostgreSQL (Supabase)'
    row3 = table.add_row().cells
    row3[0].text = 'Frontend'
    row3[1].text = 'Node.js 16+, Vue.js 3, Vite'

    doc.add_heading('INSTALACIÓN DEL BACKEND', 2)
    doc.add_paragraph('El backend está desarrollado con Python y FastAPI.')
    doc.add_paragraph('1. Navegue al directorio del backend e instale dependencias:')
    doc.add_paragraph('python -m venv venv\nvenv\\Scripts\\activate\npip install -r requirements.txt', style='Normal')
    doc.add_paragraph('2. Descargue el modelo de spaCy:')
    doc.add_paragraph('python -m spacy download es_core_news_md', style='Normal')
    doc.add_paragraph('3. Configure las variables de entorno para PostgreSQL (Supabase) y Gemini API en el archivo .env.')
    doc.add_paragraph('4. Inicie el servidor backend:')
    doc.add_paragraph('uvicorn app.main:app --reload --port 8000', style='Normal')

    doc.add_heading('INSTALACIÓN DEL FRONTEND', 2)
    doc.add_paragraph('El frontend fue desarrollado con Vue.js 3 y Vite.')
    doc.add_paragraph('1. Navegue al directorio del frontend e instale dependencias:')
    doc.add_paragraph('npm install', style='Normal')
    doc.add_paragraph('2. Ejecute el servidor de desarrollo:')
    doc.add_paragraph('npm run dev', style='Normal')
    
    doc.add_page_break()

    # CAPÍTULO 3
    doc.add_heading('CAPÍTULO 3', 1)
    doc.add_heading('MANUAL DE USO POR MÓDULOS', 1)

    doc.add_heading('ACCESO AL SISTEMA', 2)
    doc.add_paragraph('Para ingresar al sistema, acceda a la URL proporcionada por el servidor frontend. Deberá introducir sus credenciales (correo electrónico y contraseña). El sistema lo redirigirá al panel principal según su rol.')
    p_img = doc.add_paragraph('[Espacio para Captura de Pantalla: Pantalla de Login]')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('MÓDULO DE GESTIÓN DE USUARIOS', 2)
    doc.add_paragraph('Permite administrar la información de los usuarios registrados en el sistema.')
    doc.add_paragraph('• Lista de Usuarios: Visualización, búsqueda y filtrado de usuarios registrados.')
    doc.add_paragraph('• Nuevo Usuario: Formulario para registrar y asignar roles a nuevos usuarios.')
    p_img = doc.add_paragraph('[Espacio para Captura de Pantalla: Pantalla de Gestión de Usuarios]')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('MÓDULO DE DIGITALIZACIÓN DE PERFILES', 2)
    doc.add_paragraph('Este módulo se encarga del procesamiento de los CVs utilizando Procesamiento de Lenguaje Natural para extraer la información estructurada.')
    doc.add_paragraph('• Subir CV: Carga de documento PDF. El sistema extrae automáticamente datos personales y competencias.')
    doc.add_paragraph('• Mi Perfil Digitalizado: Visualización de los datos extraídos.')
    doc.add_paragraph('• Editar Perfil: Corrección y adición manual de información del perfil.')
    doc.add_paragraph('• Buscar Perfiles: (Para administradores/operadores) búsqueda de perfiles de estudiantes y titulados.')
    p_img = doc.add_paragraph('[Espacio para Captura de Pantalla: Pantalla de Digitalización de Perfiles]')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('MÓDULO DE PERFILES INSTITUCIONALES', 2)
    doc.add_paragraph('Permite gestionar la información de las instituciones u organizaciones que ofrecen pasantías y convocatorias laborales.')
    doc.add_paragraph('• Lista de Perfiles Institucionales: Registro de instituciones asociadas.')
    doc.add_paragraph('• Nuevo Perfil: Creación de información de una nueva institución.')
    p_img = doc.add_paragraph('[Espacio para Captura de Pantalla: Pantalla de Perfiles Institucionales]')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('MÓDULO DE GESTIÓN DE CONVOCATORIAS LABORALES', 2)
    doc.add_paragraph('Centraliza la creación y administración de oportunidades de empleo y pasantías.')
    doc.add_paragraph('• Lista de Convocatorias: Visualización de oportunidades activas e históricas.')
    doc.add_paragraph('• Nueva Convocatoria: Formulario donde se especifican los requerimientos técnicos y competencias esperadas.')
    p_img = doc.add_paragraph('[Espacio para Captura de Pantalla: Pantalla de Convocatorias]')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('MÓDULO DE EVALUACIÓN DE PERFILES', 2)
    doc.add_paragraph('El núcleo analítico del sistema. Realiza la evaluación de correspondencia (Match) entre el perfil del candidato y los requisitos de la convocatoria utilizando NLP y Machine Learning.')
    doc.add_paragraph('• Correspondencia entre Perfiles: Muestra el nivel de afinidad entre un candidato y ofertas.')
    doc.add_paragraph('• Historial de Postulaciones: Seguimiento de las ofertas a las que se ha aplicado.')
    doc.add_paragraph('• Evaluación de Candidatos: Herramienta de ranking para que los administradores identifiquen los mejores candidatos para una convocatoria específica.')
    p_img = doc.add_paragraph('[Espacio para Captura de Pantalla: Pantalla de Evaluación/Ranking]')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('MÓDULO DE INFORMES Y REPORTES', 2)
    doc.add_paragraph('Sección de analítica y estadísticas globales del sistema.')
    doc.add_paragraph('• Reportes de Usuarios y Convocatorias: Gráficos de cumplimiento, perfiles creados y tipos de instituciones.')
    p_img = doc.add_paragraph('[Espacio para Captura de Pantalla: Pantalla de Reportes]')
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(r'docs\Manual_de_Usuario_ICONDORIC.docx')

if __name__ == "__main__":
    create_manual()
