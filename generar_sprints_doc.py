"""
Generador de Documentación de Sprints - Trabajo de Grado
Autor: Ivan Condori Choquehuanca
Genera un archivo .docx con la documentación completa de los Sprints 1-3
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

AUTOR = "Ivan Condori Choquehuanca"
SISTEMA = "Sistema de Intermediación Laboral con PLN"

SPRINTS = {
    1: {"nombre": "Gestión de Usuarios", "inicio": "08/09/2025", "fin": "28/09/2025"},
    2: {"nombre": "Digitalización de Perfiles", "inicio": "29/09/2025", "fin": "19/10/2025"},
    3: {"nombre": "Evaluación de Perfiles", "inicio": "20/10/2025", "fin": "09/11/2025"},
}

# Contadores globales para figuras y tablas
tabla_counter = {"value": 0}
figura_counter = {"value": 0}


# ============================================================
# FUNCIONES AUXILIARES DE FORMATO
# ============================================================

def set_cell_shading(cell, color_hex):
    """Aplica sombreado a una celda de tabla"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(table):
    """Aplica bordes a toda la tabla"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_table(doc, headers, rows, sprint_num, caption, nota, fuente="Elaboración propia"):
    """Agrega una tabla formateada con caption, nota y fuente debajo."""
    tabla_counter["value"] += 1
    num = tabla_counter["value"]

    # Crear tabla
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_borders(table)

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            # Alternar color de fila
            if r_idx % 2 == 1:
                set_cell_shading(cell, "D6E4F0")

    # Caption debajo
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Tabla {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    # Nota
    nota_p = doc.add_paragraph()
    nota_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    # Fuente
    fuente_p = doc.add_paragraph()
    fuente_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()  # Espacio
    return num


def add_figure_placeholder(doc, descripcion_captura, caption, nota, fuente="Elaboración propia"):
    """Agrega un placeholder de figura con instrucciones de captura."""
    figura_counter["value"] += 1
    num = figura_counter["value"]

    # Placeholder box
    p_box = doc.add_paragraph()
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_box.add_run(f"[INSERTAR CAPTURA DE PANTALLA]\n{descripcion_captura}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 0, 0)
    run.italic = True

    # Caption debajo
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    # Nota
    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    # Fuente
    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_code_block(doc, code_text, caption, nota, fuente="Elaboración propia"):
    """Agrega un bloque de código con formato monoespaciado."""
    figura_counter["value"] += 1
    num = figura_counter["value"]

    # Código
    lines = code_text.strip().split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        # Set font for East Asian and complex scripts too
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Consolas" w:hAnsi="Consolas" w:cs="Consolas"/>')
        rPr.append(rFonts)

    # Caption
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = caption_p.add_run(f"Figura {num}. ")
    run_label.bold = True
    run_label.font.size = Pt(10)
    run_text = caption_p.add_run(caption)
    run_text.font.size = Pt(10)

    # Nota
    nota_p = doc.add_paragraph()
    run_nota_label = nota_p.add_run("Nota: ")
    run_nota_label.bold = True
    run_nota_label.italic = True
    run_nota_label.font.size = Pt(10)
    run_nota_text = nota_p.add_run(nota)
    run_nota_text.italic = True
    run_nota_text.font.size = Pt(10)

    # Fuente
    fuente_p = doc.add_paragraph()
    run_fuente_label = fuente_p.add_run("Fuente: ")
    run_fuente_label.bold = True
    run_fuente_label.italic = True
    run_fuente_label.font.size = Pt(10)
    run_fuente_text = fuente_p.add_run(fuente)
    run_fuente_text.italic = True
    run_fuente_text.font.size = Pt(10)

    doc.add_paragraph()
    return num


def add_heading_sprint(doc, sprint_num, section_name, level=2):
    """Agrega un encabezado de sección de sprint."""
    h = doc.add_heading(f"3.5.{sprint_num}. Sprint {sprint_num}: {SPRINTS[sprint_num]['nombre']}", level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)


def add_subsection(doc, sprint_num, sub_idx, title, level=2):
    """Agrega un subtítulo dentro de un sprint."""
    prefix = f"3.5.{sprint_num}.{sub_idx}"
    h = doc.add_heading(f"{prefix}. {title}", level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)


# ============================================================
# SPRINT 1: GESTIÓN DE USUARIOS
# ============================================================

def generar_sprint1(doc):
    sprint_num = 1
    add_heading_sprint(doc, sprint_num, "Gestión de Usuarios")

    doc.add_paragraph(
        f"El Sprint 1 se desarrolló del {SPRINTS[1]['inicio']} al {SPRINTS[1]['fin']} "
        "y se centró en la implementación del módulo de Gestión de Usuarios, que incluye "
        "el registro, autenticación, administración de cuentas y control de acceso basado en roles (RBAC). "
        "Este módulo constituye la base fundamental del sistema, ya que todos los demás módulos "
        "dependen de la identidad y autorización del usuario."
    )

    # --- 1. Sprint Backlog ---
    add_subsection(doc, 1, 1, "Sprint Backlog")
    doc.add_paragraph(
        "A continuación se presenta el Sprint Backlog con las historias de usuario "
        "y tareas planificadas para este sprint."
    )

    backlog_headers = ["ID", "Historia de Usuario", "Tarea", "Prioridad", "Estado", "Responsable", "Est. (h)"]
    backlog_rows = [
        ["HU-01", "Como usuario, quiero registrarme en el sistema para acceder a las funcionalidades", "Diseñar formulario de registro (Vue 3)", "Alta", "Completado", AUTOR, "8"],
        ["HU-01", "", "Implementar endpoint POST /api/auth/register", "Alta", "Completado", AUTOR, "12"],
        ["HU-01", "", "Implementar hash de contraseñas con bcrypt", "Alta", "Completado", AUTOR, "4"],
        ["HU-01", "", "Crear tabla 'usuarios' en Supabase (PostgreSQL)", "Alta", "Completado", AUTOR, "4"],
        ["HU-01", "", "Crear perfil profesional vacío al registrar usuario", "Alta", "Completado", AUTOR, "3"],
        ["HU-02", "Como usuario, quiero iniciar sesión para acceder a mi cuenta", "Diseñar formulario de login", "Alta", "Completado", AUTOR, "6"],
        ["HU-02", "", "Implementar endpoint POST /api/auth/login", "Alta", "Completado", AUTOR, "8"],
        ["HU-02", "", "Implementar generación y validación de JWT (HS256)", "Alta", "Completado", AUTOR, "10"],
        ["HU-02", "", "Implementar interceptor Axios para token Bearer", "Alta", "Completado", AUTOR, "4"],
        ["HU-02", "", "Implementar store de autenticación (Pinia)", "Alta", "Completado", AUTOR, "6"],
        ["HU-03", "Como admin, quiero gestionar usuarios del sistema", "Implementar listado paginado GET /api/users/", "Media", "Completado", AUTOR, "10"],
        ["HU-03", "", "Implementar búsqueda y filtrado por rol/email", "Media", "Completado", AUTOR, "6"],
        ["HU-03", "", "Implementar detalle de usuario GET /api/users/{id}", "Media", "Completado", AUTOR, "4"],
        ["HU-03", "", "Implementar edición de usuario PUT /api/users/{id}", "Media", "Completado", AUTOR, "6"],
        ["HU-03", "", "Implementar eliminación DELETE /api/users/{id}", "Media", "Completado", AUTOR, "6"],
        ["HU-03", "", "Diseñar vista administrativa de usuarios (Vue 3)", "Media", "Completado", AUTOR, "12"],
        ["HU-04", "Como usuario, quiero configurar mi cuenta", "Implementar GET /api/users/me/account", "Baja", "Completado", AUTOR, "4"],
        ["HU-04", "", "Implementar PUT /api/users/me (actualizar nombre/email)", "Baja", "Completado", AUTOR, "6"],
        ["HU-04", "", "Implementar PUT /api/users/me/password", "Baja", "Completado", AUTOR, "6"],
        ["HU-04", "", "Diseñar vista de configuración de cuenta", "Baja", "Completado", AUTOR, "8"],
        ["HU-05", "Como sistema, necesito control de acceso basado en roles", "Implementar middleware de verificación JWT", "Alta", "Completado", AUTOR, "8"],
        ["HU-05", "", "Implementar guards de rol (admin, operador)", "Alta", "Completado", AUTOR, "6"],
        ["HU-05", "", "Implementar guards de ruta en Vue Router", "Alta", "Completado", AUTOR, "6"],
        ["HU-05", "", "Manejar respuestas 401/403 en interceptor Axios", "Alta", "Completado", AUTOR, "4"],
    ]

    add_formatted_table(doc, backlog_headers, backlog_rows, 1,
                        "Sprint Backlog – Sprint 1: Gestión de Usuarios",
                        "Se detallan las historias de usuario y tareas técnicas planificadas para el módulo de gestión de usuarios, "
                        "incluyendo registro, autenticación, administración y control de acceso.")

    # --- 2. Bitácora de temas importantes ---
    add_subsection(doc, 1, 2, "Bitácora de Identificación de Temas Importantes")
    doc.add_paragraph(
        "Durante el desarrollo del Sprint 1, se identificaron los siguientes temas importantes "
        "que impactaron las decisiones de diseño e implementación."
    )

    bitacora_headers = ["Fecha", "Tema Identificado", "Impacto", "Decisión Tomada"]
    bitacora_rows = [
        ["08/09/2025", "Selección de método de autenticación", "Alto", "Se decidió usar JWT con algoritmo HS256 y expiración de 24 horas, almacenando el token en localStorage del navegador."],
        ["10/09/2025", "Esquema de hash de contraseñas", "Alto", "Se seleccionó bcrypt (passlib) por su resistencia a ataques de fuerza bruta y su factor de costo configurable."],
        ["12/09/2025", "Diseño de roles del sistema", "Alto", "Se definieron 4 roles: estudiante, titulado, administrador, operador. Los roles determinan permisos y contenido visible."],
        ["15/09/2025", "Base de datos: Supabase vs PostgreSQL local", "Alto", "Se eligió Supabase (PostgreSQL managed) por su API REST integrada, panel de administración y facilidad de despliegue."],
        ["18/09/2025", "Patrón de manejo de estado frontend", "Medio", "Se adoptó Pinia con Composition API (setup function) para los stores de autenticación."],
        ["20/09/2025", "Estrategia de manejo de errores HTTP", "Medio", "Se implementaron interceptores Axios que redirigen automáticamente: 401 → login, 403 → dashboard."],
        ["25/09/2025", "Creación automática de perfil vacío", "Medio", "Al registrar un usuario, se crea automáticamente un registro vacío en perfiles_profesionales para evitar errores en módulos posteriores."],
    ]

    add_formatted_table(doc, bitacora_headers, bitacora_rows, 1,
                        "Bitácora de temas importantes – Sprint 1",
                        "Registro de decisiones técnicas y de diseño relevantes durante el Sprint 1.")

    # --- 3. Actualización tabla de tareas ---
    add_subsection(doc, 1, 3, "Actualización de la Tabla de Tareas")
    doc.add_paragraph(
        "La siguiente tabla muestra el estado final de las tareas al cierre del Sprint 1, "
        "incluyendo las horas reales invertidas."
    )

    tareas_headers = ["Tarea", "Estado", "Horas Est.", "Horas Real", "Observaciones"]
    tareas_rows = [
        ["Crear tabla usuarios en Supabase", "Completado", "4", "3", "Configuración directa desde panel Supabase"],
        ["Endpoint registro POST /api/auth/register", "Completado", "12", "14", "Se añadió validación de email duplicado"],
        ["Endpoint login POST /api/auth/login", "Completado", "8", "7", "—"],
        ["Generación y validación JWT", "Completado", "10", "12", "Se requirió investigar python-jose"],
        ["Hash bcrypt con passlib", "Completado", "4", "3", "—"],
        ["Formulario registro Vue 3", "Completado", "8", "10", "Se añadió selección de rol con validación"],
        ["Formulario login Vue 3", "Completado", "6", "5", "—"],
        ["Store autenticación Pinia", "Completado", "6", "8", "Se implementó persistencia en localStorage"],
        ["Interceptor Axios (Bearer token)", "Completado", "4", "5", "Manejo especial de 401 y 403"],
        ["CRUD usuarios (admin)", "Completado", "22", "26", "Incluye paginación y búsqueda"],
        ["Vista admin de usuarios", "Completado", "12", "14", "Integración con paginación y filtros"],
        ["Endpoints cuenta personal", "Completado", "16", "15", "GET, PUT cuenta y cambio contraseña"],
        ["Vista configuración cuenta", "Completado", "8", "9", "—"],
        ["Guards de ruta Vue Router", "Completado", "6", "7", "Redirección según rol"],
        ["Dependencias FastAPI (RBAC)", "Completado", "14", "13", "verify_admin_role, verify_operator_access"],
    ]

    add_formatted_table(doc, tareas_headers, tareas_rows, 1,
                        "Actualización de tareas – Sprint 1",
                        "Estado final de las tareas con comparación entre horas estimadas y reales. "
                        "Total estimado: 140h, Total real: 151h. Desviación: +7.8%.")

    # --- 4. Diagrama de caso de uso expandido ---
    add_subsection(doc, 1, 4, "Diagrama de Caso de Uso Expandido del Módulo")
    doc.add_paragraph(
        "El diagrama de caso de uso expandido presenta las interacciones entre los actores "
        "(Usuario, Administrador, Sistema) y las funcionalidades del módulo de Gestión de Usuarios."
    )

    add_figure_placeholder(doc,
        "Tomar captura del diagrama de caso de uso del módulo de Gestión de Usuarios.\n"
        "Actores: Usuario (Estudiante/Titulado), Administrador, Sistema.\n"
        "Casos de uso: Registrarse, Iniciar Sesión, Cerrar Sesión, Ver Perfil de Cuenta,\n"
        "Editar Cuenta, Cambiar Contraseña, Listar Usuarios (admin), Editar Usuario (admin),\n"
        "Eliminar Usuario (admin), Validar Token JWT (sistema), Verificar Rol (sistema).",
        "Diagrama de caso de uso expandido – Módulo de Gestión de Usuarios",
        "El diagrama muestra las interacciones de los actores con el sistema para operaciones de registro, "
        "autenticación y administración de usuarios. Los casos de uso del administrador extienden la gestión "
        "básica con operaciones CRUD sobre todos los usuarios.")

    # Caso de uso expandido textual
    doc.add_paragraph(
        "A continuación se describe el caso de uso expandido más relevante del módulo:"
    )

    cu_headers = ["Campo", "Descripción"]
    cu_rows = [
        ["Caso de Uso", "CU-01: Registrar Usuario"],
        ["Actor Principal", "Usuario (Estudiante/Titulado)"],
        ["Precondiciones", "El usuario no tiene una cuenta registrada en el sistema"],
        ["Postcondiciones", "Se crea una cuenta con hash bcrypt, un perfil vacío y se retorna un JWT válido"],
        ["Flujo Principal", "1. El usuario accede al formulario de registro\n"
                           "2. Ingresa email, contraseña, nombre completo y selecciona rol\n"
                           "3. El sistema valida que el email no esté registrado\n"
                           "4. El sistema genera el hash bcrypt de la contraseña\n"
                           "5. Se inserta el registro en la tabla 'usuarios'\n"
                           "6. Se crea un perfil profesional vacío en 'perfiles_profesionales'\n"
                           "7. Se genera un JWT con user_id y rol\n"
                           "8. Se retorna el token y datos del usuario"],
        ["Flujo Alternativo", "3a. Si el email ya existe: retorna error 400 'Email ya registrado'\n"
                              "4a. Si el rol no es válido: retorna error 400 'Rol inválido'\n"
                              "5a. Si hay error de BD: retorna error 500"],
        ["Requisitos Especiales", "Contraseña mínima 6 caracteres. JWT expira en 24 horas. Algoritmo HS256."],
    ]

    add_formatted_table(doc, cu_headers, cu_rows, 1,
                        "Caso de uso expandido CU-01: Registrar Usuario",
                        "Descripción detallada del flujo de registro de usuario incluyendo la creación "
                        "automática de perfil profesional vacío.")

    # --- 5. Diseño de colección y documentos ---
    add_subsection(doc, 1, 5, "Diseño de la Colección y Documentos")
    doc.add_paragraph(
        "El módulo de Gestión de Usuarios utiliza la base de datos PostgreSQL hospedada en Supabase. "
        "A continuación se detalla el diseño de las tablas involucradas en este sprint."
    )

    # Tabla usuarios
    col_headers = ["Campo", "Tipo", "Restricciones", "Descripción"]
    col_rows = [
        ["id", "UUID", "PRIMARY KEY, auto-generado", "Identificador único del usuario"],
        ["email", "TEXT", "UNIQUE, NOT NULL", "Correo electrónico (usado para login)"],
        ["password_hash", "TEXT", "NOT NULL", "Hash bcrypt de la contraseña"],
        ["rol", "TEXT", "NOT NULL", "Rol: estudiante, titulado, administrador, operador"],
        ["nombre_completo", "TEXT", "NULLABLE", "Nombre completo del usuario"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de creación del registro"],
    ]

    add_formatted_table(doc, col_headers, col_rows, 1,
                        "Estructura de la tabla 'usuarios'",
                        "Tabla principal que almacena las credenciales y datos básicos de cada usuario del sistema. "
                        "El campo password_hash almacena el hash bcrypt, nunca la contraseña en texto plano.")

    # Tabla perfiles_profesionales (campos base)
    col_rows2 = [
        ["id", "UUID", "PRIMARY KEY, auto-generado", "Identificador único del perfil"],
        ["usuario_id", "UUID", "FOREIGN KEY → usuarios(id)", "Referencia al usuario propietario"],
        ["gemini_extraction", "JSONB", "DEFAULT '{}'", "Datos crudos extraídos del CV por Gemini"],
        ["hard_skills", "JSONB (array)", "DEFAULT '[]'", "Lista de habilidades técnicas normalizadas"],
        ["soft_skills", "JSONB (array)", "DEFAULT '[]'", "Lista de habilidades blandas normalizadas"],
        ["education_level", "TEXT", "NULLABLE", "Nivel educativo más alto detectado"],
        ["experience_years", "NUMERIC", "DEFAULT 0", "Años totales de experiencia laboral"],
        ["languages", "JSONB (array)", "DEFAULT '[]'", "Idiomas con nivel de dominio"],
        ["cv_filename", "TEXT", "NULLABLE", "Nombre del archivo CV subido"],
        ["cv_uploaded_at", "TIMESTAMPTZ", "NULLABLE", "Fecha de subida del CV"],
        ["is_complete", "BOOLEAN", "DEFAULT false", "Indica si el perfil supera el 70% de completitud"],
        ["completeness_score", "NUMERIC", "DEFAULT 0", "Score de completitud (0.0 a 1.0)"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de creación"],
        ["updated_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de última actualización"],
    ]

    add_formatted_table(doc, col_headers, col_rows2, 1,
                        "Estructura de la tabla 'perfiles_profesionales'",
                        "Tabla que almacena los perfiles profesionales de los usuarios. Utiliza campos JSONB "
                        "para flexibilidad en el almacenamiento de datos extraídos del CV. Esta tabla se crea vacía "
                        "durante el registro del usuario (Sprint 1) y se llena en el Sprint 2.")

    # --- 6. Mockups ---
    add_subsection(doc, 1, 6, "Diseño e Interfaz del Módulo (Mockups)")
    doc.add_paragraph(
        "A continuación se presentan las interfaces de usuario implementadas para el módulo "
        "de Gestión de Usuarios."
    )

    add_figure_placeholder(doc,
        "Tomar captura de la pantalla de REGISTRO (RegisterView.vue).\n"
        "Debe mostrar: campos email, contraseña, nombre completo, selector de rol\n"
        "(estudiante/titulado), botón 'Registrarse'.",
        "Interfaz de registro de usuario",
        "Formulario de registro con selección de rol. El usuario puede elegir entre 'Estudiante' y 'Titulado'. "
        "Los campos son validados antes de enviar la solicitud al backend.")

    add_figure_placeholder(doc,
        "Tomar captura de la pantalla de LOGIN (LoginView.vue).\n"
        "Debe mostrar: campos email y contraseña, botón 'Iniciar Sesión',\n"
        "enlace a registro.",
        "Interfaz de inicio de sesión",
        "Formulario de autenticación que envía credenciales al endpoint /api/auth/login y "
        "almacena el JWT retornado en localStorage.")

    add_figure_placeholder(doc,
        "Tomar captura de la vista ADMIN DE USUARIOS (UsersAdminView.vue).\n"
        "Debe mostrar: tabla de usuarios con columnas (nombre, email, rol, fecha),\n"
        "filtro por rol, barra de búsqueda, paginación.",
        "Vista administrativa de gestión de usuarios",
        "Panel administrativo que permite listar, buscar, filtrar por rol y paginar usuarios. "
        "Cada fila muestra el estado del perfil y permite acceder a acciones de edición y eliminación.")

    add_figure_placeholder(doc,
        "Tomar captura de la vista de CONFIGURACIÓN DE CUENTA.\n"
        "Debe mostrar: campos nombre, email (editables), sección cambiar contraseña.",
        "Interfaz de configuración de cuenta del usuario",
        "Vista que permite al usuario actualizar su información personal (nombre, email) y "
        "cambiar su contraseña. Requiere la contraseña actual para realizar el cambio.")

    # --- 7. Código fuente ---
    add_subsection(doc, 1, 7, "Código Fuente del Módulo (Incremento del Producto)")
    doc.add_paragraph(
        "A continuación se presentan los fragmentos de código fuente más relevantes "
        "del módulo de Gestión de Usuarios."
    )

    # auth.py - register
    code_register = '''@router.post("/register", response_model=Token)
async def register(user: UserRegister):
    # 1. Verificar email duplicado
    check = supabase.table("usuarios").select("id")\\
        .eq("email", user.email).execute()
    if check.data:
        raise HTTPException(status_code=400,
                          detail="Email ya registrado")

    # 2. Crear usuario con hash bcrypt
    hashed_password = get_password_hash(user.password)
    new_user_data = {
        "email": user.email,
        "password_hash": hashed_password,
        "rol": user.rol,
        "nombre_completo": user.nombre_completo
    }
    response = supabase.table("usuarios")\\
        .insert(new_user_data).execute()
    user_id = response.data[0]["id"]

    # 3. Crear perfil profesional vacío
    supabase.table("perfiles_profesionales")\\
        .insert({"usuario_id": user_id}).execute()

    # 4. Generar JWT
    access_token = create_access_token(user_id, user.rol)
    return {"access_token": access_token, ...}'''

    add_code_block(doc, code_register,
                   "Código fuente del endpoint de registro de usuario (auth.py)",
                   "Endpoint POST /api/auth/register que crea un nuevo usuario con hash bcrypt, "
                   "genera un perfil profesional vacío y retorna un JWT válido por 24 horas.")

    # security.py
    code_security = '''pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

def verify_password(plain_password: str, hashed: str) -> bool:
    return pwd_context.verify(plain_password, hashed)

def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)

def create_access_token(subject, role: str) -> str:
    expire = datetime.utcnow() + timedelta(
        minutes=settings.ACCESS_TOKEN_EXPIRE_MINUTES  # 1440 = 24h
    )
    to_encode = {"sub": str(subject), "role": role, "exp": expire}
    return jwt.encode(to_encode, settings.SECRET_KEY,
                     algorithm=settings.ALGORITHM)  # HS256'''

    add_code_block(doc, code_security,
                   "Módulo de seguridad: hash bcrypt y generación JWT (security.py)",
                   "Funciones de seguridad que gestionan el hash de contraseñas con bcrypt y la "
                   "generación de tokens JWT con algoritmo HS256 y expiración configurable.")

    # dependencies.py (RBAC)
    code_rbac = '''async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(bearer)
):
    token = credentials.credentials
    payload = jwt.decode(token, settings.SECRET_KEY,
                        algorithms=[settings.ALGORITHM])
    user_id = payload.get("sub")
    role = payload.get("role")
    return {"user_id": user_id, "role": role}

async def verify_admin_role(current_user = Depends(get_current_user)):
    if current_user["role"] != "administrador":
        raise HTTPException(status_code=403,
                          detail="Acceso denegado")
    return current_user

async def verify_operator_access(current_user = Depends(get_current_user)):
    if current_user["role"] not in ["operador", "administrador"]:
        raise HTTPException(status_code=403,
                          detail="Acceso denegado")
    return current_user'''

    add_code_block(doc, code_rbac,
                   "Sistema de control de acceso basado en roles (dependencies.py)",
                   "Funciones de dependencia de FastAPI que implementan RBAC: extracción del usuario desde "
                   "JWT y verificación de roles para proteger endpoints administrativos.")

    # --- 8. Pruebas unitarias ---
    add_subsection(doc, 1, 8, "Resultado de Pruebas Unitarias")
    doc.add_paragraph(
        "Se realizaron pruebas unitarias para validar el correcto funcionamiento "
        "de los endpoints de autenticación y gestión de usuarios."
    )

    pruebas_headers = ["ID Prueba", "Descripción", "Entrada", "Resultado Esperado", "Estado"]
    pruebas_rows = [
        ["PU-1.01", "Registro exitoso con datos válidos", "email, password, rol='estudiante', nombre", "Token JWT + user_id + status 200", "Aprobado"],
        ["PU-1.02", "Registro con email duplicado", "email ya existente", "Error 400: 'Email ya registrado'", "Aprobado"],
        ["PU-1.03", "Registro con rol inválido", "rol='superadmin'", "Error 400: 'Rol inválido'", "Aprobado"],
        ["PU-1.04", "Login exitoso", "email y password correctos", "Token JWT + datos usuario + status 200", "Aprobado"],
        ["PU-1.05", "Login con contraseña incorrecta", "email correcto, password incorrecto", "Error 400: 'Email o contraseña incorrectos'", "Aprobado"],
        ["PU-1.06", "Login con email inexistente", "email no registrado", "Error 400: 'Email o contraseña incorrectos'", "Aprobado"],
        ["PU-1.07", "Acceso a ruta protegida sin token", "Sin header Authorization", "Error 401: Unauthorized", "Aprobado"],
        ["PU-1.08", "Acceso admin con rol estudiante", "Token JWT con rol='estudiante'", "Error 403: Acceso denegado", "Aprobado"],
        ["PU-1.09", "Listar usuarios como admin", "GET /api/users/ con token admin", "Lista paginada + status 200", "Aprobado"],
        ["PU-1.10", "Eliminar usuario como admin", "DELETE /api/users/{id} con token admin", "Mensaje éxito + perfil eliminado", "Aprobado"],
        ["PU-1.11", "Actualizar cuenta propia", "PUT /api/users/me con nuevo nombre", "Datos actualizados + status 200", "Aprobado"],
        ["PU-1.12", "Cambiar contraseña con actual incorrecta", "old_password incorrecto", "Error 400: 'Contraseña actual incorrecta'", "Aprobado"],
    ]

    add_formatted_table(doc, pruebas_headers, pruebas_rows, 1,
                        "Resultados de pruebas unitarias – Sprint 1",
                        "Se ejecutaron 12 pruebas unitarias cubriendo los flujos principales y alternativos "
                        "del módulo de autenticación y gestión de usuarios. Todas las pruebas pasaron exitosamente.")

    add_figure_placeholder(doc,
        "Tomar captura de la terminal ejecutando las pruebas unitarias del Sprint 1.\n"
        "Puede ser con pytest mostrando todos los tests PASSED o con el script\n"
        "de pruebas mostrando resultados exitosos.",
        "Ejecución de pruebas unitarias del Sprint 1",
        "Captura de la terminal mostrando la ejecución exitosa de las 12 pruebas unitarias "
        "del módulo de Gestión de Usuarios.")

    # --- 9. Bitácora de cambios ---
    add_subsection(doc, 1, 9, "Bitácora de Cambios al Sprint Backlog")
    doc.add_paragraph(
        "Durante el Sprint 1 se realizaron los siguientes cambios al backlog original."
    )

    cambios_headers = ["Fecha", "Cambio", "Motivo", "Impacto"]
    cambios_rows = [
        ["12/09/2025", "Se añadió tarea: 'Crear perfil vacío al registrar'", "Se identificó que los módulos posteriores (Sprint 2) requieren un registro previo en perfiles_profesionales", "Bajo: +3 horas"],
        ["18/09/2025", "Se incrementó estimación de CRUD admin de 16h a 22h", "La paginación con Supabase requirió manejo especial de count y range", "Medio: +6 horas"],
        ["22/09/2025", "Se añadió tarea: 'Manejar 401/403 en interceptor Axios'", "Se detectó que las respuestas de error no redirigían correctamente al usuario", "Bajo: +4 horas"],
        ["26/09/2025", "Se añadió rol 'operador' además de los 3 iniciales", "Necesidad identificada de un rol intermedio entre admin y estudiante", "Medio: +4 horas en guards"],
    ]

    add_formatted_table(doc, cambios_headers, cambios_rows, 1,
                        "Bitácora de cambios al Sprint Backlog – Sprint 1",
                        "Cambios realizados al Sprint Backlog durante la ejecución del Sprint 1. "
                        "Los cambios se debieron principalmente a requisitos emergentes identificados durante la implementación.")

    doc.add_page_break()


# ============================================================
# SPRINT 2: DIGITALIZACIÓN DE PERFILES
# ============================================================

def generar_sprint2(doc):
    sprint_num = 2
    add_heading_sprint(doc, sprint_num, "Digitalización de Perfiles")

    doc.add_paragraph(
        f"El Sprint 2 se desarrolló del {SPRINTS[2]['inicio']} al {SPRINTS[2]['fin']} "
        "y se centró en la implementación del módulo de Digitalización de Perfiles, que permite "
        "a los usuarios subir su Curriculum Vitae en formato PDF, extraer automáticamente la información "
        "mediante Google Gemini (LLM), y almacenar los datos estructurados en su perfil profesional. "
        "Adicionalmente, se implementó la edición manual de campos y el cálculo de completitud del perfil."
    )

    # --- 1. Sprint Backlog ---
    add_subsection(doc, 2, 1, "Sprint Backlog")
    doc.add_paragraph(
        "A continuación se presenta el Sprint Backlog con las historias de usuario "
        "y tareas planificadas para el módulo de Digitalización de Perfiles."
    )

    backlog_headers = ["ID", "Historia de Usuario", "Tarea", "Prioridad", "Estado", "Responsable", "Est. (h)"]
    backlog_rows = [
        ["HU-06", "Como usuario, quiero subir mi CV en PDF para digitalizar mi perfil", "Diseñar interfaz de carga de CV (SubirCVView)", "Alta", "Completado", AUTOR, "10"],
        ["HU-06", "", "Implementar componente CVUploadSection", "Alta", "Completado", AUTOR, "8"],
        ["HU-06", "", "Implementar endpoint POST /api/profile/upload-cv", "Alta", "Completado", AUTOR, "12"],
        ["HU-06", "", "Integrar pdfplumber para extracción de texto del PDF", "Alta", "Completado", AUTOR, "6"],
        ["HU-06", "", "Integrar Google Gemini API para extracción estructurada", "Alta", "Completado", AUTOR, "16"],
        ["HU-06", "", "Implementar normalización de skills extraídos", "Alta", "Completado", AUTOR, "8"],
        ["HU-06", "", "Validar tipo y tamaño de archivo (solo PDF, max 10MB)", "Media", "Completado", AUTOR, "4"],
        ["HU-07", "Como usuario, quiero ver un resumen de los datos extraídos de mi CV", "Implementar componente DigitalizationSummary", "Alta", "Completado", AUTOR, "8"],
        ["HU-07", "", "Implementar CompetenciasCard (hard/soft skills)", "Media", "Completado", AUTOR, "6"],
        ["HU-07", "", "Implementar EducationExperienceGrid", "Media", "Completado", AUTOR, "6"],
        ["HU-08", "Como usuario, quiero editar mi perfil manualmente", "Implementar endpoint PUT /api/profile/me", "Media", "Completado", AUTOR, "8"],
        ["HU-08", "", "Diseñar ProfileEditor y EditProfileModal", "Media", "Completado", AUTOR, "10"],
        ["HU-08", "", "Implementar edición de skills, educación, experiencia, idiomas", "Media", "Completado", AUTOR, "8"],
        ["HU-09", "Como usuario, quiero ver el porcentaje de completitud de mi perfil", "Implementar endpoint GET /api/profile/completeness", "Media", "Completado", AUTOR, "6"],
        ["HU-09", "", "Implementar cálculo de completitud con pesos ponderados", "Media", "Completado", AUTOR, "8"],
        ["HU-09", "", "Implementar visualización con GaugeChart/ProgressBar", "Baja", "Completado", AUTOR, "6"],
        ["HU-10", "Como usuario, quiero eliminar mi perfil para subirlo de nuevo", "Implementar endpoint DELETE /api/profile/me", "Baja", "Completado", AUTOR, "4"],
        ["HU-10", "", "Implementar DeleteConfirmModal", "Baja", "Completado", AUTOR, "4"],
        ["HU-11", "Como sistema, necesito un servicio de perfiles", "Implementar ProfileService (Singleton)", "Alta", "Completado", AUTOR, "16"],
        ["HU-11", "", "Implementar extracción de nivel educativo más alto", "Alta", "Completado", AUTOR, "6"],
        ["HU-11", "", "Implementar cálculo de años de experiencia (regex)", "Alta", "Completado", AUTOR, "8"],
    ]

    add_formatted_table(doc, backlog_headers, backlog_rows, 2,
                        "Sprint Backlog – Sprint 2: Digitalización de Perfiles",
                        "Se detallan las historias de usuario y tareas técnicas para la digitalización de CV, "
                        "extracción con Gemini, edición manual y cálculo de completitud del perfil.")

    # --- 2. Bitácora de temas importantes ---
    add_subsection(doc, 2, 2, "Bitácora de Identificación de Temas Importantes")

    bitacora_headers = ["Fecha", "Tema Identificado", "Impacto", "Decisión Tomada"]
    bitacora_rows = [
        ["29/09/2025", "Selección de LLM para extracción de CV", "Alto", "Se eligió Google Gemini por su capacidad de procesamiento de documentos PDF y generación de JSON estructurado."],
        ["01/10/2025", "Formato de almacenamiento de datos extraídos", "Alto", "Se adoptó JSONB en PostgreSQL para almacenar los datos del CV, permitiendo consultas flexibles sin migraciones de esquema."],
        ["03/10/2025", "Normalización de habilidades técnicas", "Medio", "Se implementó normalización a minúsculas y mapeo de sinónimos (ej: 'JS' → 'javascript', 'C#' → 'c#') para mejorar el matching."],
        ["06/10/2025", "Cálculo de años de experiencia", "Medio", "Se implementaron múltiples patrones regex para parsear duraciones: 'X años', 'X meses', 'YYYY-YYYY', 'YYYY-presente'."],
        ["09/10/2025", "Umbral de completitud del perfil", "Alto", "Se definió 70% como umbral mínimo para acceder a recomendaciones. Los pesos se distribuyeron: hard_skills 25%, experience 25%, education 20%, soft_skills 15%, languages 10%, summary 5%."],
        ["14/10/2025", "Patrón Singleton para ProfileService", "Medio", "Se adoptó Singleton para evitar múltiples instancias del servicio y mantener consistencia en caché de datos."],
        ["17/10/2025", "Límite de tamaño de CV", "Bajo", "Se estableció 10MB como límite máximo para archivos PDF, considerando CVs con imágenes y gráficos."],
    ]

    add_formatted_table(doc, bitacora_headers, bitacora_rows, 2,
                        "Bitácora de temas importantes – Sprint 2",
                        "Registro de decisiones técnicas durante el desarrollo del módulo de digitalización, "
                        "incluyendo la integración con Gemini y el diseño del sistema de completitud.")

    # --- 3. Actualización tabla de tareas ---
    add_subsection(doc, 2, 3, "Actualización de la Tabla de Tareas")

    tareas_headers = ["Tarea", "Estado", "Horas Est.", "Horas Real", "Observaciones"]
    tareas_rows = [
        ["Interfaz SubirCVView + CVUploadSection", "Completado", "18", "20", "Se añadió modal de progreso de carga"],
        ["Endpoint POST /api/profile/upload-cv", "Completado", "12", "14", "Validación de PDF y tamaño"],
        ["Integración pdfplumber", "Completado", "6", "5", "Extracción directa de texto"],
        ["Integración Google Gemini API", "Completado", "16", "22", "Prompt engineering iterativo para JSON correcto"],
        ["Normalización de skills", "Completado", "8", "9", "Mapeo de sinónimos y normalización"],
        ["DigitalizationSummary + cards", "Completado", "14", "13", "Componentes de visualización"],
        ["Endpoint PUT /api/profile/me", "Completado", "8", "8", "Edición manual de campos"],
        ["ProfileEditor + EditProfileModal", "Completado", "10", "12", "Formularios dinámicos"],
        ["Cálculo de completitud", "Completado", "14", "16", "Pesos ponderados + recomendaciones"],
        ["ProfileService (Singleton)", "Completado", "16", "18", "Incluye regex para experiencia"],
        ["Endpoint DELETE /api/profile/me", "Completado", "4", "3", "Reset a vacío, no eliminación"],
        ["DeleteConfirmModal", "Completado", "4", "3", "Modal de confirmación"],
        ["Endpoint GET /api/profile/preview", "Completado", "4", "5", "Vista previa para recomendaciones"],
    ]

    add_formatted_table(doc, tareas_headers, tareas_rows, 2,
                        "Actualización de tareas – Sprint 2",
                        "Total estimado: 134h, Total real: 148h. Desviación: +10.4%. "
                        "La mayor desviación se dio en la integración con Gemini API debido al prompt engineering iterativo.")

    # --- 4. Diagrama de caso de uso expandido ---
    add_subsection(doc, 2, 4, "Diagrama de Caso de Uso Expandido del Módulo")
    doc.add_paragraph(
        "El diagrama de caso de uso expandido presenta las interacciones entre los actores "
        "y las funcionalidades del módulo de Digitalización de Perfiles."
    )

    add_figure_placeholder(doc,
        "Tomar captura del diagrama de caso de uso del módulo de Digitalización de Perfiles.\n"
        "Actores: Usuario Autenticado, Google Gemini (sistema externo), Sistema.\n"
        "Casos de uso: Subir CV (PDF), Extraer Datos con Gemini, Ver Resumen de Extracción,\n"
        "Editar Perfil Manualmente, Consultar Completitud, Eliminar Datos de Perfil,\n"
        "Calcular Completitud (sistema), Normalizar Skills (sistema).",
        "Diagrama de caso de uso expandido – Módulo de Digitalización de Perfiles",
        "El diagrama muestra el flujo desde la carga del PDF hasta la persistencia de datos, "
        "incluyendo la interacción con Google Gemini como sistema externo para la extracción estructurada.")

    cu_headers = ["Campo", "Descripción"]
    cu_rows = [
        ["Caso de Uso", "CU-06: Subir CV y Digitalizar Perfil"],
        ["Actor Principal", "Usuario Autenticado (Estudiante/Titulado)"],
        ["Actor Secundario", "Google Gemini API (LLM externo)"],
        ["Precondiciones", "El usuario está autenticado y tiene un perfil (vacío o con datos previos)"],
        ["Postcondiciones", "El perfil se actualiza con hard_skills, soft_skills, education, experience, languages extraídos del CV. Se recalcula el completeness_score."],
        ["Flujo Principal", "1. El usuario accede a la vista 'Subir CV'\n"
                           "2. Selecciona un archivo PDF (máx 10MB)\n"
                           "3. El sistema lee el PDF con pdfplumber\n"
                           "4. El texto se envía a Gemini API con prompt estructurado\n"
                           "5. Gemini retorna JSON: {hard_skills, soft_skills, education, experience, personal_info}\n"
                           "6. El sistema normaliza los skills (minúsculas, sinónimos)\n"
                           "7. Se extrae el nivel educativo más alto y los años de experiencia\n"
                           "8. Se calcula el completeness_score con pesos ponderados\n"
                           "9. Se actualiza perfiles_profesionales en la BD\n"
                           "10. Se muestra el resumen de extracción al usuario"],
        ["Flujo Alternativo", "2a. Archivo no es PDF → Error 400\n"
                              "2b. Archivo supera 10MB → Error 400\n"
                              "4a. Gemini no disponible → Error 500\n"
                              "5a. JSON de Gemini incompleto → Se procesan campos disponibles"],
    ]

    add_formatted_table(doc, cu_headers, cu_rows, 2,
                        "Caso de uso expandido CU-06: Subir CV y Digitalizar Perfil",
                        "Descripción del flujo completo desde la carga del archivo PDF hasta "
                        "la actualización del perfil con datos extraídos por Gemini.")

    # --- 5. Diseño de colección ---
    add_subsection(doc, 2, 5, "Diseño de la Colección y Documentos")
    doc.add_paragraph(
        "El Sprint 2 utiliza la tabla 'perfiles_profesionales' creada en el Sprint 1, "
        "pero ahora pobla los campos JSONB con datos reales del CV. "
        "A continuación se muestra la estructura del campo JSONB 'gemini_extraction'."
    )

    json_headers = ["Campo", "Tipo", "Ejemplo", "Descripción"]
    json_rows = [
        ["hard_skills", "Array[string]", '["Python", "React", "SQL"]', "Habilidades técnicas detectadas"],
        ["soft_skills", "Array[string]", '["Liderazgo", "Trabajo en equipo"]', "Habilidades blandas detectadas"],
        ["education", "Array[object]", '{degree, institution, year}', "Lista de formación académica"],
        ["education[].degree", "string", '"Ingeniería de Sistemas"', "Título obtenido"],
        ["education[].institution", "string", '"Escuela Militar de Ingeniería"', "Institución educativa"],
        ["education[].year", "string", '"2021"', "Año de graduación"],
        ["experience", "Array[object]", '{role, company, duration, description}', "Experiencia laboral"],
        ["experience[].role", "string", '"Full Stack Developer"', "Cargo desempeñado"],
        ["experience[].company", "string", '"TechBolivia"', "Empresa"],
        ["experience[].duration", "string", '"2 años"', "Duración en el cargo"],
        ["personal_info", "object", '{summary, languages}', "Información personal"],
        ["personal_info.summary", "string", '"Ingeniero con 3 años..."', "Resumen profesional"],
        ["personal_info.languages", "Array[string]", '["Español (Nativo)", "Inglés (B2)"]', "Idiomas con nivel"],
    ]

    add_formatted_table(doc, json_headers, json_rows, 2,
                        "Estructura del campo JSONB 'gemini_extraction'",
                        "Estructura de los datos extraídos por Google Gemini del CV del usuario. "
                        "Esta estructura se almacena completa en el campo gemini_extraction "
                        "y los datos normalizados se copian a los campos individuales del perfil.")

    # Pesos de completitud
    pesos_headers = ["Dimensión", "Peso (%)", "Condición para Completar"]
    pesos_rows = [
        ["Habilidades técnicas (hard_skills)", "25%", "Al menos 3 habilidades registradas"],
        ["Experiencia laboral (experience)", "25%", "Al menos 1 experiencia o experience_years > 0"],
        ["Formación académica (education)", "20%", "Nivel educativo detectado o al menos 1 registro"],
        ["Habilidades blandas (soft_skills)", "15%", "Al menos 2 habilidades registradas"],
        ["Idiomas (languages)", "10%", "Al menos 1 idioma registrado"],
        ["Resumen profesional (summary)", "5%", "Texto con más de 20 caracteres"],
    ]

    add_formatted_table(doc, pesos_headers, pesos_rows, 2,
                        "Pesos ponderados para el cálculo de completitud del perfil",
                        "El perfil se considera completo cuando el score supera 0.70 (70%). "
                        "Los pesos fueron diseñados priorizando habilidades técnicas y experiencia, "
                        "que son las dimensiones más relevantes para el matching laboral.")

    # --- 6. Mockups ---
    add_subsection(doc, 2, 6, "Diseño e Interfaz del Módulo (Mockups)")

    add_figure_placeholder(doc,
        "Tomar captura de la vista SUBIR CV (SubirCVView.vue).\n"
        "Debe mostrar: zona de arrastrar/soltar PDF, botón 'Seleccionar archivo',\n"
        "indicador de progreso, restricciones (PDF, max 10MB).",
        "Interfaz de carga de CV en formato PDF",
        "Vista que permite al usuario subir su CV arrastrando el archivo o seleccionándolo. "
        "Muestra restricciones de formato y tamaño máximo permitido.")

    add_figure_placeholder(doc,
        "Tomar captura del RESUMEN DE EXTRACCIÓN (DigitalizationSummary).\n"
        "Debe mostrar: lista de hard skills y soft skills detectados,\n"
        "educación, experiencia laboral, idiomas extraídos del CV.",
        "Resumen de datos extraídos del CV por Gemini",
        "Componente que muestra al usuario los datos estructurados que fueron extraídos "
        "automáticamente de su CV mediante Google Gemini.")

    add_figure_placeholder(doc,
        "Tomar captura de MI PERFIL (MiPerfilView.vue).\n"
        "Debe mostrar: tarjetas de competencias, educación, experiencia,\n"
        "barra/gauge de completitud, botón editar.",
        "Vista 'Mi Perfil' con datos del perfil profesional",
        "Vista completa del perfil profesional del usuario mostrando todas las dimensiones "
        "extraídas del CV, el score de completitud y opciones de edición manual.")

    add_figure_placeholder(doc,
        "Tomar captura del MODAL DE EDICIÓN (EditProfileModal).\n"
        "Debe mostrar: campos editables para skills, educación, idiomas,\n"
        "botones guardar/cancelar.",
        "Modal de edición manual del perfil profesional",
        "Modal que permite al usuario modificar manualmente los datos de su perfil, "
        "incluyendo agregar/eliminar habilidades, actualizar nivel educativo e idiomas.")

    # --- 7. Código fuente ---
    add_subsection(doc, 2, 7, "Código Fuente del Módulo (Incremento del Producto)")

    code_upload = '''@router.post("/upload-cv", response_model=CVUploadResponse)
async def upload_cv(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    # Validar tipo de archivo
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(400, "Solo se permiten archivos PDF")

    # Validar tamaño (max 10MB)
    contents = await file.read()
    if len(contents) > 10 * 1024 * 1024:
        raise HTTPException(400, "Archivo excede 10MB")

    # Codificar a base64 y extraer con Gemini
    pdf_base64 = base64.b64encode(contents).decode('utf-8')
    gemini_output = ml_service.extract_cv_with_gemini(pdf_base64)

    # Actualizar perfil con datos extraídos
    profile = profile_service.update_profile_from_cv(
        current_user['user_id'], gemini_output,
        cv_filename=file.filename
    )
    return CVUploadResponse(message="CV procesado exitosamente",
                           perfil=profile, ...)'''

    add_code_block(doc, code_upload,
                   "Endpoint de carga y procesamiento de CV (profile.py)",
                   "Endpoint POST /api/profile/upload-cv que recibe un archivo PDF, lo valida, "
                   "lo codifica en base64, lo envía a Gemini para extracción y actualiza el perfil.")

    code_profile_service = '''class ProfileService:
    COMPLETENESS_WEIGHTS = {
        'hard_skills': 0.25, 'soft_skills': 0.15,
        'education': 0.20, 'experience': 0.25,
        'languages': 0.10, 'summary': 0.05
    }

    def update_profile_from_cv(self, user_id, gemini_output, cv_filename=None):
        hard_skills = gemini_output.get('hard_skills', [])
        soft_skills = gemini_output.get('soft_skills', [])
        education_level = self._extract_highest_education(
            gemini_output.get('education', []))
        experience_years = self._calculate_experience_years(
            gemini_output.get('experience', []))
        languages = gemini_output.get('personal_info', {}).get('languages', [])

        update_data = {
            'gemini_extraction': gemini_output,
            'hard_skills': hard_skills,
            'soft_skills': soft_skills,
            'education_level': education_level,
            'experience_years': experience_years,
            'languages': languages,
            'cv_filename': cv_filename,
            'cv_uploaded_at': datetime.utcnow().isoformat()
        }
        completeness = self.calculate_completeness(update_data)
        update_data['completeness_score'] = completeness['score']
        update_data['is_complete'] = completeness['is_complete']

        response = supabase.table("perfiles_profesionales")\\
            .update(update_data).eq("usuario_id", user_id).execute()
        return response.data[0]'''

    add_code_block(doc, code_profile_service,
                   "ProfileService: actualización de perfil desde CV (profile_service.py)",
                   "Método principal del servicio de perfiles que procesa la salida de Gemini, "
                   "extrae el nivel educativo, calcula años de experiencia y actualiza el perfil en BD.")

    code_completeness = '''def calculate_completeness(self, profile):
    score = 0.0
    missing_fields = []

    # Hard skills (25%) - mínimo 3
    hard_skills = profile.get('hard_skills', [])
    if hard_skills and len(hard_skills) >= 3:
        score += self.COMPLETENESS_WEIGHTS['hard_skills']
    else:
        missing_fields.append('hard_skills')

    # Education (20%) - nivel detectado
    if profile.get('education_level') or \\
       len(profile.get('gemini_extraction', {}).get('education', [])) >= 1:
        score += self.COMPLETENESS_WEIGHTS['education']

    # Experience (25%) - al menos 1 experiencia
    if profile.get('experience_years', 0) > 0:
        score += self.COMPLETENESS_WEIGHTS['experience']

    # ... soft_skills (15%), languages (10%), summary (5%)
    is_complete = score >= 0.70  # Umbral 70%
    return {'score': round(score, 2), 'is_complete': is_complete,
            'missing_fields': missing_fields}'''

    add_code_block(doc, code_completeness,
                   "Cálculo de completitud del perfil con pesos ponderados",
                   "Método que evalúa la completitud del perfil verificando cada dimensión contra "
                   "un mínimo requerido y aplicando pesos ponderados. El umbral de 70% habilita el acceso a recomendaciones.")

    # --- 8. Pruebas unitarias ---
    add_subsection(doc, 2, 8, "Resultado de Pruebas Unitarias")

    pruebas_headers = ["ID Prueba", "Descripción", "Entrada", "Resultado Esperado", "Estado"]
    pruebas_rows = [
        ["PU-2.01", "Subida exitosa de CV PDF", "Archivo PDF válido < 10MB", "Perfil actualizado + resumen de extracción", "Aprobado"],
        ["PU-2.02", "Rechazo de archivo no PDF", "Archivo .docx", "Error 400: 'Solo se permiten archivos PDF'", "Aprobado"],
        ["PU-2.03", "Rechazo de archivo > 10MB", "PDF de 15MB", "Error 400: 'Archivo excede 10MB'", "Aprobado"],
        ["PU-2.04", "Extracción de hard_skills con Gemini", "CV con skills técnicos", "Array de skills normalizados", "Aprobado"],
        ["PU-2.05", "Extracción de educación y nivel", "CV con título universitario", "education_level = 'Ingeniería'", "Aprobado"],
        ["PU-2.06", "Cálculo de años de experiencia (patrón 'X años')", "'2 años', '1 año'", "experience_years = 3.0", "Aprobado"],
        ["PU-2.07", "Cálculo de años de experiencia (patrón YYYY-YYYY)", "'2020-2023'", "experience_years = 3.0", "Aprobado"],
        ["PU-2.08", "Completitud perfil completo (>70%)", "Perfil con todos los campos", "is_complete = True, score > 0.70", "Aprobado"],
        ["PU-2.09", "Completitud perfil vacío", "Perfil sin datos", "is_complete = False, score = 0.0", "Aprobado"],
        ["PU-2.10", "Edición manual de hard_skills", "PUT con nuevas skills", "Skills actualizadas + recálculo completitud", "Aprobado"],
        ["PU-2.11", "Eliminación de perfil (reset)", "DELETE /api/profile/me", "Perfil reseteado a valores vacíos", "Aprobado"],
        ["PU-2.12", "Vista previa de perfil para ML", "GET /api/profile/preview", "Formato compatible con sistema ML", "Aprobado"],
    ]

    add_formatted_table(doc, pruebas_headers, pruebas_rows, 2,
                        "Resultados de pruebas unitarias – Sprint 2",
                        "Se ejecutaron 12 pruebas unitarias cubriendo carga de CV, extracción con Gemini, "
                        "cálculo de completitud y edición manual. Todas las pruebas pasaron exitosamente.")

    add_figure_placeholder(doc,
        "Tomar captura de la terminal ejecutando las pruebas del Sprint 2.\n"
        "Puede ser pytest o el script test_feature_engineering.py mostrando PASS.",
        "Ejecución de pruebas unitarias del Sprint 2",
        "Captura de la terminal mostrando la ejecución exitosa de las pruebas unitarias "
        "del módulo de Digitalización de Perfiles.")

    # --- 9. Bitácora de cambios ---
    add_subsection(doc, 2, 9, "Bitácora de Cambios al Sprint Backlog")

    cambios_headers = ["Fecha", "Cambio", "Motivo", "Impacto"]
    cambios_rows = [
        ["01/10/2025", "Se aumentó estimación de integración Gemini de 12h a 16h", "El prompt engineering requirió múltiples iteraciones para obtener JSON consistente", "Medio: +4h que luego subieron a +6h reales"],
        ["06/10/2025", "Se añadió tarea: regex para cálculo de experiencia", "Gemini retorna duraciones en formatos variados que requieren parsing", "Bajo: +8h"],
        ["10/10/2025", "Se añadió endpoint GET /api/profile/preview", "El Sprint 3 necesita el perfil en formato compatible con ML", "Bajo: +4h"],
        ["15/10/2025", "Se rediseñó DigitalizationSummary como múltiples componentes", "El componente original era demasiado grande; se separó en CompetenciasCard + EducationExperienceGrid", "Bajo: mismas horas, mejor organización"],
    ]

    add_formatted_table(doc, cambios_headers, cambios_rows, 2,
                        "Bitácora de cambios al Sprint Backlog – Sprint 2",
                        "Los cambios se debieron principalmente a la complejidad de la integración con Gemini "
                        "y a la preparación de datos para el Sprint 3 (Evaluación).")

    doc.add_page_break()


# ============================================================
# SPRINT 3: EVALUACIÓN DE PERFILES
# ============================================================

def generar_sprint3(doc):
    sprint_num = 3
    add_heading_sprint(doc, sprint_num, "Evaluación de Perfiles")

    doc.add_paragraph(
        f"El Sprint 3 se desarrolló del {SPRINTS[3]['inicio']} al {SPRINTS[3]['fin']} "
        "y se centró en la implementación del módulo de Evaluación de Perfiles, que permite evaluar "
        "un CV contra perfiles institucionales utilizando un pipeline de Machine Learning. El sistema "
        "extrae un vector de 18 dimensiones a partir del CV y la configuración institucional, y utiliza "
        "un modelo Ridge Regression para predecir el score de compatibilidad. Los resultados incluyen "
        "clasificación (Excelente/Bueno/Regular/Deficiente), fortalezas y debilidades principales."
    )

    # --- 1. Sprint Backlog ---
    add_subsection(doc, 3, 1, "Sprint Backlog")

    backlog_headers = ["ID", "Historia de Usuario", "Tarea", "Prioridad", "Estado", "Responsable", "Est. (h)"]
    backlog_rows = [
        ["HU-12", "Como usuario, quiero evaluar mi CV contra un perfil institucional", "Diseñar EvaluationView (interfaz de evaluación)", "Alta", "Completado", AUTOR, "10"],
        ["HU-12", "", "Implementar endpoint POST /api/ml/evaluate-cv", "Alta", "Completado", AUTOR, "16"],
        ["HU-12", "", "Implementar MLIntegrationService (Singleton)", "Alta", "Completado", AUTOR, "20"],
        ["HU-12", "", "Integrar modelo Ridge Regression (.joblib)", "Alta", "Completado", AUTOR, "12"],
        ["HU-13", "Como sistema, necesito extraer features del CV para ML", "Implementar FeatureExtractor (orquestador)", "Alta", "Completado", AUTOR, "14"],
        ["HU-13", "", "Implementar HardSkillsScorer (Jaccard + TF-IDF)", "Alta", "Completado", AUTOR, "10"],
        ["HU-13", "", "Implementar SoftSkillsScorer (categorización + matching)", "Alta", "Completado", AUTOR, "8"],
        ["HU-13", "", "Implementar EducationScorer (escala 1-5)", "Media", "Completado", AUTOR, "6"],
        ["HU-13", "", "Implementar ExperienceScorer (años + clasificación)", "Media", "Completado", AUTOR, "8"],
        ["HU-13", "", "Implementar LanguagesScorer (matching + nivel)", "Media", "Completado", AUTOR, "6"],
        ["HU-14", "Como usuario, quiero ver los resultados detallados de la evaluación", "Implementar componente EvaluationResult", "Alta", "Completado", AUTOR, "10"],
        ["HU-14", "", "Implementar ScoreChart (visualización de scores)", "Media", "Completado", AUTOR, "8"],
        ["HU-14", "", "Implementar clasificación (Excelente/Bueno/Regular/Deficiente)", "Media", "Completado", AUTOR, "4"],
        ["HU-14", "", "Calcular y mostrar top fortalezas y debilidades", "Media", "Completado", AUTOR, "6"],
        ["HU-15", "Como usuario, quiero ver mi historial de evaluaciones", "Implementar endpoint GET /api/ml/user-evaluations", "Baja", "Completado", AUTOR, "6"],
        ["HU-15", "", "Diseñar HistoryView para evaluaciones", "Baja", "Completado", AUTOR, "8"],
        ["HU-15", "", "Guardar evaluaciones en tabla 'evaluaciones'", "Media", "Completado", AUTOR, "6"],
        ["HU-16", "Como sistema, necesito un modelo ML entrenado", "Entrenar modelo Ridge con datos sintéticos", "Alta", "Completado", AUTOR, "16"],
        ["HU-16", "", "Implementar InstitutionalMatchModel", "Alta", "Completado", AUTOR, "10"],
        ["HU-16", "", "Serializar modelo con joblib (ridge_v1.joblib)", "Alta", "Completado", AUTOR, "4"],
        ["HU-16", "", "Implementar endpoint GET /api/ml/model-info", "Baja", "Completado", AUTOR, "4"],
    ]

    add_formatted_table(doc, backlog_headers, backlog_rows, 3,
                        "Sprint Backlog – Sprint 3: Evaluación de Perfiles",
                        "Se detallan las historias de usuario y tareas técnicas para el pipeline de evaluación ML, "
                        "incluyendo feature engineering, modelo Ridge, clasificación y visualización de resultados.")

    # --- 2. Bitácora de temas importantes ---
    add_subsection(doc, 3, 2, "Bitácora de Identificación de Temas Importantes")

    bitacora_headers = ["Fecha", "Tema Identificado", "Impacto", "Decisión Tomada"]
    bitacora_rows = [
        ["20/10/2025", "Selección del modelo de ML", "Alto", "Se eligió Ridge Regression por su interpretabilidad (coeficientes explicables), regularización L2, y buen rendimiento con datos de pocas dimensiones (18 features)."],
        ["22/10/2025", "Dimensionalidad del vector de features", "Alto", "Se definieron 18 features: 5 scores base del CV, 5 pesos institucionales, 5 interacciones (score×peso), 3 features de contexto (experiencia total, requerida, delta)."],
        ["24/10/2025", "Datos de entrenamiento", "Alto", "Se generaron datos sintéticos basados en combinaciones de perfiles CV y configuraciones institucionales, dado que no se disponía de datos históricos reales."],
        ["27/10/2025", "Método de scoring de hard skills", "Medio", "Se combinó similitud Jaccard (exact match) con TF-IDF (similitud semántica) para mejorar la cobertura de matching de habilidades técnicas."],
        ["30/10/2025", "Escala de educación", "Medio", "Se definió escala 1-5: bachillerato/técnico (1-2), licenciatura/ingeniería (3), maestría (4), doctorado (5). Permite comparación numérica con requisitos institucionales."],
        ["02/11/2025", "Umbrales de clasificación", "Medio", "Se definieron: Excelente >0.80, Bueno 0.60-0.80, Regular 0.40-0.60, Deficiente <0.40. Configurables por perfil institucional."],
        ["05/11/2025", "Almacenamiento de evaluaciones", "Medio", "Las evaluaciones se guardan opcionalmente (solo si el usuario está autenticado) en la tabla 'evaluaciones' para historial."],
    ]

    add_formatted_table(doc, bitacora_headers, bitacora_rows, 3,
                        "Bitácora de temas importantes – Sprint 3",
                        "Registro de decisiones técnicas del pipeline de ML, incluyendo selección del modelo, "
                        "diseño del vector de features y umbrales de clasificación.")

    # --- 3. Actualización tabla de tareas ---
    add_subsection(doc, 3, 3, "Actualización de la Tabla de Tareas")

    tareas_headers = ["Tarea", "Estado", "Horas Est.", "Horas Real", "Observaciones"]
    tareas_rows = [
        ["EvaluationView (interfaz)", "Completado", "10", "12", "Incluye selección de perfil institucional"],
        ["Endpoint POST /api/ml/evaluate-cv", "Completado", "16", "18", "Integra todo el pipeline ML"],
        ["MLIntegrationService (Singleton)", "Completado", "20", "24", "Caché de perfiles + gestión de modelo"],
        ["Modelo Ridge (entrenamiento)", "Completado", "16", "20", "Generación de datos sintéticos + cross-validation"],
        ["InstitutionalMatchModel", "Completado", "10", "10", "Wrapper sobre Ridge de sklearn"],
        ["FeatureExtractor (orquestador)", "Completado", "14", "14", "Coordina los 5 scorers"],
        ["HardSkillsScorer", "Completado", "10", "12", "Jaccard + TF-IDF scoring"],
        ["SoftSkillsScorer", "Completado", "8", "7", "Categorización + matching"],
        ["EducationScorer", "Completado", "6", "5", "Escala numérica 1-5"],
        ["ExperienceScorer", "Completado", "8", "9", "Regex + clasificación"],
        ["LanguagesScorer", "Completado", "6", "5", "Matching de idiomas"],
        ["EvaluationResult + ScoreChart", "Completado", "18", "20", "Visualización de resultados"],
        ["Historial de evaluaciones", "Completado", "14", "13", "CRUD + vista de historial"],
        ["Serialización modelo (.joblib)", "Completado", "4", "3", "joblib dump/load"],
        ["Endpoint GET /api/ml/model-info", "Completado", "4", "3", "Información del modelo"],
    ]

    add_formatted_table(doc, tareas_headers, tareas_rows, 3,
                        "Actualización de tareas – Sprint 3",
                        "Total estimado: 164h, Total real: 175h. Desviación: +6.7%. "
                        "La mayor desviación se dio en el entrenamiento del modelo ML y la integración del servicio.")

    # --- 4. Diagrama de caso de uso expandido ---
    add_subsection(doc, 3, 4, "Diagrama de Caso de Uso Expandido del Módulo")

    add_figure_placeholder(doc,
        "Tomar captura del diagrama de caso de uso del módulo de Evaluación de Perfiles.\n"
        "Actores: Usuario (Autenticado/Anónimo), Google Gemini, Modelo ML Ridge, Sistema.\n"
        "Casos de uso: Evaluar CV contra Perfil Institucional, Ver Resultado de Evaluación,\n"
        "Ver Historial de Evaluaciones, Extraer Features del CV (sistema),\n"
        "Predecir Match Score (sistema), Clasificar Candidato (sistema),\n"
        "Consultar Info del Modelo.",
        "Diagrama de caso de uso expandido – Módulo de Evaluación de Perfiles",
        "El diagrama muestra el flujo de evaluación que involucra la extracción con Gemini, "
        "el feature engineering con 5 scorers especializados y la predicción con el modelo Ridge.")

    cu_headers = ["Campo", "Descripción"]
    cu_rows = [
        ["Caso de Uso", "CU-12: Evaluar CV contra Perfil Institucional"],
        ["Actor Principal", "Usuario (puede ser autenticado o anónimo)"],
        ["Actores Secundarios", "Google Gemini API, Modelo Ridge Regression"],
        ["Precondiciones", "El modelo ML está cargado (ridge_v1.joblib). Existe al menos un perfil institucional activo."],
        ["Postcondiciones", "Se retorna match_score, clasificación, top fortalezas/debilidades. Si el usuario está autenticado, se guarda en la tabla evaluaciones."],
        ["Flujo Principal", "1. El usuario sube un CV en PDF (base64) y selecciona un perfil institucional\n"
                           "2. El sistema extrae el CV con Gemini → JSON estructurado\n"
                           "3. Se carga la configuración del perfil institucional (requisitos + pesos)\n"
                           "4. FeatureExtractor orquesta los 5 scorers:\n"
                           "   4a. HardSkillsScorer: Jaccard + TF-IDF\n"
                           "   4b. SoftSkillsScorer: categorización\n"
                           "   4c. EducationScorer: escala 1-5\n"
                           "   4d. ExperienceScorer: años + clasificación\n"
                           "   4e. LanguagesScorer: matching\n"
                           "5. Se construye el vector de 18 features\n"
                           "6. El modelo Ridge predice match_score [0-1]\n"
                           "7. Se clasifica: Excelente/Bueno/Regular/Deficiente\n"
                           "8. Se calculan top 3 fortalezas y debilidades\n"
                           "9. Se retorna resultado completo al frontend"],
        ["Flujo Alternativo", "1a. CV inválido → Error 400\n"
                              "3a. Perfil institucional no encontrado → Error 404\n"
                              "6a. Modelo no cargado → Error 503"],
    ]

    add_formatted_table(doc, cu_headers, cu_rows, 3,
                        "Caso de uso expandido CU-12: Evaluar CV contra Perfil Institucional",
                        "Descripción del flujo completo de evaluación ML que transforma un CV en PDF "
                        "en un score de compatibilidad con explicación de fortalezas y debilidades.")

    # --- 5. Diseño de colección ---
    add_subsection(doc, 3, 5, "Diseño de la Colección y Documentos")
    doc.add_paragraph(
        "El Sprint 3 introduce dos nuevas tablas para almacenar los perfiles institucionales "
        "y las evaluaciones realizadas."
    )

    # Tabla evaluaciones
    col_headers = ["Campo", "Tipo", "Restricciones", "Descripción"]
    col_rows = [
        ["id", "UUID", "PRIMARY KEY", "Identificador único de la evaluación"],
        ["user_id", "UUID", "FK → usuarios(id), NULLABLE", "Usuario que realizó la evaluación (null si anónimo)"],
        ["institutional_profile_id", "UUID", "FK → perfiles_institucionales(id)", "Perfil institucional evaluado"],
        ["match_score", "NUMERIC", "NOT NULL", "Score de compatibilidad (0.0 a 1.0)"],
        ["classification", "TEXT", "NOT NULL", "Clasificación: Excellent, Good, Fair, Poor"],
        ["cv_scores", "JSONB", "NOT NULL", "Scores individuales por dimensión"],
        ["top_strengths", "JSONB", "NOT NULL", "Top 3 fortalezas del candidato"],
        ["top_weaknesses", "JSONB", "NOT NULL", "Top 3 debilidades del candidato"],
        ["gemini_extraction", "JSONB", "NOT NULL", "Datos crudos extraídos del CV"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de la evaluación"],
    ]

    add_formatted_table(doc, col_headers, col_rows, 3,
                        "Estructura de la tabla 'evaluaciones'",
                        "Tabla que almacena los resultados de cada evaluación de CV contra un perfil institucional. "
                        "Incluye el score total, scores por dimensión, clasificación y fortalezas/debilidades.")

    # Tabla perfiles_institucionales
    col_rows2 = [
        ["id", "UUID", "PRIMARY KEY", "Identificador único del perfil"],
        ["institution_name", "TEXT", "NOT NULL", "Nombre de la institución/empresa"],
        ["sector", "TEXT", "NOT NULL", "Sector: tecnología, finanzas, salud, etc."],
        ["weights", "JSONB", "NOT NULL", "Pesos por dimensión (deben sumar 1.0)"],
        ["requirements", "JSONB", "NOT NULL", "Requisitos: skills, educación, experiencia, idiomas"],
        ["thresholds", "JSONB", "DEFAULT '{\"apto\": 0.70, \"considerado\": 0.50}'", "Umbrales de clasificación"],
        ["is_active", "BOOLEAN", "DEFAULT true", "Si el perfil está activo"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Fecha de creación"],
    ]

    add_formatted_table(doc, col_headers, col_rows2, 3,
                        "Estructura de la tabla 'perfiles_institucionales'",
                        "Tabla que define los perfiles de evaluación con pesos y requisitos específicos por "
                        "institución. Cada perfil configura qué tan importante es cada dimensión para el matching.")

    # Vector de features
    doc.add_paragraph(
        "El vector de features extraído por el FeatureExtractor consta de 18 dimensiones:"
    )

    feat_headers = ["Índice", "Feature", "Tipo", "Descripción"]
    feat_rows = [
        ["0", "hard_skills_score", "CV Score [0-1]", "Similitud Jaccard + TF-IDF de habilidades técnicas"],
        ["1", "soft_skills_score", "CV Score [0-1]", "Score de matching de habilidades blandas"],
        ["2", "experience_score", "CV Score [0-1]", "Score de experiencia (años vs requeridos)"],
        ["3", "education_score", "CV Score [0-1]", "Score de nivel educativo (escala 1-5)"],
        ["4", "languages_score", "CV Score [0-1]", "Score de matching de idiomas"],
        ["5-9", "weight_*", "Peso inst. [0-1]", "Pesos institucionales (hard, soft, exp, edu, lang)"],
        ["10-14", "interaction_*", "Score × Peso", "Productos cruzados (score × peso por dimensión)"],
        ["15", "total_experience_years", "Contexto", "Años totales de experiencia del candidato"],
        ["16", "min_required_years", "Contexto", "Años mínimos requeridos por la institución"],
        ["17", "experience_delta", "Contexto", "Diferencia: total_years - min_required_years"],
    ]

    add_formatted_table(doc, feat_headers, feat_rows, 3,
                        "Vector de 18 features para el modelo Ridge Regression",
                        "El vector combina scores del CV, pesos institucionales, interacciones y features de "
                        "contexto para alimentar el modelo de predicción. La dimensionalidad fue diseñada "
                        "para balance entre expresividad e interpretabilidad.")

    # --- 6. Mockups ---
    add_subsection(doc, 3, 6, "Diseño e Interfaz del Módulo (Mockups)")

    add_figure_placeholder(doc,
        "Tomar captura de la vista EVALUACIÓN (EvaluationView.vue).\n"
        "Debe mostrar: zona de carga de CV, selector de perfil institucional,\n"
        "botón 'Evaluar'.",
        "Interfaz de evaluación de CV",
        "Vista que permite al usuario subir un CV en PDF y seleccionar un perfil institucional "
        "para obtener una evaluación de compatibilidad basada en ML.")

    add_figure_placeholder(doc,
        "Tomar captura del RESULTADO DE EVALUACIÓN (EvaluationResult).\n"
        "Debe mostrar: score general (gauge o porcentaje), clasificación\n"
        "(Excelente/Bueno/Regular/Deficiente), gráfico de scores por dimensión,\n"
        "lista de fortalezas, lista de debilidades.",
        "Resultado detallado de evaluación de CV",
        "Vista de resultados que presenta el match score total, la clasificación del candidato, "
        "un desglose visual por dimensión (radar/barras) y las principales fortalezas y debilidades.")

    add_figure_placeholder(doc,
        "Tomar captura del HISTORIAL DE EVALUACIONES (HistoryView.vue).\n"
        "Debe mostrar: lista de evaluaciones previas con fecha, perfil evaluado,\n"
        "score obtenido, clasificación.",
        "Historial de evaluaciones del usuario",
        "Vista que lista las evaluaciones previas del usuario con fecha, perfil institucional evaluado, "
        "score de compatibilidad y clasificación obtenida.")

    # --- 7. Código fuente ---
    add_subsection(doc, 3, 7, "Código Fuente del Módulo (Incremento del Producto)")

    code_evaluate = '''@router.post("/evaluate-cv", response_model=CVEvaluationResponse)
async def evaluate_cv(request: CVEvaluationRequest,
    current_user = Depends(get_current_user_optional),
    ml_service = Depends(verify_ml_model_loaded)):

    # 1. Extraer CV con Gemini
    gemini_output = ml_service.extract_cv_with_gemini(request.cv_file)

    # 2. Cargar perfil institucional
    profile = ml_service.load_institutional_profile(
        request.institutional_profile_id)

    # 3. Evaluar CV (Feature Engineering + Ridge)
    evaluation = ml_service.evaluate_cv(gemini_output, profile)

    # 4. Guardar si usuario autenticado
    evaluation_id = None
    if current_user:
        evaluation_id = ml_service.save_evaluation(
            user_id=current_user['user_id'],
            profile_id=request.institutional_profile_id,
            evaluation_result=evaluation,
            gemini_extraction=gemini_output)

    return CVEvaluationResponse(
        match_score=evaluation['match_score'],
        classification=evaluation['classification'],
        cv_scores=evaluation['cv_scores'],
        top_strengths=evaluation['top_strengths'],
        top_weaknesses=evaluation['top_weaknesses'], ...)'''

    add_code_block(doc, code_evaluate,
                   "Endpoint de evaluación de CV con ML (ml_predictions.py)",
                   "Endpoint POST /api/ml/evaluate-cv que orquesta el pipeline completo: "
                   "extracción con Gemini, feature engineering, predicción Ridge y almacenamiento de resultados.")

    code_extractor = '''def extract_features(gemini_output, institutional_config):
    # PASO 1: Extraer datos del CV
    cv_hard_skills = extract_hard_skills_from_gemini(gemini_output)
    cv_soft_skills = extract_soft_skills_from_gemini(gemini_output)
    cv_education = extract_education_from_gemini(gemini_output)
    cv_experience = extract_experience_from_gemini(gemini_output)
    cv_languages = extract_languages_from_gemini(gemini_output)

    # PASO 2: Extraer requisitos institucionales
    requirements = institutional_config.get('requirements', {})
    weights = institutional_config.get('weights', {})

    # PASO 3: Calcular scores individuales [0-1]
    hard_result = calculate_hard_skills_score(cv_hard_skills, ...)
    soft_result = calculate_soft_skills_score(cv_soft_skills, ...)
    edu_result = calculate_education_score(cv_education, ...)
    exp_result = calculate_experience_score_from_cv(cv_experience, ...)
    lang_result = calculate_languages_score_from_cv(cv_languages, ...)

    # PASO 4: Construir vector de 18 features
    feature_vector = [
        hard_result['score'], soft_result['score'],  # 0-1
        exp_result['score'], edu_result['score'],    # 2-3
        lang_result['score'],                         # 4
        *weights.values(),                            # 5-9
        *(s*w for s,w in zip(scores, weights)),      # 10-14
        exp_result['total_years'],                    # 15
        min_experience_years,                         # 16
        exp_result['total_years'] - min_years         # 17
    ]
    return {'feature_vector': feature_vector,
            'cv_scores': cv_scores, ...}'''

    add_code_block(doc, code_extractor,
                   "Feature Extractor: construcción del vector de 18 features (feature_extractor.py)",
                   "Función principal del módulo de feature engineering que orquesta los 5 scorers "
                   "especializados y construye el vector numérico de 18 dimensiones para el modelo Ridge.")

    code_hard = '''def calculate_hard_skills_score(cv_skills, required_skills,
                                preferred_skills):
    # Normalizar (minúsculas, sinónimos)
    cv_norm = {normalize(s) for s in cv_skills}
    req_norm = {normalize(s) for s in required_skills}

    # Jaccard: intersección / unión
    if req_norm:
        matched = cv_norm & req_norm
        required_ratio = len(matched) / len(req_norm)
    else:
        required_ratio = 1.0

    # TF-IDF para skills preferidos
    pref_norm = {normalize(s) for s in preferred_skills}
    pref_matched = cv_norm & pref_norm
    preferred_ratio = len(pref_matched) / max(len(pref_norm), 1)

    # Score final: 70% required + 30% preferred
    score = 0.7 * required_ratio + 0.3 * preferred_ratio
    return {'score': min(score, 1.0), 'matched_required': matched,
            'missing_required': req_norm - cv_norm, ...}'''

    add_code_block(doc, code_hard,
                   "Scorer de habilidades técnicas con Jaccard y TF-IDF (hard_skills_scorer.py)",
                   "Función que calcula el score de habilidades técnicas combinando similitud Jaccard "
                   "para skills requeridos (70%) y TF-IDF para skills preferidos (30%).")

    # --- 8. Pruebas unitarias ---
    add_subsection(doc, 3, 8, "Resultado de Pruebas Unitarias")

    pruebas_headers = ["ID Prueba", "Descripción", "Entrada", "Resultado Esperado", "Estado"]
    pruebas_rows = [
        ["PU-3.01", "Hard Skills Scorer: matching parcial", "CV: [Python, React, SQL], Req: [Python, Java, SQL]", "Score entre 0.5-0.8, matched: [Python, SQL]", "Aprobado"],
        ["PU-3.02", "Hard Skills Scorer: matching total", "CV contiene todos los required", "Score >= 0.7", "Aprobado"],
        ["PU-3.03", "Soft Skills Scorer: categorización", "CV: [Liderazgo, Comunicación], Req: [Trabajo en equipo]", "Score > 0, categorías correctas", "Aprobado"],
        ["PU-3.04", "Education Scorer: nivel superior", "CV: Ingeniería, Req: Licenciatura", "Score = 1.0, meets_requirement = True", "Aprobado"],
        ["PU-3.05", "Education Scorer: nivel inferior", "CV: Técnico, Req: Maestría", "Score < 1.0, meets_requirement = False", "Aprobado"],
        ["PU-3.06", "Experience Scorer: supera mínimo", "CV: 3 años, Req: 1.5 años", "Score alto, meets_minimum = True", "Aprobado"],
        ["PU-3.07", "Languages Scorer: matching", "CV: [Español, Inglés B2], Req: [Inglés]", "Score > 0, matched: [Inglés]", "Aprobado"],
        ["PU-3.08", "Feature Extractor: vector completo", "CV ejemplo + config institucional", "Vector de 18 features numéricos", "Aprobado"],
        ["PU-3.09", "Feature Extractor: score ponderado", "Features extraídos", "Score final entre 0 y 1", "Aprobado"],
        ["PU-3.10", "Clasificación: Excelente (>0.8)", "Score = 0.85", "Clasificación = 'APTO'", "Aprobado"],
        ["PU-3.11", "Clasificación: Deficiente (<0.4)", "Score = 0.25", "Clasificación = 'NO_APTO'", "Aprobado"],
        ["PU-3.12", "Endpoint evaluate-cv completo", "PDF base64 + profile_id", "match_score, classification, cv_scores", "Aprobado"],
        ["PU-3.13", "Modelo Ridge: predicción", "Vector 18 features", "Predicción [0-1] numérica", "Aprobado"],
        ["PU-3.14", "Historial de evaluaciones", "GET /api/ml/user-evaluations", "Lista de evaluaciones previas", "Aprobado"],
    ]

    add_formatted_table(doc, pruebas_headers, pruebas_rows, 3,
                        "Resultados de pruebas unitarias – Sprint 3",
                        "Se ejecutaron 14 pruebas unitarias cubriendo los 5 scorers individuales, el feature extractor, "
                        "la clasificación, el endpoint completo y el modelo Ridge. Todas las pruebas pasaron exitosamente.")

    add_figure_placeholder(doc,
        "Tomar captura de la terminal ejecutando test_feature_engineering.py.\n"
        "El script muestra los resultados de cada scorer y el resultado final.\n"
        "Comando: python backend/tests/test_feature_engineering.py\n"
        "Debe verse: 'TODAS LAS PRUEBAS PASARON EXITOSAMENTE'",
        "Ejecución de pruebas unitarias del Sprint 3 (Feature Engineering)",
        "Captura de la terminal mostrando la ejecución exitosa del script de pruebas "
        "test_feature_engineering.py con resultados de los 5 scorers y el feature extractor completo.")

    # --- 9. Bitácora de cambios ---
    add_subsection(doc, 3, 9, "Bitácora de Cambios al Sprint Backlog")

    cambios_headers = ["Fecha", "Cambio", "Motivo", "Impacto"]
    cambios_rows = [
        ["22/10/2025", "Se añadió tarea: generación de datos sintéticos", "No se disponía de datos históricos reales para entrenar el modelo Ridge", "Alto: +12h de trabajo en generación y validación"],
        ["27/10/2025", "Se aumentó estimación de HardSkillsScorer de 8h a 10h", "La implementación de TF-IDF requirió investigación adicional sobre sklearn", "Bajo: +2h"],
        ["30/10/2025", "Se añadió evaluación sin autenticación", "Se identificó que la evaluación ad-hoc debe estar disponible para usuarios no registrados como demo", "Medio: +4h (get_current_user_optional)"],
        ["05/11/2025", "Se añadió caché de perfiles institucionales (5 min)", "Múltiples evaluaciones consecutivas consultaban la BD innecesariamente", "Bajo: +3h mejora de rendimiento"],
        ["07/11/2025", "Se incrementó de 12 a 14 pruebas unitarias", "Se añadieron pruebas para clasificación y modelo Ridge que no estaban inicialmente planificadas", "Bajo: +4h"],
    ]

    add_formatted_table(doc, cambios_headers, cambios_rows, 3,
                        "Bitácora de cambios al Sprint Backlog – Sprint 3",
                        "Los cambios principales se debieron a la necesidad de datos sintéticos para entrenamiento "
                        "y a la habilitación de evaluación sin autenticación como funcionalidad de demostración.")


# ============================================================
# GENERADOR PRINCIPAL
# ============================================================

def generar_documento():
    doc = Document()

    # Configurar estilos base
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.54)

    # Título principal
    title = doc.add_heading("3.5. Desarrollo de Sprints", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    doc.add_paragraph(
        "En esta sección se documenta el desarrollo de los sprints del proyecto, "
        "siguiendo la metodología ágil Scrum adaptada para un equipo de un solo desarrollador. "
        "Cada sprint tiene una duración de 3 semanas e incluye: Sprint Backlog, bitácora de temas "
        "importantes, actualización de tareas, diagrama de caso de uso expandido, diseño de base de "
        "datos, diseño de interfaces, código fuente relevante, resultados de pruebas unitarias "
        "y bitácora de cambios."
    )
    doc.add_paragraph(
        f"Desarrollador: {AUTOR}"
    )
    doc.add_paragraph(
        f"Período total: {SPRINTS[1]['inicio']} – {SPRINTS[3]['fin']}"
    )

    doc.add_page_break()

    # Generar cada sprint
    generar_sprint1(doc)
    generar_sprint2(doc)
    generar_sprint3(doc)

    # Guardar
    output_path = os.path.join(os.path.dirname(__file__), "Documentacion_Sprints_1_3.docx")
    doc.save(output_path)
    print(f"Documento generado exitosamente: {output_path}")
    print(f"Total tablas: {tabla_counter['value']}")
    print(f"Total figuras: {figura_counter['value']}")
    return output_path


if __name__ == "__main__":
    generar_documento()
